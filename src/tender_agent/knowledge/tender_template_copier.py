"""从本次上传的招标书中复制附件/格式范本到投标书。

★ V3 修复(2026-05-11):
   修复前面所有版本的问题,兼容两类招标书:

   类型 A(规范型):附件用 Heading 样式
     如海洋馆:'一、投标函及投标函附录' Heading 2

   类型 B(自由型):附件用 Normal/List Paragraph 样式
     如嘉诚国际:'投 标 函' List Paragraph,'(必填)' Normal

   关键改动:
   1. _detect_heading_level 严格化:
      数字编号 '1.' '2.' 不再被识别为 heading(除非 Heading 样式)
      防止投标函正文的列表"1、2、3、4、5、6、7、"被误判
      中文编号 '一、二、' 也要求短(<25)+ 不以标点结尾

   2. _collect_heading_candidates 加"模板词上下文兜底":
      短段落(<15字)+ 模板词后缀(函/书/表/...)+ 后续段含"必填/格式自拟/盖章"
      → 识别为 level=2 范本标题
      支持嘉诚国际"投 标 函" 这种非 Heading 样式但实际是范本的情况

   3. _find_attachment_title_match 评分增加 (3-level)*1500:
      优先匹配父级标题(如"一、投标函及投标函附录"而不是"(一)投标函")

   4. _is_child_template_heading 兼容全/半角括号 + 中文/阿拉伯数字

   5. _clean_title 剥离"一、二、" 中文编号

   6. _extract_template_tokens 加"附录" 作为模板词后缀

   7. ★ 新增:节点 token 与候选 token 做"双向包含匹配"
      解决"授权委托书" 节点 vs "三、法定代表人投标授权委托书" 候选 的匹配
"""
from copy import deepcopy
from difflib import SequenceMatcher
from io import BytesIO
import re
from pathlib import Path
from typing import Any, Dict, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from tender_agent.knowledge.docx_ooxml import (
    append_body_block as _append_body_block,
    fit_copied_tender_template_tables_to_page,
    remove_element,
    strip_page_break_marks as _strip_page_break_marks,
    strip_section_properties,
)
from tender_agent.parsing.attachment_refs import (
    attachment_reference as _attachment_reference,
    is_bare_attachment_label as _is_bare_attachment_label,
    strip_leading_attachment_label,
)


def _resolve_tender_template_block_span(
    node: Dict[str, Any],
    source_doc,
    *,
    blocks=None,
    candidates=None,
) -> Optional[Dict[str, Any]]:
    """Resolve the raw DOCX block span used by the renderer's title fallback."""
    blocks = list(blocks) if blocks is not None else list(source_doc.element.body.iterchildren())
    candidates = list(candidates) if candidates is not None else _collect_heading_candidates(source_doc, blocks)
    match = _find_attachment_title_match(node.get("name", ""), source_doc, blocks)
    if match is None:
        match = _best_template_match(node.get("name", ""), candidates)
    if match is None:
        return None

    start = match["block_idx"]
    end = _find_section_end(candidates, match, len(blocks))
    if _should_copy_representative_authorization_group(node.get("name", "")):
        end = _extend_authorization_template_end(candidates, match, end, len(blocks))
    end = _clip_end_at_cover_boundary(source_doc, blocks, start, end)
    paragraph_by_element = {paragraph._element: paragraph for paragraph in source_doc.paragraphs}
    narrowed_span = _narrow_generic_parent_template_span(
        blocks,
        paragraph_by_element,
        str(node.get("name", "")),
        match,
        start,
        end,
    )
    if narrowed_span is not None:
        start, end = narrowed_span
    elif not _should_preserve_internal_template_headings(str(node.get("name", ""))):
        end = _clip_template_span_at_internal_sibling_heading(blocks, paragraph_by_element, start, end)
    skip_indexes = _leading_duplicate_title_indexes(
        blocks,
        paragraph_by_element,
        start,
        min(end, start + 3),
        str(node.get("name", "")),
    )
    return {
        "blocks": blocks,
        "start": start,
        "end": end,
        "skip_indexes": skip_indexes,
        "source_anchor": str(match.get("title") or node.get("name") or ""),
    }


def _resolved_span_to_anchors(resolved: Dict[str, Any], source_doc) -> Optional[Dict[str, str]]:
    """Convert a raw OOXML span into the parser's non-empty block anchors."""
    paragraph_by_element = {paragraph._element: paragraph for paragraph in source_doc.paragraphs}
    anchor_index = -1
    included: list[int] = []
    for raw_index, block in enumerate(resolved["blocks"]):
        is_anchor_block = block.tag.endswith("}tbl")
        if block.tag.endswith("}p"):
            paragraph = paragraph_by_element.get(block)
            is_anchor_block = paragraph is not None and bool(paragraph.text.strip())
        if not is_anchor_block:
            continue
        anchor_index += 1
        if resolved["start"] <= raw_index < resolved["end"] and raw_index not in resolved["skip_indexes"]:
            included.append(anchor_index)
    if not included:
        return None
    return {
        "anchor_start": f"p{included[0]}",
        "anchor_end": f"p{included[-1]}",
        "source_anchor": str(resolved["source_anchor"]),
        "copy_method": "renderer_title_match",
    }


def resolve_tender_template_spans_by_nodes(
    nodes: list[Dict[str, Any]],
    tender_doc_path: str,
) -> Dict[str, Dict[str, str]]:
    """Resolve multiple outline nodes while opening and indexing the source DOCX only once."""
    source_path = Path(tender_doc_path)
    if not source_path.exists() or source_path.suffix.lower() != ".docx":
        return {}
    source_doc = Document(str(source_path))
    blocks = list(source_doc.element.body.iterchildren())
    candidates = _collect_heading_candidates(source_doc, blocks)
    spans: Dict[str, Dict[str, str]] = {}
    for node in nodes or []:
        node_id = str(node.get("id") or node.get("node_id") or "").strip()
        if not node_id:
            continue
        resolved = _resolve_tender_template_block_span(
            node,
            source_doc,
            blocks=blocks,
            candidates=candidates,
        )
        if resolved is None:
            continue
        anchors = _resolved_span_to_anchors(resolved, source_doc)
        if anchors:
            spans[node_id] = anchors
    return spans


def resolve_tender_template_span_by_node(node: Dict[str, Any], tender_doc_path: str) -> Optional[Dict[str, str]]:
    """Expose the renderer's resolved span as parser-style non-empty block anchors."""
    source_path = Path(tender_doc_path)
    if not source_path.exists() or source_path.suffix.lower() != ".docx":
        return None
    start_anchor = _anchor_block_number(node.get("anchor_start"))
    end_anchor = _anchor_block_number(node.get("anchor_end"))
    copy_method = str(node.get("copy_method") or "").strip()
    if (
        copy_method != "renderer_title_match"
        and start_anchor is not None
        and end_anchor is not None
        and end_anchor >= start_anchor
    ):
        return {
            "anchor_start": f"p{start_anchor}",
            "anchor_end": f"p{end_anchor}",
            "source_anchor": str(node.get("source_anchor") or node.get("name") or ""),
            "copy_method": str(node.get("copy_method") or "located_section"),
        }

    source_doc = Document(str(source_path))
    resolved = _resolve_tender_template_block_span(node, source_doc)
    if resolved is None:
        return None
    return _resolved_span_to_anchors(resolved, source_doc)


def copy_tender_template_by_node(
    node: Dict[str, Any],
    tender_doc_path: str,
    target_doc,
) -> bool:
    """按目录节点名在招标书中查找附件范本并复制。返回 True 表示找到并复制了内容。"""
    source_path = Path(tender_doc_path)
    if not source_path.exists() or source_path.suffix.lower() != ".docx":
        return False

    source_doc = Document(str(source_path))
    # renderer_title_match anchors are cached search results. Re-resolve them so
    # improvements to structural boundary detection also repair old projects.
    prefer_resolved_span = str(node.get("copy_method") or "") == "renderer_title_match"
    if not prefer_resolved_span and _copy_tender_template_by_anchors(node, source_doc, target_doc):
        return True
    resolved = _resolve_tender_template_block_span(node, source_doc)
    if resolved is None:
        return _copy_tender_template_by_anchors(node, source_doc, target_doc)
    blocks = resolved["blocks"]
    start = resolved["start"]
    end = resolved["end"]
    skip_indexes = resolved["skip_indexes"]
    copied = 0
    for idx, block in enumerate(blocks[start:end], start=start):
        if idx in skip_indexes:
            continue
        if block.tag.endswith("sectPr"):
            continue
        copied_block = deepcopy(block)
        _strip_toc_heading_marks(copied_block)
        fit_copied_tender_template_tables_to_page(copied_block, target_doc)
        _copy_block_image_relationships(copied_block, source_doc, target_doc)
        _append_body_block(target_doc, copied_block)
        copied += 1

    return copied > 0


def _anchor_block_number(value: Any) -> Optional[int]:
    match = re.fullmatch(r"p(\d+)", str(value or "").strip())
    return int(match.group(1)) if match else None


def _copy_tender_template_by_anchors(node: Dict[str, Any], source_doc, target_doc) -> bool:
    start = _anchor_block_number(node.get("anchor_start"))
    end = _anchor_block_number(node.get("anchor_end"))
    if start is None or end is None or end < start:
        return False

    paragraph_by_element = {paragraph._element: paragraph for paragraph in source_doc.paragraphs}
    anchored_blocks = []
    for block in source_doc.element.body.iterchildren():
        if block.tag.endswith("}p"):
            paragraph = paragraph_by_element.get(block)
            if paragraph is None or not paragraph.text.strip():
                continue
            anchored_blocks.append(block)
        elif block.tag.endswith("}tbl"):
            anchored_blocks.append(block)
    if start >= len(anchored_blocks):
        return False
    end = min(end, len(anchored_blocks) - 1)

    copied = 0
    target_title = _template_title_core(_clean_title(str(node.get("name") or "")))
    allow_partial_title = not _should_preserve_internal_template_headings(str(node.get("name") or ""))
    for offset, block in enumerate(anchored_blocks[start : end + 1]):
        if offset == 0 and block.tag.endswith("}p"):
            paragraph = paragraph_by_element.get(block)
            if paragraph is not None and _is_pure_duplicate_template_title(
                paragraph.text,
                target_title,
                allow_partial=allow_partial_title,
            ):
                continue
        copied_block = deepcopy(block)
        _strip_toc_heading_marks(copied_block)
        fit_copied_tender_template_tables_to_page(copied_block, target_doc)
        _copy_block_image_relationships(copied_block, source_doc, target_doc)
        _append_body_block(target_doc, copied_block)
        copied += 1
    return copied > 0


def _leading_duplicate_title_indexes(
    blocks,
    paragraph_by_element,
    start: int,
    end: int,
    node_name: str,
) -> set[int]:
    """跳过已由投标书标题表示的源表单开头标题行。

    仅跳过位于开头、内容纯粹为标题或附件标签的段落。像“报价函\n广州越秀...”
    这类混合段落包含实际表单正文，不能跳过。
    """
    target = _template_title_core(_clean_title(node_name))
    allow_partial_match = not _should_preserve_internal_template_headings(node_name)
    skipped: set[int] = set()
    for idx in range(start, min(end, len(blocks))):
        block = blocks[idx]
        if not block.tag.endswith("}p"):
            break
        paragraph = paragraph_by_element.get(block)
        text = paragraph.text.strip() if paragraph is not None else ""
        if not text:
            continue
        if _is_pure_duplicate_template_title(text, target, allow_partial=allow_partial_match):
            skipped.add(idx)
            continue
        break
    return skipped


def _is_pure_duplicate_template_title(
    text: str,
    target_title: str,
    *,
    allow_partial: bool = True,
) -> bool:
    lines = [line.strip() for line in str(text or "").splitlines() if line.strip()]
    if not lines:
        return False
    saw_title = False
    for line in lines:
        if _is_bare_attachment_label(line):
            continue
        core = _template_title_core(_clean_title(line))
        if core and (
            core == target_title
            or (allow_partial and (core in target_title or target_title in core))
        ):
            saw_title = True
            continue
        return False
    return saw_title or all(_is_bare_attachment_label(line) for line in lines)


def _collect_heading_candidates(source_doc, blocks) -> list[dict]:
    """收集所有可能的范本标题候选"""
    candidates = []
    paragraph_by_element = {p._element: p for p in source_doc.paragraphs}
    for idx, block in enumerate(blocks):
        if not block.tag.endswith("}p"):
            continue
        paragraph = paragraph_by_element.get(block)
        if paragraph is None:
            continue
        text = paragraph.text.strip()
        if not text:
            continue
        if _is_attachment_directory_item(blocks, paragraph_by_element, idx):
            continue

        # 严格识别 heading
        level = _detect_heading_level(paragraph, text)

        if _is_bare_attachment_label(text):
            next_title = _next_nonempty_paragraph_text(blocks, paragraph_by_element, idx)
            if next_title:
                candidates.append(
                    {
                        "block_idx": idx,
                        "title": f"{text.strip()} {next_title.strip()}",
                        "clean_title": _clean_title(next_title),
                        "level": 1,
                        "attachment_label": text.strip(),
                    }
                )
                continue

        # ★ V3 兜底:短段落 + 模板词后缀 + 后续有"必填/盖章"等
        if level == 0 and _is_template_title_with_context(blocks, paragraph_by_element, idx, text):
            level = 2

        if level == 0:
            continue
        candidates.append(
            {
                "block_idx": idx,
                "title": text,
                "clean_title": _clean_title(text),
                "level": level,
            }
        )
    return candidates


def _is_template_title_with_context(blocks, paragraph_by_element, idx, text) -> bool:
    """模板词上下文兜底:看后续段落是否含'必填/格式自拟/盖章/签字'等"""
    text_no_space = re.sub(r"\s+", "", text)
    template_suffixes = ['函', '书', '表', '声明', '承诺', '协议', '清单', '证明', '范本', '模板', '附录']
    has_suffix = any(text_no_space.endswith(s) for s in template_suffixes)
    if not has_suffix or len(text_no_space) > 15:
        return False

    # 排除明显的引用("投标函(格式见附件)" 这种是说明,不是范本)
    if any(marker in text for marker in ('格式见附件', '格式见后', '详见', '参见', '参考')):
        return False

    # 看后续 5 段,有"必填"等关键标记 → 确认是范本
    for i in range(idx + 1, min(idx + 6, len(blocks))):
        block = blocks[i]
        if not block.tag.endswith('}p'):
            continue
        p = paragraph_by_element.get(block)
        if p is None:
            continue
        next_text = p.text.strip()
        if not next_text:
            continue
        if any(m in next_text for m in ('必填', '格式自拟', '盖单位公章', '盖章', '签章', '签字', '本人')):
            return True
        # 长正文就停(说明不是范本)
        if len(next_text) > 10:
            break
    return False


def _best_template_match(node_name: str, candidates: list[dict]) -> Optional[dict]:
    keywords = _template_keywords(node_name)
    if not keywords:
        return None

    target_titles = _target_template_titles(node_name)
    title_priority = {title: idx for idx, title in enumerate(target_titles)}
    exact_candidates = [
        candidate
        for candidate in candidates
        if candidate["clean_title"] in title_priority
        and not _is_instruction_title(candidate["title"])
    ]
    if exact_candidates:
        return min(
            exact_candidates,
            key=lambda item: (title_priority[item["clean_title"]], -item["block_idx"]),
        )

    best = None
    best_score = 0
    for candidate in candidates:
        title = candidate["clean_title"]
        raw_title = candidate["title"]
        if _is_instruction_title(raw_title):
            continue
        score = 0
        matched = 0
        for keyword in keywords:
            if keyword and keyword in title:
                matched += 1
                score += 20
        if matched == 0:
            continue
        if matched == len(keywords):
            score += 30
        if any(word in raw_title for word in ("附件", "格式", "范本", "模板")):
            score += 40
        if title in target_titles or title == _clean_title(node_name):
            score += 30
        score += (3 - min(candidate["level"], 3)) * 30
        score += min(candidate["block_idx"] // 80, 20)

        if score > best_score:
            best_score = score
            best = candidate

    return best if best_score >= 20 else None


def _find_attachment_title_match(node_name: str, source_doc, blocks) -> Optional[dict]:
    """按节点标题通用匹配附件正文标题"""
    referenced_match = _find_explicit_attachment_reference_match(node_name, source_doc, blocks)
    if referenced_match is not None:
        return referenced_match

    name = _clean_title(node_name)
    target_titles = _title_variants(name)
    target_tokens = _extract_template_tokens(name)
    if not target_titles and not target_tokens:
        return None

    paragraph_by_element = {p._element: p for p in source_doc.paragraphs}
    matches = []
    for idx, block in enumerate(blocks):
        if not block.tag.endswith("}p"):
            continue
        paragraph = paragraph_by_element.get(block)
        if paragraph is None:
            continue
        raw_title = paragraph.text.strip()
        if not raw_title:
            continue
        if _is_attachment_directory_item(blocks, paragraph_by_element, idx):
            continue
        if _is_bare_attachment_label(raw_title):
            next_title = _next_nonempty_paragraph_text(blocks, paragraph_by_element, idx)
            if next_title:
                raw_title = f"{raw_title.strip()} {next_title.strip()}"
        if _is_numbered_instruction_item(raw_title):
            continue

        # 严格识别 + 兜底
        level = _detect_heading_level(paragraph, raw_title)
        if level == 0:
            if not _is_template_title_with_context(blocks, paragraph_by_element, idx, raw_title):
                continue
            level = 2

        clean_title = _clean_title(raw_title)
        clean_variants = _title_variants(clean_title)
        clean_tokens = _extract_template_tokens(clean_title)

        title_overlap = bool(target_titles & clean_variants)
        token_overlap = bool(target_tokens & clean_tokens)
        # ★ V3 新增:双向包含匹配,解决"授权委托书" vs "三、法定代表人投标授权委托书"
        token_substring = _has_substring_overlap(target_tokens, clean_tokens)
        similarity = _template_similarity(name, clean_title)

        if (
            not title_overlap
            and not token_overlap
            and not token_substring
            and similarity < 0.58
        ):
            continue

        # 评分
        score = idx
        if _attachment_reference(raw_title):
            score += 10000
        score += (3 - min(level, 3)) * 1500  # level 高的优先(父级标题)
        if title_overlap:
            score += 5000
        if target_tokens and clean_tokens and target_tokens == clean_tokens:
            score += 2000
        if token_substring and not token_overlap:
            score += 500  # 子串匹配也给分,但低于精确
        score += int(similarity * 3500)
        if similarity >= 0.72:
            score += 2000
        if _should_copy_representative_authorization_group(node_name) and _is_representative_proof_title(clean_title):
            score += 6000
        # 复合标题加分
        if any(sep in raw_title for sep in ("及", "与", "和")):
            score += 1000

        matches.append(
            {
                "block_idx": idx,
                "title": raw_title,
                "clean_title": clean_title,
                "level": level,
                "_score": score,
                "_similarity": round(similarity, 3),
            }
        )

    if not matches:
        return None
    if _should_copy_representative_authorization_group(node_name):
        proof_matches = [
            item for item in matches if _is_representative_proof_title(item.get("clean_title", ""))
        ]
        if proof_matches:
            match = max(proof_matches, key=lambda item: item.get("_score", 0))
            match.pop("_score", None)
            match.pop("_similarity", None)
            return match
    leaf_match = _best_specific_leaf_template_match(node_name, matches)
    if leaf_match is not None:
        leaf_match.pop("_score", None)
        leaf_match.pop("_similarity", None)
        return leaf_match

    match = max(matches, key=lambda item: item["_score"])
    match.pop("_score", None)
    match.pop("_similarity", None)
    return match


def _best_specific_leaf_template_match(node_name: str, matches: list[dict]) -> Optional[dict]:
    target_core = _template_title_core(_clean_title(node_name))
    if not target_core or _is_composite_template_title(target_core):
        return None

    target_tokens = _extract_template_tokens(target_core)
    best: Optional[dict] = None
    best_score = 0
    for item in matches:
        raw_title = str(item.get("title") or "")
        clean_title = _template_title_core(_clean_title(str(item.get("clean_title") or raw_title)))
        if not clean_title:
            continue

        if _is_composite_template_title(raw_title) and clean_title != target_core:
            continue

        score = _specific_child_template_score(target_core, clean_title)
        clean_tokens = _extract_template_tokens(clean_title)
        if target_tokens and clean_tokens and target_tokens == clean_tokens:
            score = max(score, 130)
        if clean_title.endswith(target_core) and len(clean_title) <= len(target_core) + 24:
            score = max(score, 110)

        if score > best_score:
            best_score = score
            best = item

    return best if best is not None and best_score >= 80 else None


def _is_composite_template_title(title: str) -> bool:
    cleaned = _clean_title(title)
    if not any(sep in cleaned for sep in ("及", "与", "和", "、", "/", "／")):
        return False
    return len(_extract_template_tokens(cleaned)) >= 2


def _find_explicit_attachment_reference_match(node_name: str, source_doc, blocks) -> Optional[dict]:
    """Resolve an explicitly referenced appendix before title similarity.

    Directory wording and the actual form title can differ. Their shared
    attachment/schedule number is the authoritative structural identifier.
    """
    target = _attachment_reference(node_name)
    if target is None:
        return None

    paragraph_by_element = {paragraph._element: paragraph for paragraph in source_doc.paragraphs}
    for idx, block in enumerate(blocks):
        if not block.tag.endswith("}p"):
            continue
        paragraph = paragraph_by_element.get(block)
        raw_label = paragraph.text.strip() if paragraph is not None else ""
        if not _is_bare_attachment_label(raw_label) or _attachment_reference(raw_label) != target:
            continue
        next_title = _next_nonempty_paragraph_text(blocks, paragraph_by_element, idx)
        return {
            "block_idx": idx,
            "title": f"{raw_label} {next_title}".strip(),
            "clean_title": _clean_title(next_title or raw_label),
            "level": 1,
            "attachment_label": raw_label,
        }
    return None


def _next_nonempty_paragraph_text(blocks, paragraph_by_element, idx: int, max_lookahead: int = 4) -> Optional[str]:
    for offset in range(1, max_lookahead + 1):
        if idx + offset >= len(blocks):
            break
        block = blocks[idx + offset]
        if not block.tag.endswith("}p"):
            break
        paragraph = paragraph_by_element.get(block)
        if paragraph is None:
            continue
        text = paragraph.text.strip()
        if text:
            for line in text.splitlines():
                line = line.strip()
                if line:
                    return line
    return None


def _is_attachment_directory_item(blocks, paragraph_by_element, idx: int) -> bool:
    """跳过“3.报价函”这类附件目录行。

    这些行出现在“附件：”附近，只是在列出附件文件名，不是实际范本正文，
    不能作为表单章节复制。
    """
    block = blocks[idx]
    if not block.tag.endswith("}p"):
        return False
    paragraph = paragraph_by_element.get(block)
    current = paragraph.text.strip() if paragraph is not None else ""
    if not re.match(r"^\s*\d+\s*[.．、]\s*\S{1,30}\s*$", current):
        return False

    saw_attachment_label = False
    numbered_neighbors = 0
    for probe in range(max(0, idx - 6), min(len(blocks), idx + 7)):
        neighbor = blocks[probe]
        if not neighbor.tag.endswith("}p"):
            continue
        neighbor_paragraph = paragraph_by_element.get(neighbor)
        text = neighbor_paragraph.text.strip() if neighbor_paragraph is not None else ""
        if not text:
            continue
        compact = re.sub(r"\s+", "", text)
        if re.match(r"^附件[:：]", compact):
            saw_attachment_label = True
        if re.match(r"^\d+\s*[.．、]\s*\S{1,30}\s*$", text):
            numbered_neighbors += 1
    return saw_attachment_label and numbered_neighbors >= 4


def _has_substring_overlap(set_a: set, set_b: set) -> bool:
    """两个 token 集合,任一对儿有双向子串关系即返回 True"""
    for a in set_a:
        for b in set_b:
            if len(a) >= 3 and len(b) >= 3:  # 太短的子串没意义
                if a in b or b in a:
                    return True
    return False


def _template_similarity(left: str, right: str) -> float:
    """短标题相似度:归一化文本 + token + 字符 ngram,避免硬编码大量同义词。"""
    left_norm = _similarity_text(left)
    right_norm = _similarity_text(right)
    if not left_norm or not right_norm:
        return 0.0

    seq_score = SequenceMatcher(None, left_norm, right_norm).ratio()
    ngram_score = _ngram_cosine(left_norm, right_norm)

    left_tokens = _extract_template_tokens(left_norm)
    right_tokens = _extract_template_tokens(right_norm)
    token_score = 0.0
    if left_tokens or right_tokens:
        intersection = left_tokens & right_tokens
        union = left_tokens | right_tokens
        token_score = len(intersection) / len(union) if union else 0.0
        if _has_substring_overlap(left_tokens, right_tokens):
            token_score = max(token_score, 0.75)

    return max(seq_score * 0.45 + ngram_score * 0.35 + token_score * 0.20, token_score)


def _similarity_text(text: str) -> str:
    normalized = _normalize_template_text(_clean_title(text))
    normalized = re.sub(r"[《》〈〉“”\"'（）()\s]", "", normalized)
    normalized = re.sub(r"(本项目|采购项目|响应文件|投标文件)", "", normalized)
    return normalized


def _ngram_cosine(left: str, right: str, n: int = 2) -> float:
    left_grams = _char_ngrams(left, n)
    right_grams = _char_ngrams(right, n)
    if not left_grams or not right_grams:
        return 0.0
    common = set(left_grams) & set(right_grams)
    numerator = sum(left_grams[g] * right_grams[g] for g in common)
    left_norm = sum(v * v for v in left_grams.values()) ** 0.5
    right_norm = sum(v * v for v in right_grams.values()) ** 0.5
    return numerator / (left_norm * right_norm) if left_norm and right_norm else 0.0


def _char_ngrams(text: str, n: int) -> dict[str, int]:
    if len(text) <= n:
        return {text: 1} if text else {}
    grams: dict[str, int] = {}
    for index in range(len(text) - n + 1):
        gram = text[index : index + n]
        grams[gram] = grams.get(gram, 0) + 1
    return grams


def _title_variants(title: str) -> set[str]:
    cleaned = _normalize_template_text(title)
    cleaned = re.sub(r"[（(].*?[）)]", "", cleaned)
    cleaned = re.sub(r"\s+", "", cleaned)
    cleaned = cleaned.strip("：:.-—_ ")
    variants = {cleaned} if cleaned else set()
    variants.update(_alias_template_titles(cleaned))
    for prefix in ("本项目", "项目", "采购项目", "响应文件", "投标文件", "投标人", "投标单位"):
        if cleaned.startswith(prefix) and len(cleaned) > len(prefix) + 1:
            variants.add(cleaned[len(prefix):])
    return variants


def _extract_template_tokens(title: str) -> set[str]:
    cleaned = _normalize_template_text(title)
    cleaned = re.sub(r"[（(].*?[）)]", "", cleaned)
    cleaned = re.sub(r"[《》〈〉“”\"']", "", cleaned)
    cleaned = re.sub(r"附件[一二三四五六七八九十\d]+", "", cleaned)
    cleaned = re.sub(r"\d+(?:\.\d+)*", "", cleaned)
    parts = re.split(r"[、，,；;\/和与及\s]+", cleaned)
    tokens = set()
    for part in parts:
        token = part.strip("：:.-—_ ")
        tokens.update(_template_token_variants(token))
    tokens.update(_alias_template_tokens(cleaned))
    return tokens


def _normalize_template_text(text: str) -> str:
    return (text or "").replace("申明", "声明").replace("授权书", "授权委托书")


def _looks_like_authorization_request(cleaned: str) -> bool:
    if "授权" not in cleaned:
        return False
    return any(word in cleaned for word in ("法定代表人", "法人", "代表", "委托", "授权文件"))


def _looks_like_confirmation_request(cleaned: str) -> bool:
    if "确认" not in cleaned:
        return False
    return any(word in cleaned for word in ("费用", "服务费", "报价", "金额", "价格", "价款", "收费"))


def _alias_template_titles(cleaned: str) -> set[str]:
    aliases: set[str] = set()
    if _looks_like_authorization_request(cleaned):
        aliases.update(
            {
                "授权委托书",
                "法人授权委托书",
                "法定代表人授权委托书",
                "法定代表人证明书",
                "法定代表人身份证明",
                "法人代表证明书",
                "授权书",
                "委托书",
            }
        )
    if _looks_like_confirmation_request(cleaned):
        aliases.update({"确认书", "确认表", "确认单"})
        if any(word in cleaned for word in ("费用", "服务费", "收费", "价款")):
            aliases.update({"费用确认书", "费用确认表", "费用确认单"})
        if any(word in cleaned for word in ("报价", "金额", "价格")):
            aliases.update({"报价确认书", "报价确认表", "报价确认单", "金额确认表"})
    return aliases


def _alias_template_tokens(cleaned: str) -> set[str]:
    aliases: set[str] = set()
    if _looks_like_authorization_request(cleaned):
        aliases.update({"授权委托书", "授权书", "委托书", "身份证明", "证明书"})
    if _looks_like_confirmation_request(cleaned):
        aliases.update({"确认书", "确认表", "确认单", "明细表"})
    return aliases


def _template_token_variants(token: str) -> set[str]:
    variants = set()
    if not (2 <= len(token) <= 24):
        return variants
    if not re.search(r"(函|书|表|单|清单|声明|承诺|证明|材料|附录|介绍)$", token):
        return variants

    variants.add(token)
    for prefix in ("投标人", "投标单位", "供应商", "报价人", "法定代表人", "法人", "项目"):
        if token.startswith(prefix) and len(token) > len(prefix) + 1:
            variants.add(token[len(prefix):])

    for keyword in (
        "资格声明函",
        "声明函",
        "报价表",
        "报价单",
        "明细表",
        "身份证明",
        "授权委托书",
        "授权书",
        "基本情况介绍",
    ):
        if keyword in token:
            variants.add(keyword)
    if "报价" in token and (token.endswith("表") or token.endswith("单")):
        variants.update({"报价表", "报价单"})
    return variants


def _template_keywords(node_name: str) -> list[str]:
    name = _clean_title(node_name)
    tokens = _extract_template_tokens(name)
    if tokens:
        return sorted(tokens)
    return sorted(title for title in _title_variants(name) if len(title) >= 3)


def _target_template_titles(node_name: str) -> list[str]:
    name = _clean_title(node_name)
    titles = _title_variants(name)
    tokens = _extract_template_tokens(name)
    return sorted(titles | tokens)


def _is_instruction_title(text: str) -> bool:
    cleaned = _clean_title(text)
    if len(cleaned) > 25:
        return True
    has_instruction_word = any(
        word in text
        for word in (
            "参见", "参考", "须", "均须", "提供", "加盖",
            "文件", "格式", "详见", "说明",
        )
    )
    return has_instruction_word and not _is_template_title_looking(cleaned)


def _is_template_title_looking(text: str) -> bool:
    """文本看着像范本标题(短 + 模板词后缀)"""
    text = _normalize_template_text(text)
    if len(text) > 25:
        return False
    text_no_space = re.sub(r"\s+", "", text)
    return any(text_no_space.endswith(s) for s in
               ['函', '书', '表', '单', '声明', '承诺', '协议', '清单', '证明', '材料', '附录', '介绍'])


def _narrow_generic_parent_template_span(
    blocks,
    paragraph_by_element,
    node_name: str,
    match: dict,
    start: int,
    end: int,
) -> Optional[tuple[int, int]]:
    """When a generic parent form matches, copy the exact child form if present.

    Some tenders group several independent forms under a wrapper such as
    "自查表". The matcher may intentionally prefer that parent heading, but a
    leaf outline node like "技术评审自查表" should not receive the whole wrapper
    group. If the matched title is only a generic suffix of the node title, look
    inside the matched range for a more specific child heading and narrow the
    copied span to that child.
    """
    if _should_copy_representative_authorization_group(node_name):
        return None

    target_core = _template_title_core(_clean_title(node_name))
    matched_core = _template_title_core(_clean_title(match.get("clean_title") or match.get("title", "")))
    if not _is_generic_parent_template_match(target_core, matched_core):
        return None

    child_start = _find_specific_child_template_heading(
        blocks,
        paragraph_by_element,
        target_core,
        start + 1,
        end,
    )
    if child_start is None:
        return None

    child_end = _find_next_child_template_boundary(
        blocks,
        paragraph_by_element,
        child_start,
        end,
    )
    if child_end <= child_start:
        return None
    return child_start, child_end


def _is_generic_parent_template_match(target_core: str, matched_core: str) -> bool:
    if not target_core or not matched_core or target_core == matched_core:
        return False
    if len(matched_core) < 2:
        return False
    return matched_core in target_core and len(target_core) - len(matched_core) >= 2


def _find_specific_child_template_heading(
    blocks,
    paragraph_by_element,
    target_core: str,
    start: int,
    end: int,
) -> Optional[int]:
    best_idx: Optional[int] = None
    best_score = 0
    for idx in range(start, end):
        block = blocks[idx]
        if not block.tag.endswith("}p"):
            continue
        paragraph = paragraph_by_element.get(block)
        raw_title = paragraph.text.strip() if paragraph is not None else ""
        clean_title = _template_title_core(_clean_title(raw_title))
        if not _looks_like_child_template_heading(raw_title, clean_title):
            continue
        score = _specific_child_template_score(target_core, clean_title)
        if score > best_score:
            best_score = score
            best_idx = idx
    return best_idx if best_score >= 80 else None


def _specific_child_template_score(target_core: str, clean_title: str) -> int:
    if not target_core or not clean_title:
        return 0
    if clean_title == target_core:
        return 120
    if len(clean_title) >= 3 and (clean_title in target_core or target_core in clean_title):
        return 100
    similarity = _template_similarity(target_core, clean_title)
    if similarity >= 0.86:
        return int(similarity * 90)
    return 0


def _find_next_child_template_boundary(
    blocks,
    paragraph_by_element,
    child_start: int,
    end: int,
) -> int:
    for idx in range(child_start + 1, end):
        block = blocks[idx]
        if not block.tag.endswith("}p"):
            continue
        paragraph = paragraph_by_element.get(block)
        raw_title = paragraph.text.strip() if paragraph is not None else ""
        clean_title = _template_title_core(_clean_title(raw_title))
        if _looks_like_child_template_heading(raw_title, clean_title):
            return idx
    return end


def _looks_like_child_template_heading(raw_title: str, clean_title: str) -> bool:
    if not clean_title or len(clean_title) > 30:
        return False
    stripped = raw_title.strip()
    if not stripped or _ends_like_sentence_or_list_item(stripped):
        return False
    if stripped.startswith(("注", "备注")):
        return False
    if len(stripped) > 45:
        return False
    return _is_template_title_looking(clean_title)


def _clip_template_span_at_internal_sibling_heading(
    blocks,
    paragraph_by_element,
    start: int,
    end: int,
) -> int:
    """Stop exact template copies at the next sibling form heading.

    Some converted .doc files do not expose every "（1）（2）（3）" form title as a
    heading candidate, so _find_section_end can leave a leaf form spanning over
    later sibling forms. A short numbered form-looking paragraph inside the span
    is a safe deterministic boundary.
    """
    start_marker = _numbered_template_heading_marker(_paragraph_text(blocks, paragraph_by_element, start))
    for idx in range(start + 1, end):
        raw_title = _paragraph_text(blocks, paragraph_by_element, idx)
        if not raw_title:
            continue
        clean_title = _template_title_core(_clean_title(raw_title))
        if not _looks_like_child_template_heading(raw_title, clean_title):
            continue
        marker = _numbered_template_heading_marker(raw_title)
        if marker and start_marker and marker == start_marker:
            return idx
        if marker and not start_marker:
            return idx
    return end


def _paragraph_text(blocks, paragraph_by_element, idx: int) -> str:
    if idx < 0 or idx >= len(blocks):
        return ""
    block = blocks[idx]
    if not block.tag.endswith("}p"):
        return ""
    paragraph = paragraph_by_element.get(block)
    return paragraph.text.strip() if paragraph is not None else ""


def _numbered_template_heading_marker(text: str) -> str:
    stripped = str(text or "").strip()
    if re.match(r"^[（(][一二三四五六七八九十]+[）)]", stripped):
        return "paren_cn"
    if re.match(r"^[（(]\d+[）)]", stripped):
        return "paren_num"
    if re.match(r"^[一二三四五六七八九十]+[、.．]", stripped):
        return "cn"
    if re.match(r"^\d+(?:\.\d+)*[、.．]", stripped):
        return "num"
    return ""


def _find_section_end(candidates: list[dict], match: dict, block_count: int) -> int:
    start = match["block_idx"]
    level = match["level"]
    parent_marker = _numbered_template_heading_marker(str(match.get("title") or ""))
    composite_parent = _should_preserve_internal_template_headings(str(match.get("title") or ""))
    for candidate in candidates:
        if candidate["block_idx"] <= start:
            continue
        if _is_child_template_heading(match, candidate):
            continue
        candidate_level = int(candidate.get("level") or 0)
        if candidate_level < level:
            return candidate["block_idx"]
        if candidate_level != level:
            continue
        if composite_parent and parent_marker:
            candidate_marker = _numbered_template_heading_marker(str(candidate.get("title") or ""))
            if candidate_marker and candidate_marker != parent_marker:
                continue
        return candidate["block_idx"]
    return block_count


def _clip_end_at_cover_boundary(source_doc, blocks, start: int, end: int) -> int:
    """Some tenders place a new response-file cover between form templates.

    That cover is not part of the preceding form. Treat short "投标文件/响应文件"
    cover blocks followed by project/company fields as a hard boundary.
    """
    paragraph_by_element = {p._element: p for p in source_doc.paragraphs}
    for idx in range(start + 2, end):
        block = blocks[idx]
        if not block.tag.endswith("}p"):
            continue
        paragraph = paragraph_by_element.get(block)
        if paragraph is None:
            continue
        text = re.sub(r"\s+", "", paragraph.text or "")
        if not _looks_like_response_cover_title(text):
            continue
        if _has_cover_context(blocks, paragraph_by_element, idx):
            return idx
    return end


def _looks_like_response_cover_title(text: str) -> bool:
    if not text or len(text) > 18:
        return False
    return text in {
        "投标文件",
        "响应文件",
        "投标响应文件",
        "报价文件",
        "商务投标文件",
        "技术投标文件",
    }


def _has_cover_context(blocks, paragraph_by_element, idx: int) -> bool:
    hits = 0
    for offset in range(1, 25):
        if idx + offset >= len(blocks):
            break
        block = blocks[idx + offset]
        if not block.tag.endswith("}p"):
            continue
        paragraph = paragraph_by_element.get(block)
        if paragraph is None:
            continue
        text = re.sub(r"\s+", "", paragraph.text or "")
        if not text:
            continue
        if any(token in text for token in ("商务技术部分", "商务部分", "技术部分", "项目名称", "投标单位", "供应商名称", "法定代表人")):
            hits += 1
    return hits >= 2


def _is_child_template_heading(parent: dict, candidate: dict) -> bool:
    """复合标题下的(一)(二)子表单不应截断复制范围"""
    parent_title = parent.get("title", "") or parent.get("clean_title", "")
    candidate_title = candidate.get("title", "") or ""
    parent_is_attachment_title = bool(parent.get("attachment_label")) or _is_bare_attachment_label(
        str(parent_title).split()[0] if str(parent_title).split() else ""
    )
    if parent_is_attachment_title:
        same_title = _template_title_core(parent.get("clean_title", "")) == _template_title_core(
            candidate.get("clean_title", "")
        )
        adjacent = int(candidate.get("block_idx", 0)) <= int(parent.get("block_idx", 0)) + 2
        if same_title and adjacent:
            return True
    if not _is_composite_template_title(parent_title):
        return False
    # 兼容全/半角括号 + 中文/阿拉伯数字
    return bool(re.match(r"^[(\uff08(][一二三四五六七八九十\d]+[)\uff09)]", candidate_title.strip()))


def _is_composite_template_title(title: str) -> bool:
    cleaned = _clean_title(title)
    tokens = _extract_template_tokens(cleaned)
    if len(tokens) >= 2:
        return True
    return any(separator in cleaned for separator in ("及", "和", "与", "、", "/"))


def _should_preserve_internal_template_headings(node_name: str) -> bool:
    """Whether one outline node intentionally spans several source form headings."""
    cleaned = _clean_title(node_name)
    if _should_copy_representative_authorization_group(cleaned):
        return True
    # Only explicit conjunctions mean the node asks for a combined template. A
    # title like "法定代表人授权委托书" has multiple keywords but is still one form.
    return any(separator in cleaned for separator in ("及", "和", "与", "、", "/", "／"))


def normalize_copied_tender_template_block(block) -> None:
    """Make copied official template blocks fit the target document layout.

    Source tender files may use section styles, large paragraph spacing, or
    keep-with-next controls that are safe on the original page but expand after
    insertion under a generated bid heading.
    """
    strip_section_properties(block)
    _strip_page_break_marks(block)
    for p_style in block.findall(".//" + qn("w:pStyle")):
        remove_element(p_style)
    for outline_level in block.findall(".//" + qn("w:outlineLvl")):
        remove_element(outline_level)
    _compact_copied_template_paragraph_layout(block)


def _strip_toc_heading_marks(block) -> None:
    normalize_copied_tender_template_block(block)


def _compact_copied_template_paragraph_layout(block) -> None:
    for p_pr in block.findall(".//" + qn("w:pPr")):
        for tag in ("w:keepNext", "w:keepLines", "w:pageBreakBefore"):
            for child in list(p_pr.findall(qn(tag))):
                p_pr.remove(child)

        spacing = p_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            p_pr.append(spacing)
        for attr in ("w:beforeAutospacing", "w:afterAutospacing"):
            spacing.attrib.pop(qn(attr), None)
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "auto")


def _detect_heading_level(paragraph, text: str) -> int:
    """严格识别 heading,避免误判正文里的列表编号"""
    style_name = (paragraph.style.name or "").strip()
    lower = style_name.lower()

    # 1. Word Heading 样式(最可靠)
    if lower.startswith("heading "):
        try:
            return int(lower.split(" ")[1])
        except (IndexError, ValueError):
            return 0
    if style_name.startswith("标题"):
        num = style_name.replace("标题", "").replace(" ", "")
        if num.isdigit():
            return int(num)

    # WPS/Word may mark visual headings with outlineLvl but no Heading style.
    outline_level = _paragraph_outline_level(paragraph)
    if outline_level is not None:
        cleaned = _template_title_core(_clean_title(text))
        if cleaned and len(cleaned) < 25 and _is_template_title_looking(cleaned):
            return outline_level

    # 2. "附件 X" 开头
    if re.match(r"^附件[一二三四五六七八九十\d]+", text):
        return 1

    # 3. "(一)" "(1)" 全/半角括号
    if re.match(r"^[(\uff08(][一二三四五六七八九十\d]+[)\uff09)]", text.strip()):
        cleaned = text.strip()
        if len(cleaned) < 25 and not _ends_like_sentence_or_list_item(cleaned):
            return 3

    # 4. ★ 关键修复:数字编号"1." "2." 不再识别为 heading
    #    投标函正文的"1、 2、 3、 4、 5、 6、 7、" 都不再被误判
    #    要识别 1.1 这种章节,必须用 Heading 样式
    #    例外:响应文件格式里短的"2.法定代表人证明书/3．法人授权委托书"
    #    这类格式范本标题需要作为复制边界,但必须看起来像函/书/表/证明等范本标题。
    if re.match(r"^\d+\s*[.．]\s*", text.strip()):
        cleaned = _template_title_core(_clean_title(text))
        if len(cleaned) < 25 and _is_template_title_looking(cleaned):
            return 2

    # 5. 中文"一、二、" 严格
    if re.match(r"^[一二三四五六七八九十]+、", text):
        stripped = text.rstrip()
        if len(stripped) < 25 and not _ends_like_sentence_or_list_item(stripped):
            return 2

    return 0



def _paragraph_outline_level(paragraph) -> Optional[int]:
    """Return Word outline level as 1-based heading level when present."""
    try:
        p_pr = getattr(getattr(paragraph, "_p", None), "pPr", None)
        outline_level = p_pr.find(qn("w:outlineLvl")) if p_pr is not None else None
        if outline_level is None:
            return None
        value = outline_level.get(qn("w:val"))
        if value is None:
            return None
        return int(value) + 1
    except (TypeError, ValueError, AttributeError):
        return None

def _clean_title(text: str) -> str:
    """剥离编号"""
    cleaned = strip_leading_attachment_label(text)
    cleaned = re.sub(r"^\d+(?:[.．]\d+)*[.．、\s]*", "", cleaned)
    cleaned = re.sub(r"^[一二三四五六七八九十]+[、.\s]+", "", cleaned)
    return re.sub(r"\s+", "", cleaned)


def _template_title_core(text: str) -> str:
    cleaned = re.sub(r"[（(].*?[）)]", "", text or "")
    return re.sub(r"\s+", "", cleaned).strip("：:.-—_ ")


def _ends_like_sentence_or_list_item(text: str) -> bool:
    return text.rstrip().endswith(("。", "；", ";", "，", ",", "：", ":"))


def _is_numbered_instruction_item(text: str) -> bool:
    stripped = text.strip()
    if not re.match(r"^(?:[(（][一二三四五六七八九十\d]+[)）]|\d+[.．、]|[一二三四五六七八九十]+、)", stripped):
        return False
    if _ends_like_sentence_or_list_item(stripped):
        return True
    if len(stripped) > 30 and any(word in stripped for word in ("应", "须", "需", "提供", "包括", "资料", "文件")):
        return True
    return False


def _should_copy_representative_authorization_group(node_name: str) -> bool:
    cleaned = _clean_title(node_name)
    if "授权" not in cleaned:
        return False
    has_representative = any(word in cleaned for word in ("法定代表人", "法人", "代表"))
    has_proof_hint = any(word in cleaned for word in ("证明", "身份", "授权文件", "及授权", "或授权"))
    return has_representative and has_proof_hint


def _is_representative_proof_title(title: str) -> bool:
    cleaned = _clean_title(title)
    if not any(word in cleaned for word in ("法定代表人", "法人", "代表人")):
        return False
    return any(word in cleaned for word in ("证明书", "身份证明", "证明"))


def _is_authorization_title(title: str) -> bool:
    cleaned = _clean_title(title)
    return "授权" in cleaned and any(word in cleaned for word in ("委托书", "授权书", "委托"))


def _extend_authorization_template_end(
    candidates: list[dict],
    match: dict,
    current_end: int,
    block_count: int,
) -> int:
    """Copy representative proof and authorization forms as one template group."""
    start = match["block_idx"]
    auth_candidate = None
    for candidate in candidates:
        if candidate["block_idx"] <= start:
            continue
        title = candidate.get("clean_title", "")
        if _is_authorization_title(title):
            auth_candidate = candidate
            break
    if auth_candidate is None:
        return current_end
    return max(current_end, _find_section_end(candidates, auth_candidate, block_count))


def _copy_block_image_relationships(block, source_doc, target_doc) -> None:
    for blip in block.findall(".//" + qn("a:blip")):
        old_rid = blip.get(qn("r:embed"))
        if not old_rid:
            continue
        image_part = source_doc.part.related_parts.get(old_rid)
        if image_part is None:
            continue
        new_rid, _ = target_doc.part.get_or_add_image(BytesIO(image_part.blob))
        blip.set(qn("r:embed"), new_rid)
