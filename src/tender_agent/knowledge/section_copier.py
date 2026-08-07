"""Knowledge material copier utilities."""

from copy import deepcopy
import hashlib
import re
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

from sqlalchemy import or_
from tender_agent.knowledge.docx_ooxml import (
    append_body_block,
    inline_floating_drawings,
    paragraph_body_text_nodes,
    paragraph_visible_text,
    strip_page_break_marks,
    strip_section_properties,
)
from sqlalchemy.orm import Session

from data.knowledge.models import Certificate, TemplateSection


def _company_scope_filter(model: Any, company_id: str):
    company_id = str(company_id or "").strip()
    if not company_id:
        return model.scope == "shared"
    return or_(model.company_id == company_id, model.scope == "shared")


def copy_certificates_by_category(
    category: str,
    target_doc,
    db_session: Session,
    max_count: int = 10,
    title_level: int = 3,
    seen_file_fingerprints: Optional[Set[str]] = None,
    query_name: Optional[str] = None,
    performance_title_level: int = 0,
    company_id: str = "",
) -> int:
    """Copy current certificates by category into target doc."""
    if db_session is None:
        target_doc.add_paragraph(f"【证书库未连接：{category}】")
        return 0

    query = (
        db_session.query(Certificate)
        .filter(
            Certificate.category == category,
            Certificate.is_current == True,
            Certificate.deleted_at.is_(None),
            _company_scope_filter(Certificate, company_id),
        )
        .order_by(Certificate.created_at.desc())
    )
    if category == "审计报告":
        rows = _select_audit_report_group(query.all(), query_name or "")
        max_count = max(max_count, len(rows))
    elif category == "合作业绩":
        max_groups = max(1, int(max_count or 1))
        rows = _select_performance_material_groups(query.all(), max_groups)
        if performance_title_level > 0:
            return _copy_performance_material_groups(
                rows=rows,
                target_doc=target_doc,
                max_groups=max_groups,
                title_level=performance_title_level,
                seen_file_fingerprints=seen_file_fingerprints,
            )
        max_count = max(max_count, len(rows))
    else:
        rows = query.limit(max(max_count * 3, max_count)).all()
    if not rows:
        target_doc.add_paragraph(f"【未找到证书素材：{category}】")
        return 0

    heading_style = f"Heading {min(max(title_level, 1), 3)}" if title_level > 0 else ""
    copied = 0
    for cert in rows:
        if copied >= max_count:
            break
        file_path = _resolve_file_path(cert.file_path or "")
        fingerprint = _file_fingerprint(file_path)
        if fingerprint and seen_file_fingerprints is not None and fingerprint in seen_file_fingerprints:
            continue

        if heading_style:
            try:
                target_doc.add_paragraph(cert.name, style=heading_style)
            except Exception:
                target_doc.add_paragraph(cert.name)

        if not file_path.exists():
            target_doc.add_paragraph(f"【证书文件不存在：{cert.file_path}】")
            continue
        try:
            target_doc.add_picture(str(file_path), width=Inches(5.8))
            copied += 1
            if fingerprint and seen_file_fingerprints is not None:
                seen_file_fingerprints.add(fingerprint)
        except Exception as exc:
            target_doc.add_paragraph(f"【证书图片插入失败：{exc}】")
    return copied


def _copy_performance_material_groups(
    rows: list[Certificate],
    target_doc,
    max_groups: int,
    title_level: int,
    seen_file_fingerprints: Optional[Set[str]] = None,
) -> int:
    if not rows:
        target_doc.add_paragraph("【未找到证书素材：合作业绩】")
        return 0

    groups: Dict[str, list[Certificate]] = {}
    first_seen: Dict[str, int] = {}
    for index, row in enumerate(rows):
        group_key = _performance_group_key(row.name or "") or str(row.name or row.id or index)
        groups.setdefault(group_key, []).append(row)
        first_seen.setdefault(group_key, index)

    heading_style = f"Heading {min(max(title_level, 1), 3)}"
    copied = 0
    copied_groups = 0
    for group_key in sorted(groups.keys(), key=lambda key: first_seen[key]):
        if copied_groups >= max(1, int(max_groups or 1)):
            break
        group_rows = sorted(groups[group_key], key=lambda row: _performance_page_index(row.name or ""))
        visible_rows: list[Certificate] = []
        for cert in group_rows:
            file_path = _resolve_file_path(cert.file_path or "")
            fingerprint = _file_fingerprint(file_path)
            if fingerprint and seen_file_fingerprints is not None and fingerprint in seen_file_fingerprints:
                continue
            visible_rows.append(cert)
        if not visible_rows:
            continue

        heading = _performance_company_heading(group_key)
        try:
            target_doc.add_paragraph(heading, style=heading_style)
        except Exception:
            target_doc.add_paragraph(heading)

        copied_groups += 1
        for cert in visible_rows:
            file_path = _resolve_file_path(cert.file_path or "")
            fingerprint = _file_fingerprint(file_path)
            if not file_path.exists():
                target_doc.add_paragraph(f"【证书文件不存在：{cert.file_path}】")
                continue
            try:
                target_doc.add_picture(str(file_path), width=Inches(5.8))
                copied += 1
                if fingerprint and seen_file_fingerprints is not None:
                    seen_file_fingerprints.add(fingerprint)
            except Exception as exc:
                target_doc.add_paragraph(f"【证书图片插入失败：{exc}】")
    return copied


def _select_audit_report_group(rows: list[Certificate], query_name: str) -> list[Certificate]:
    """Audit reports are stored page-by-page; copy the whole requested year group."""
    if not rows:
        return []
    preferred_year = _audit_report_year(query_name)
    groups: Dict[str, list[Certificate]] = {}
    for row in rows:
        year = _audit_report_year(row.name or "")
        if year:
            groups.setdefault(year, []).append(row)
    if not groups:
        return rows
    year = preferred_year if preferred_year in groups else sorted(groups.keys(), reverse=True)[0]
    return sorted(groups[year], key=lambda row: _audit_report_page_index(row.name or ""))


def _select_performance_material_groups(rows: list[Certificate], max_groups: int) -> list[Certificate]:
    """Performance materials are page-level rows; copy complete company/material groups."""
    if not rows:
        return []
    groups: Dict[str, list[Certificate]] = {}
    first_seen: Dict[str, int] = {}
    for index, row in enumerate(rows):
        group_key = _performance_group_key(row.name or "")
        if not group_key:
            group_key = str(row.name or row.id or index)
        groups.setdefault(group_key, []).append(row)
        first_seen.setdefault(group_key, index)

    selected_keys = sorted(groups.keys(), key=lambda key: first_seen[key])[: max(1, int(max_groups or 1))]
    selected: list[Certificate] = []
    for key in selected_keys:
        selected.extend(sorted(groups[key], key=lambda row: _performance_page_index(row.name or "")))
    return selected


def _performance_group_key(name: str) -> str:
    """Only a trailing underscore-number is a page suffix; year ranges stay in the material name."""
    return re.sub(r"_\d+\s*$", "", str(name or "").strip())


def _performance_page_index(name: str) -> int:
    match = re.search(r"_(\d+)\s*$", str(name or ""))
    return int(match.group(1)) if match else 1


def _performance_company_heading(name: str) -> str:
    value = _performance_group_key(name)
    value = re.sub(r"[_\-—–]+\s*$", "", value).strip()
    value = re.sub(r"\s*(?:20\d{2}\s*年?\s*[-至到~—–]+\s*)?20\d{2}\s*年?\s*$", "", value).strip()
    value = re.sub(r"\s*20\d{2}\s*[-至到~—–]+\s*20\d{2}\s*$", "", value).strip()
    value = re.sub(r"\s*[（(]?\s*20\d{2}\s*[-至到~—–]+\s*20\d{2}\s*[）)]?\s*$", "", value).strip()
    return value or _performance_group_key(name) or "合作业绩证明材料"


def _audit_report_year(text: str) -> str:
    match = re.search(r"(20\d{2})\s*年", str(text or ""))
    return match.group(1) if match else ""


def _audit_report_page_index(name: str) -> int:
    match = re.search(r"_(\d+)\s*$", str(name or ""))
    return int(match.group(1)) if match else 1


def extract_section_text_from_master(
    section: TemplateSection,
    master_path: str,
    max_chars: int = 6000,
) -> str:
    """Extract plain text for a template section using its block range."""
    if section is None or section.start_block_idx is None or section.end_block_idx is None:
        return ""
    master = _resolve_master_docx_path(master_path)
    if not master.exists():
        return ""

    source_doc = Document(str(master))
    blocks = list(source_doc.element.body.iterchildren())
    start = max(section.start_block_idx - 1, 0)
    end = min(section.end_block_idx, len(blocks))
    parts: list[str] = []
    char_count = 0
    for block in blocks[start:end]:
        text = _block_plain_text(block)
        if not text:
            continue
        parts.append(text)
        char_count += len(text)
        if char_count >= max_chars:
            break
    return "\n".join(parts).strip()[:max_chars]


def _block_plain_text(block) -> str:
    tag = block.tag.rsplit("}", 1)[-1]
    if tag == "p":
        return paragraph_visible_text(block)
    if tag == "tbl":
        rows: list[str] = []
        for tr in block.findall(".//" + qn("w:tr")):
            cells: list[str] = []
            for tc in tr.findall(".//" + qn("w:tc")):
                text = "".join(node.text or "" for node in tc.iter() if node.tag.endswith("}t")).strip()
                if text:
                    cells.append(text)
            if cells:
                rows.append(" | ".join(cells))
        return "\n".join(rows).strip()
    return ""


def _resolve_master_docx_path(master_path: str) -> Path:
    """Resolve a knowledge master docx even if the caller passed mojibake."""
    raw = Path(str(master_path or ""))
    candidates = [raw]
    if not raw.is_absolute():
        candidates.extend(
            [
                Path.cwd() / raw,
                Path(__file__).resolve().parents[3] / raw,
            ]
        )
    for candidate in candidates:
        if candidate.exists() and candidate.suffix.lower() == ".docx":
            return candidate

    master_dir_candidates = []
    if str(raw.parent) not in {"", "."}:
        master_dir_candidates.append(raw.parent)
        if not raw.parent.is_absolute():
            master_dir_candidates.extend(
                [
                    Path.cwd() / raw.parent,
                    Path(__file__).resolve().parents[3] / raw.parent,
                ]
            )
    master_dir_candidates.extend(
        [
            Path(__file__).resolve().parents[3] / "data/knowledge/master",
            Path.cwd() / "data/knowledge/master",
        ]
    )

    seen_dirs: set[str] = set()
    for master_dir in master_dir_candidates:
        key = str(master_dir.resolve()) if master_dir.exists() else str(master_dir)
        if key in seen_dirs or not master_dir.exists():
            continue
        seen_dirs.add(key)
        docx_files = [path for path in master_dir.glob("*.docx") if not path.name.startswith("~$")]
        if not docx_files:
            continue
        technical = [path for path in docx_files if "技术" in path.name]
        if technical:
            return sorted(technical, key=lambda item: item.name)[0]
        non_business = [path for path in docx_files if "商务" not in path.name]
        if non_business:
            return sorted(non_business, key=lambda item: item.name)[0]
        return sorted(docx_files, key=lambda item: item.name)[0]

    return raw


def _resolve_existing_docx_path(file_path: str) -> Optional[Path]:
    if not str(file_path or "").strip():
        return None
    raw = Path(str(file_path or ""))
    candidates = [raw]
    if not raw.is_absolute():
        candidates.extend(
            [
                Path.cwd() / raw,
                Path(__file__).resolve().parents[3] / raw,
            ]
        )
    for candidate in candidates:
        if candidate.exists() and candidate.suffix.lower() == ".docx":
            return candidate
    return None


def copy_section_from_master(
    chapter_id: str,
    master_path: str,
    target_doc,
    db_session: Session,
    heading_base_level: int = 2,
    parent_number: str = "",
    max_heading_depth: int = 1,
    company_id: str = "",
    section_id: str = "",
) -> int:
    """Copy a chapter block range from technical master docx."""
    if db_session is None:
        target_doc.add_paragraph(f"【技术章节库未连接：{chapter_id}】")
        return 0

    query = (
        db_session.query(TemplateSection)
        .filter(
            TemplateSection.is_current == True,
            TemplateSection.deleted_at.is_(None),
            _company_scope_filter(TemplateSection, company_id),
        )
    )
    if section_id:
        query = query.filter(TemplateSection.id == section_id)
    else:
        query = query.filter(TemplateSection.chapter_id == chapter_id)
    rows: List[TemplateSection] = query.order_by(TemplateSection.start_block_idx.asc()).all()
    section = next(
        (row for row in rows if company_id and str(getattr(row, "company_id", "") or "") == company_id),
        rows[0] if rows else None,
    )
    if section is None:
        target_doc.add_paragraph(f"【未找到技术章节：{chapter_id}】")
        return 0

    source_file_path = str((getattr(section, "metadata_info", {}) or {}).get("source_file_path") or "")
    master = _resolve_existing_docx_path(source_file_path) or _resolve_master_docx_path(master_path)
    if not master.exists():
        target_doc.add_paragraph(f"【技术母版不存在：{master_path}】")
        return 0
    if section.start_block_idx is None or section.end_block_idx is None:
        target_doc.add_paragraph(f"【技术章节缺少块索引：{chapter_id}】")
        return 0

    source_doc = Document(str(master))
    blocks = list(source_doc.element.body.iterchildren())
    start = max(section.start_block_idx - 1, 0)
    end = min(section.end_block_idx, len(blocks))
    copied = 0
    heading_state: Dict[str, Any] = {
        "source_base_level": int(section.level or 1),
        "source_heading_style_levels": _source_heading_style_levels(source_doc),
        "parent_number": str(parent_number or "").strip(),
        "max_heading_depth": max(0, int(max_heading_depth or 0)),
        "counters": {},
    }
    for block in blocks[start:end]:
        if block.tag.endswith("sectPr"):
            continue
        copied_block = deepcopy(block)
        if not _clean_copied_block(copied_block, heading_base_level, heading_state, target_doc):
            continue
        _copy_block_image_relationships(copied_block, source_doc, target_doc)
        append_body_block(target_doc, copied_block)
        copied += 1

    if copied == 0:
        target_doc.add_paragraph(f"【技术章节为空：{chapter_id} {section.title}】")
    return copied



def copy_section_range_from_master(
    chapter_start: str,
    chapter_end: str,
    master_path: str,
    target_doc,
    db_session: Session,
    heading_base_level: int = 2,
    parent_number: str = "",
    max_heading_depth: int = 1,
    company_id: str = "",
) -> int:
    """Copy direct technical master chapters in a chapter_id range."""
    if db_session is None:
        target_doc.add_paragraph(f"【技术章节库未连接：{chapter_start}-{chapter_end}】")
        return 0

    start_parts = _chapter_id_parts(chapter_start)
    end_parts = _chapter_id_parts(chapter_end)
    if not start_parts or not end_parts or len(start_parts) != len(end_parts):
        target_doc.add_paragraph(f"【技术章节范围无效：{chapter_start}-{chapter_end}】")
        return 0

    prefix = ".".join(str(part) for part in start_parts[:-1])
    rows: List[TemplateSection] = (
        db_session.query(TemplateSection)
        .filter(
            TemplateSection.is_current == True,
            TemplateSection.deleted_at.is_(None),
            _company_scope_filter(TemplateSection, company_id),
        )
        .order_by(TemplateSection.start_block_idx.asc())
        .all()
    )
    sections = [
        row
        for row in rows
        if _chapter_id_in_direct_range(str(row.chapter_id or ""), prefix, start_parts, end_parts)
    ]
    if not sections:
        target_doc.add_paragraph(f"【未找到技术章节范围：{chapter_start}-{chapter_end}】")
        return 0

    master = _resolve_master_docx_path(master_path)
    if not master.exists():
        target_doc.add_paragraph(f"【技术母版不存在：{master_path}】")
        return 0

    source_doc = Document(str(master))
    blocks = list(source_doc.element.body.iterchildren())
    source_base_level = min(int(section.level or 1) for section in sections)
    heading_state: Dict[str, Any] = {
        "source_base_level": source_base_level,
        "source_heading_style_levels": _source_heading_style_levels(source_doc),
        "parent_number": str(parent_number or "").strip(),
        "max_heading_depth": max(0, int(max_heading_depth or 0)),
        "counters": {},
    }

    copied = 0
    for section in sections:
        if section.start_block_idx is None or section.end_block_idx is None:
            continue
        if _looks_like_technical_toc_section(section, blocks):
            continue
        start = max(section.start_block_idx - 1, 0)
        end = min(section.end_block_idx, len(blocks))
        for block in blocks[start:end]:
            if block.tag.endswith("sectPr"):
                continue
            copied_block = deepcopy(block)
            if not _clean_copied_block(copied_block, heading_base_level, heading_state, target_doc):
                continue
            _copy_block_image_relationships(copied_block, source_doc, target_doc)
            append_body_block(target_doc, copied_block)
            copied += 1

    if copied == 0:
        target_doc.add_paragraph(f"【技术章节范围为空：{chapter_start}-{chapter_end}】")
    return copied


def _looks_like_technical_toc_section(section: TemplateSection, blocks: List[Any]) -> bool:
    """Skip table-of-contents/index pages inside technical master files."""
    title = re.sub(r"\s+", "", str(section.title or ""))
    if title in {"目录", "技术文件目录", "技术方案目录", "技术响应文件目录"}:
        return True
    if section.start_block_idx is None or section.end_block_idx is None:
        return False

    start = max(section.start_block_idx - 1, 0)
    end = min(section.end_block_idx, len(blocks))
    lines: List[str] = []
    for block in blocks[start:end]:
        text = _block_plain_text(block)
        if text:
            lines.extend(line.strip() for line in text.splitlines() if line.strip())
        if len(lines) >= 30:
            break

    if not lines:
        return False
    compact_lines = [re.sub(r"\s+", "", line) for line in lines[:30]]
    if compact_lines and compact_lines[0] in {"目录", "技术文件目录", "技术方案目录", "技术响应文件目录"}:
        return True
    toc_like = 0
    for line in compact_lines:
        if re.search(r"(?:\.{2,}|…{2,}|·{2,})\d{1,3}$", line):
            toc_like += 1
        elif re.search(r"第?\d+页$", line) and len(line) <= 45:
            toc_like += 1
        elif re.match(r"^\d+(?:\.\d+){1,4}[\u4e00-\u9fa5A-Za-z].{0,32}\d{1,3}$", line):
            toc_like += 1
    return len(compact_lines) >= 5 and toc_like >= max(3, len(compact_lines) // 2)


def _chapter_id_parts(chapter_id: str) -> Tuple[int, ...]:
    parts = str(chapter_id or "").strip().split(".")
    if not parts or any(not part.isdigit() for part in parts):
        return ()
    return tuple(int(part) for part in parts)


def _chapter_id_in_direct_range(
    chapter_id: str,
    prefix: str,
    start_parts: Tuple[int, ...],
    end_parts: Tuple[int, ...],
) -> bool:
    value = str(chapter_id or "").strip()
    if not value:
        return False
    if prefix and not value.startswith(f"{prefix}."):
        return False
    parts = _chapter_id_parts(value)
    if len(parts) != len(start_parts):
        return False
    return start_parts <= parts <= end_parts

def _clean_copied_block(
    block,
    heading_base_level: int,
    heading_state: Optional[Dict[str, Any]] = None,
    target_doc=None,
) -> bool:
    if _block_contains_raw_ooxml_text(block):
        return False
    strip_section_properties(block)
    strip_page_break_marks(block)
    inline_floating_drawings(block)
    _normalize_copied_heading(block, heading_base_level, heading_state, target_doc)
    for outline_level in block.findall(".//" + qn("w:outlineLvl")):
        parent = outline_level.getparent()
        if parent is not None:
            parent.remove(outline_level)
    for num_pr in block.findall(".//" + qn("w:numPr")):
        parent = num_pr.getparent()
        if parent is not None:
            parent.remove(num_pr)
    return True


def _block_contains_raw_ooxml_text(block) -> bool:
    text = "".join(node.text or "" for node in paragraph_body_text_nodes(block))
    if not text:
        return False
    compact = text.strip()
    if len(compact) < 12:
        return False
    return bool(re.search(r"(?:<|&lt;)/?(?:w|wp|a|r|mc|v|o):[A-Za-z]", compact))


def _normalize_copied_heading(
    block,
    heading_base_level: int,
    heading_state: Optional[Dict[str, Any]],
    target_doc,
) -> None:
    paragraphs = [block] if block.tag.endswith("}p") else list(block.findall(".//" + qn("w:p")))
    for paragraph in paragraphs:
        p_style = paragraph.find("./" + qn("w:pPr") + "/" + qn("w:pStyle"))
        if p_style is None:
            continue
        style_id = p_style.get(qn("w:val"))
        style_levels: Dict[str, int] = (heading_state or {}).get("source_heading_style_levels", {})
        source_level = style_levels.get(str(style_id or ""))
        if source_level is None:
            source_level = _heading_level_from_style_id(style_id, allow_numeric=False)
        if source_level is None:
            remove_element(p_style)
            continue
        if not paragraph_visible_text(paragraph):
            remove_element(p_style)
            continue

        state = heading_state or {}
        source_base_level = int(state.get("source_base_level") or source_level)
        depth = max(1, source_level - source_base_level + 1)
        max_depth = max(0, int(state.get("max_heading_depth") or 0))

        if max_depth and depth <= max_depth:
            counters: Dict[int, int] = state.setdefault("counters", {})
            counters[depth] = int(counters.get(depth, 0)) + 1
            for key in list(counters.keys()):
                if key > depth:
                    counters.pop(key, None)
            parent_number = str(state.get("parent_number") or "").strip()
            local_number = ".".join(str(counters[idx]) for idx in range(1, depth + 1) if counters.get(idx))
            heading_number = ".".join(part for part in (parent_number, local_number) if part)
            if heading_number:
                _prefix_block_text(paragraph, heading_number)
            target_level = min(max(int(heading_base_level or 2) + depth - 1, 1), 9)
            p_style.set(qn("w:val"), _target_heading_style_id(target_doc, target_level))
            continue

        remove_element(p_style)


def remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _source_heading_style_levels(source_doc: Document) -> Dict[str, int]:
    levels: Dict[str, int] = {}
    for style in source_doc.styles:
        style_id = str(getattr(style, "style_id", "") or "")
        style_name = str(getattr(style, "name", "") or "").strip()
        if not style_id or not style_name:
            continue
        level = _heading_level_from_style_name(style_name)
        if level is not None:
            levels[style_id] = level
    return levels


def _heading_level_from_style_name(style_name: str) -> Optional[int]:
    normalized = str(style_name or "").strip()
    lower = normalized.lower().replace(" ", "")
    match = re.match(r"heading(\d+)$", lower)
    if match:
        return int(match.group(1))
    match = re.match(r"标题\s*(\d+)$", normalized)
    if match:
        return int(match.group(1))
    return None


def _heading_level_from_style_id(style_id: Optional[str], allow_numeric: bool = True) -> Optional[int]:
    value = str(style_id or "").strip()
    if not value:
        return None
    normalized = value.lower().replace(" ", "")
    match = re.match(r"heading(\d+)$", normalized)
    if match:
        return int(match.group(1))
    match = re.match(r"标题(\d+)$", value)
    if match:
        return int(match.group(1))
    if allow_numeric and value.isdigit():
        numeric = int(value)
        if 3 <= numeric <= 11:
            return numeric - 2
    return None


def _target_heading_style_id(target_doc, level: int) -> str:
    if target_doc is not None:
        try:
            return target_doc.styles[f"Heading {level}"].style_id
        except Exception:
            pass
    return f"Heading{level}"


def _prefix_block_text(block, heading_number: str) -> None:
    prefix = f"{heading_number} "
    text_nodes = paragraph_body_text_nodes(block)
    if not text_nodes:
        return
    current = text_nodes[0].text or ""
    compact = current.strip()
    if compact.startswith(heading_number):
        return
    if re.match(r"^\d+(?:\.\d+)*\s+", compact):
        text_nodes[0].text = re.sub(r"^\s*\d+(?:\.\d+)*\s+", prefix, current, count=1)
    else:
        text_nodes[0].text = prefix + current


def _copy_block_image_relationships(block, source_doc, target_doc) -> None:
    for blip in block.findall(".//" + qn("a:blip")):
        old_rid = blip.get(qn("r:embed"))
        if not old_rid:
            continue
        image_part = source_doc.part.related_parts.get(old_rid)
        if image_part is None:
            continue
        new_rid, _ = target_doc.part.get_or_add_image(BytesIO(image_part.blob))
        blip.set(qn("r:embed"), new_rid)


def _resolve_file_path(file_path: str) -> Path:
    path = Path(file_path)
    if path.is_absolute():
        return path
    candidates = [
        Path.cwd() / path,
        Path(__file__).resolve().parents[3] / path,
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[0]


def _file_fingerprint(path: Path) -> str:
    if not path.exists() or not path.is_file():
        return ""
    try:
        stat = path.stat()
        digest = hashlib.md5(path.read_bytes()).hexdigest()
        return f"{digest}:{stat.st_size}"
    except Exception:
        return str(path.resolve()).lower()
