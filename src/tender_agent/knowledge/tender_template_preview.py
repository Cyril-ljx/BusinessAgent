"""Convert an anchored tender-template DOCX span into editable HTML."""

from __future__ import annotations

import html
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn


def _element_value(element: Any, child_tag: str, default: str = "") -> str:
    child = element.find(qn(child_tag)) if element is not None else None
    return str(child.get(qn("w:val")) if child is not None else default)


def _alpha_number(value: int, uppercase: bool = False) -> str:
    output = ""
    current = max(1, value)
    while current:
        current, remainder = divmod(current - 1, 26)
        output = chr(ord("A" if uppercase else "a") + remainder) + output
    return output


def _roman_number(value: int, uppercase: bool = False) -> str:
    pairs = (
        (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
    )
    current = max(1, value)
    output: list[str] = []
    for amount, symbol in pairs:
        while current >= amount:
            output.append(symbol)
            current -= amount
    result = "".join(output)
    return result if uppercase else result.lower()


def _chinese_number(value: int) -> str:
    if value <= 0:
        return "零"
    if value >= 10000:
        return str(value)
    digits = "零一二三四五六七八九"
    units = ("", "十", "百", "千")
    parts: list[str] = []
    zero_pending = False
    for position in range(3, -1, -1):
        base = 10 ** position
        digit = value // base
        value %= base
        if digit:
            if zero_pending and parts:
                parts.append("零")
            if not (position == 1 and digit == 1 and not parts):
                parts.append(digits[digit])
            parts.append(units[position])
            zero_pending = False
        elif parts and value:
            zero_pending = True
    return "".join(parts)


def _format_number(value: int, number_format: str) -> str:
    if number_format == "lowerLetter":
        return _alpha_number(value)
    if number_format == "upperLetter":
        return _alpha_number(value, uppercase=True)
    if number_format == "lowerRoman":
        return _roman_number(value)
    if number_format == "upperRoman":
        return _roman_number(value, uppercase=True)
    if number_format in {"chineseCounting", "chineseCountingThousand", "ideographTraditional"}:
        return _chinese_number(value)
    if number_format == "decimalZero" and value < 10:
        return f"0{value}"
    return str(value)


class _NumberingResolver:
    def __init__(self, document: Any) -> None:
        try:
            numbering = document.part.numbering_part.element
        except (KeyError, NotImplementedError):
            # Some valid DOCX producers use only literal numbering and omit
            # word/numbering.xml. The rest of the formatting is still usable.
            numbering = None
        self.abstracts = {
            str(item.get(qn("w:abstractNumId"))): item
            for item in (numbering.findall(qn("w:abstractNum")) if numbering is not None else [])
        }
        self.numbers = {
            str(item.get(qn("w:numId"))): item
            for item in (numbering.findall(qn("w:num")) if numbering is not None else [])
        }
        self.counters: dict[tuple[str, int], int] = {}

    @staticmethod
    def _paragraph_num_pr(paragraph: Any) -> Any:
        paragraph_properties = paragraph._p.pPr
        if paragraph_properties is not None and paragraph_properties.numPr is not None:
            return paragraph_properties.numPr
        style = getattr(paragraph, "style", None)
        while style is not None:
            style_properties = style._element.pPr
            if style_properties is not None and style_properties.numPr is not None:
                return style_properties.numPr
            style = getattr(style, "base_style", None)
        return None

    def _level(self, number_id: str, level_index: int) -> Any:
        number = self.numbers.get(number_id)
        abstract_id = _element_value(number, "w:abstractNumId")
        abstract = self.abstracts.get(abstract_id)
        if abstract is None:
            return None
        for level in abstract.findall(qn("w:lvl")):
            if int(level.get(qn("w:ilvl"), "0")) == level_index:
                return level
        return None

    def _start_value(self, number_id: str, level_index: int, level: Any) -> int:
        number = self.numbers.get(number_id)
        if number is not None:
            for override in number.findall(qn("w:lvlOverride")):
                if int(override.get(qn("w:ilvl"), "0")) == level_index:
                    value = _element_value(override, "w:startOverride")
                    if value.isdigit():
                        return int(value)
        value = _element_value(level, "w:start", "1")
        return int(value) if value.isdigit() else 1

    def next_label(self, paragraph: Any) -> str:
        number_properties = self._paragraph_num_pr(paragraph)
        if number_properties is None or number_properties.numId is None:
            return ""
        number_id = str(number_properties.numId.val)
        if number_id == "0":
            return ""
        level_index = int(number_properties.ilvl.val) if number_properties.ilvl is not None else 0
        level = self._level(number_id, level_index)
        if level is None:
            return ""
        counter_key = (number_id, level_index)
        start = self._start_value(number_id, level_index, level)
        self.counters[counter_key] = self.counters.get(counter_key, start - 1) + 1
        for key in list(self.counters):
            if key[0] == number_id and key[1] > level_index:
                self.counters.pop(key, None)

        label = _element_value(level, "w:lvlText", f"%{level_index + 1}.")
        for placeholder_level in range(9):
            placeholder = f"%{placeholder_level + 1}"
            if placeholder not in label:
                continue
            placeholder_definition = self._level(number_id, placeholder_level)
            placeholder_start = self._start_value(number_id, placeholder_level, placeholder_definition)
            value = self.counters.get((number_id, placeholder_level), placeholder_start)
            number_format = _element_value(placeholder_definition, "w:numFmt", "decimal")
            label = label.replace(placeholder, _format_number(value, number_format))
        return label


def freeze_numbering_labels(document: Any, kept_elements: list[Any]) -> None:
    """Turn Word automatic numbering into literal text before cropping a DOCX."""
    kept_ids = {id(element) for element in kept_elements}
    numbering = _NumberingResolver(document)
    paragraphs = {id(paragraph._element): paragraph for paragraph in document.paragraphs}

    for element in document.element.body.iterchildren():
        if not element.tag.endswith("}p"):
            continue
        paragraph = paragraphs.get(id(element))
        if paragraph is None or not paragraph.text.strip():
            continue
        label = numbering.next_label(paragraph)
        if not label or id(element) not in kept_ids:
            continue

        text_nodes = element.findall(".//" + qn("w:t"))
        if text_nodes:
            current = text_nodes[0].text or ""
            if not current.strip().startswith(label):
                text_nodes[0].text = f"{label}\t{current}"

        paragraph_properties = paragraph._p.get_or_add_pPr()
        number_properties = paragraph_properties.get_or_add_numPr()
        number_properties.get_or_add_numId().val = 0


def _anchor_number(value: Any) -> int | None:
    match = re.fullmatch(r"p(\d+)", str(value or "").strip())
    return int(match.group(1)) if match else None


def _points(value: Any) -> float | None:
    if value is None:
        return None
    try:
        return round(float(value.pt), 2)
    except (AttributeError, TypeError, ValueError):
        return None


def _paragraph_value(paragraph: Any, name: str) -> Any:
    value = getattr(paragraph.paragraph_format, name, None)
    if value is not None:
        return value
    style = getattr(paragraph, "style", None)
    while style is not None:
        value = getattr(style.paragraph_format, name, None)
        if value is not None:
            return value
        style = getattr(style, "base_style", None)
    return None


def _paragraph_css(paragraph: Any) -> str:
    declarations = ["margin: 0", "white-space: pre-wrap"]
    alignment = str(_paragraph_value(paragraph, "alignment") or "").lower()
    if "center" in alignment:
        declarations.append("text-align: center")
    elif "right" in alignment:
        declarations.append("text-align: right")
    elif "justify" in alignment or "distribute" in alignment:
        declarations.append("text-align: justify")
    for source_name, css_name in (
        ("first_line_indent", "text-indent"),
        ("left_indent", "margin-left"),
        ("right_indent", "margin-right"),
        ("space_before", "margin-top"),
        ("space_after", "margin-bottom"),
    ):
        points = _points(_paragraph_value(paragraph, source_name))
        if points is not None:
            declarations.append(f"{css_name}: {points:g}pt")
    line_spacing = _paragraph_value(paragraph, "line_spacing")
    spacing_points = _points(line_spacing)
    if spacing_points is not None:
        declarations.append(f"line-height: {spacing_points:g}pt")
    elif isinstance(line_spacing, (int, float)):
        declarations.append(f"line-height: {float(line_spacing):g}")
    return "; ".join(declarations)


def _preserve_spaces(text: str) -> str:
    escaped = html.escape(text or "")
    escaped = escaped.replace("\t", "&nbsp;&nbsp;&nbsp;&nbsp;")
    escaped = re.sub(r" {2,}", lambda item: "&nbsp;" * len(item.group(0)), escaped)
    return escaped.replace("\n", "<br>")


def _run_html(run: Any) -> str:
    content = _preserve_spaces(run.text or "")
    if not content:
        return ""
    styles: list[str] = []
    font_size = _points(run.font.size)
    if font_size is not None:
        styles.append(f"font-size: {font_size:g}pt")
    try:
        color = run.font.color.rgb
    except (AttributeError, ValueError):
        color = None
    if color is not None:
        styles.append(f"color: #{color}")
    run_properties = run._r.rPr
    run_fonts = run_properties.rFonts if run_properties is not None else None
    font_name = run.font.name or (run_fonts.get(qn("w:eastAsia")) if run_fonts is not None else None)
    if font_name:
        styles.append(f"font-family: '{html.escape(str(font_name), quote=True)}'")
    if styles:
        content = f'<span style="{"; ".join(styles)}">{content}</span>'
    if run.underline:
        content = f"<u>{content}</u>"
    if run.italic:
        content = f"<em>{content}</em>"
    if run.bold:
        content = f"<strong>{content}</strong>"
    return content


def _paragraph_html(paragraph: Any, numbering_label: str = "") -> str:
    body = "".join(_run_html(run) for run in paragraph.runs)
    if not body:
        body = _preserve_spaces(paragraph.text or "")
    if numbering_label:
        body = f'<span data-word-numbering="true">{_preserve_spaces(numbering_label)}&nbsp;</span>{body}'
    return f'<p style="{_paragraph_css(paragraph)}">{body or "<br>"}</p>'


def _table_html(table: Any) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells: list[str] = []
        for cell in row.cells:
            content = "".join(_paragraph_html(paragraph) for paragraph in cell.paragraphs)
            cells.append(f'<td style="border: 1px solid #334155; padding: 6px 8px; vertical-align: middle">{content}</td>')
        rows.append(f"<tr>{''.join(cells)}</tr>")
    return '<table style="width: 100%; border-collapse: collapse; table-layout: fixed"><tbody>' + "".join(rows) + "</tbody></table>"


def tender_template_html(source_path: str | Path, anchor_start: str, anchor_end: str) -> str:
    """Return the same non-empty DOCX blocks used by the source-template renderer."""
    start = _anchor_number(anchor_start)
    end = _anchor_number(anchor_end)
    path = Path(source_path)
    if start is None or end is None or end < start or not path.exists() or path.suffix.lower() != ".docx":
        return ""

    document = Document(str(path))
    numbering = _NumberingResolver(document)
    paragraphs = {paragraph._element: paragraph for paragraph in document.paragraphs}
    tables = {table._element: table for table in document.tables}
    blocks: list[tuple[str, Any]] = []
    for element in document.element.body.iterchildren():
        if element.tag.endswith("}p"):
            paragraph = paragraphs.get(element)
            if paragraph is not None and paragraph.text.strip():
                blocks.append(("paragraph", (paragraph, numbering.next_label(paragraph))))
        elif element.tag.endswith("}tbl"):
            table = tables.get(element)
            if table is not None:
                blocks.append(("table", table))
    if start >= len(blocks):
        return ""

    rendered: list[str] = []
    for kind, block in blocks[start : min(end, len(blocks) - 1) + 1]:
        if kind == "paragraph":
            paragraph, numbering_label = block
            rendered.append(_paragraph_html(paragraph, numbering_label))
        else:
            rendered.append(_table_html(block))
    return "".join(rendered)
