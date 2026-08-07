"""Material-source renderers used by the DOCX renderer."""

from __future__ import annotations

import hashlib
from pathlib import Path
from typing import Any, Dict, List, Optional, Set

from docx.shared import Inches

from data.knowledge.models import TemplateSection
from tender_agent.knowledge.section_copier import (
    copy_certificates_by_category,
    copy_section_from_master,
    copy_section_range_from_master,
)
from tender_agent.rendering.certificate_rendering_policy import (
    cap_certificate_count,
    certificate_item_heading_level,
    supporting_certificate_heading,
)


def flatten_materials(assignments: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    flat_materials: List[Dict[str, Any]] = []
    for assignment in assignments:
        mats = assignment.get("materials")
        if isinstance(mats, list):
            flat_materials.extend([m for m in mats if isinstance(m, dict)])
        elif isinstance(assignment, dict):
            flat_materials.append(assignment)
    return flat_materials


def _material_source(material: Dict[str, Any]) -> str:
    source = str(material.get("source", "")).strip()
    if source == "knowledge_certificate":
        return "certificate"
    if source == "knowledge_tech_section":
        return "tech_section"
    return source


def has_tech_section_material(assignments: List[Dict[str, Any]]) -> bool:
    return any(
        (
            _material_source(material) == "tech_section"
            and "." in str(material.get("chapter_id") or "")
        )
        or _material_source(material) == "tech_section_range"
        for material in flatten_materials(assignments)
    )


def render_structured_materials(
    doc,
    node: Dict[str, Any],
    assignments: List[Dict[str, Any]],
    db_session: Any,
    tech_master_path: Optional[str],
    level: int,
    render_context: Optional[Dict[str, Any]] = None,
    suppress_headings: bool = False,
) -> bool:
    materials = flatten_materials(assignments)
    if not materials:
        return False

    context = render_context or {
        "certificate_categories": set(),
        "file_fingerprints": set(),
        "tech_sections": set(),
        "company_id": "",
    }
    if suppress_headings:
        context = dict(context)
        context["suppress_material_headings"] = True
    copied = False
    for material in materials[:8]:
        source = _material_source(material)
        renderer = MATERIAL_RENDERERS.get(source)
        if renderer is None:
            continue
        if renderer(doc, node, material, db_session, tech_master_path, level, context):
            copied = True
    return copied


def copy_template_supporting_certificates_as_subsections(
    doc,
    node: Dict[str, Any],
    assignments: List[Dict[str, Any]],
    db_session: Any,
    level: int,
    render_context: Optional[Dict[str, Any]] = None,
) -> bool:
    materials = [material for material in flatten_materials(assignments) if _material_source(material) == "certificate"]
    if not materials or db_session is None:
        return False
    context = render_context or {
        "certificate_categories": set(),
        "file_fingerprints": set(),
        "tech_sections": set(),
        "company_id": "",
    }
    copied_any = False
    for material in materials[:6]:
        category = str(material.get("category") or "").strip()
        if not category:
            continue
        max_count = cap_certificate_count(category, int(material.get("max_count") or 10))
        if max_count <= 0:
            continue
        if category in context["certificate_categories"]:
            continue
        performance_title_level = min(level + 1, 3) if category == "合作业绩" else 0
        item_title_level = certificate_item_heading_level(node, category, max_count, level)
        if category != "合作业绩" and item_title_level <= 0:
            heading = supporting_certificate_heading(str(node.get("name") or ""), category)
            try:
                doc.add_paragraph(heading, style=f"Heading {min(level + 1, 3)}")
            except Exception:
                doc.add_paragraph(heading)
        copied_count = copy_certificates_by_category(
            category=category,
            target_doc=doc,
            db_session=db_session,
            max_count=max_count,
            title_level=item_title_level,
            seen_file_fingerprints=context["file_fingerprints"],
            query_name=str(node.get("name") or ""),
            performance_title_level=performance_title_level,
            company_id=str(context.get("company_id") or ""),
        )
        if copied_count > 0:
            context["certificate_categories"].add(category)
            copied_any = True
    return copied_any


def _render_tech_section_range(
    doc,
    node: Dict[str, Any],
    material: Dict[str, Any],
    db_session: Any,
    tech_master_path: Optional[str],
    level: int,
    context: Dict[str, Any],
) -> bool:
    chapter_start = str(material.get("chapter_start") or "").strip()
    chapter_end = str(material.get("chapter_end") or "").strip()
    range_key = f"{chapter_start}-{chapter_end}"
    if not (chapter_start and chapter_end and tech_master_path and db_session is not None):
        return False
    if range_key in context["tech_sections"]:
        return True
    max_heading_depth = 0 if context.get("suppress_material_headings") else 3
    copied_count = copy_section_range_from_master(
        chapter_start=chapter_start,
        chapter_end=chapter_end,
        master_path=tech_master_path,
        target_doc=doc,
        db_session=db_session,
        heading_base_level=min(level + 1, 4),
        parent_number=str(node.get("id") or ""),
        max_heading_depth=max_heading_depth,
        company_id=str(context.get("company_id") or ""),
    )
    if copied_count > 0:
        context["tech_sections"].add(range_key)
        return True
    return False


def _render_tech_section(
    doc,
    node: Dict[str, Any],
    material: Dict[str, Any],
    db_session: Any,
    tech_master_path: Optional[str],
    level: int,
    context: Dict[str, Any],
) -> bool:
    chapter_id = str(material.get("chapter_id") or "")
    section_id = str(material.get("section_id") or "")
    if not (chapter_id and tech_master_path and db_session is not None):
        return False
    copy_full_section = bool(material.get("copy_full_section"))
    if "." not in chapter_id and not copy_full_section:
        return False
    if not copy_full_section and not section_id:
        chapter_id = _resolve_safe_tech_chapter_id(
            chapter_id,
            str(node.get("name", "")),
            db_session,
            company_id=str(context.get("company_id") or ""),
        )
    if not chapter_id:
        return False
    section_key = section_id or chapter_id
    if section_key in context["tech_sections"]:
        return True
    max_heading_depth = 0 if context.get("suppress_material_headings") else 3
    copied_count = copy_section_from_master(
        chapter_id=chapter_id,
        master_path=tech_master_path,
        target_doc=doc,
        db_session=db_session,
        heading_base_level=min(level + 1, 4),
        parent_number=str(node.get("id") or ""),
        max_heading_depth=max_heading_depth,
        company_id=str(context.get("company_id") or ""),
        section_id=section_id,
    )
    if copied_count > 0:
        context["tech_sections"].add(section_key)
        return True
    return False


def _render_certificate(
    doc,
    node: Dict[str, Any],
    material: Dict[str, Any],
    db_session: Any,
    tech_master_path: Optional[str],
    level: int,
    context: Dict[str, Any],
) -> bool:
    if material.get("file_path"):
        return append_single_material_attachment(doc, material, context["file_fingerprints"])
    if not material.get("category"):
        return False
    category = str(material.get("category"))
    max_count = cap_certificate_count(category, int(material.get("max_count") or 10))
    if max_count <= 0:
        return False
    performance_title_level = min(level + 1, 3) if category == "合作业绩" else 0
    copied_count = copy_certificates_by_category(
        category=category,
        target_doc=doc,
        db_session=db_session,
        max_count=max_count,
        title_level=certificate_item_heading_level(node, category, max_count, level),
        seen_file_fingerprints=context["file_fingerprints"],
        query_name=str(node.get("name") or ""),
        performance_title_level=performance_title_level,
        company_id=str(context.get("company_id") or ""),
    )
    if copied_count > 0:
        context["certificate_categories"].add(category)
        return True
    return False


MATERIAL_RENDERERS = {
    "tech_section_range": _render_tech_section_range,
    "tech_section": _render_tech_section,
    "certificate": _render_certificate,
}


def _resolve_safe_tech_chapter_id(chapter_id: str, node_name: str, db_session: Any, company_id: str = "") -> str:
    section = _get_template_section(db_session, chapter_id, company_id=company_id)
    if section is None:
        return chapter_id
    span = _section_span(section)
    if span <= 180:
        return chapter_id

    replacement = _best_child_section(db_session, chapter_id, node_name, company_id=company_id)
    return replacement.chapter_id if replacement is not None and replacement.chapter_id else ""


def _knowledge_scope_filter(company_id: str):
    from sqlalchemy import or_

    company_id = str(company_id or "").strip()
    if not company_id:
        return TemplateSection.scope == "shared"
    return or_(TemplateSection.company_id == company_id, TemplateSection.scope == "shared")


def _get_template_section(db_session: Any, chapter_id: str, company_id: str = "") -> Optional[TemplateSection]:
    try:
        return (
            db_session.query(TemplateSection)
            .filter(
                TemplateSection.chapter_id == chapter_id,
                TemplateSection.is_current == True,
                TemplateSection.deleted_at.is_(None),
                _knowledge_scope_filter(company_id),
            )
            .order_by(TemplateSection.start_block_idx.asc())
            .first()
        )
    except Exception:
        return None


def _best_child_section(db_session: Any, chapter_id: str, node_name: str, company_id: str = "") -> Optional[TemplateSection]:
    try:
        rows = (
            db_session.query(TemplateSection)
            .filter(
                TemplateSection.chapter_id.like(f"{chapter_id}.%"),
                TemplateSection.is_current == True,
                TemplateSection.deleted_at.is_(None),
                _knowledge_scope_filter(company_id),
            )
            .order_by(TemplateSection.start_block_idx.asc())
            .all()
        )
    except Exception:
        return None

    candidates = [row for row in rows if 0 < _section_span(row) <= 180]
    if not candidates:
        return None

    preferred = _preferred_tech_title(node_name)
    best: Optional[TemplateSection] = None
    best_score = -10_000
    for row in candidates:
        text = f"{row.title or ''} {row.full_path or ''}"
        score = _tech_text_score(node_name, text)
        if preferred and preferred in text:
            score += 100
        score -= max(_section_span(row) - 80, 0) // 20
        if score > best_score:
            best = row
            best_score = score
    return best


def _preferred_tech_title(node_name: str) -> str:
    rules = [
        ("运营管理", "企业项目管理体系"),
        ("管理梯队", "企业项目管理体系"),
        ("人员履历", "员工管理准则"),
        ("招聘", "人员招聘方案"),
        ("高峰", "服务质量保证"),
        ("保障", "项目保障体系"),
        ("应急", "突发事件应急预案"),
        ("风险", "风险管控制度"),
        ("项目实施方案", "项目整体规划"),
        ("实施方案书", "项目整体规划"),
        ("实质性内容响应", "项目整体规划"),
        ("项目理解", "项目整体规划"),
        ("服务/货物清单", "项目整体规划"),
        ("服务模式", "项目整体规划"),
        ("团队介绍", "项目整体规划"),
        ("服务方案", "项目整体规划"),
        ("技术方案", "项目整体规划"),
        ("整体方案", "项目整体规划"),
    ]
    for keyword, title in rules:
        if keyword in node_name:
            return title
    return ""


def _tech_text_score(query: str, text: str) -> int:
    return sum(1 for token in _lookup_tokens(query) if token and token in text)


def _lookup_tokens(text: str) -> List[str]:
    cleaned = "".join(ch if "\u4e00" <= ch <= "\u9fff" else " " for ch in text or "")
    tokens: List[str] = []
    for part in cleaned.split():
        tokens.append(part)
        for size in (2, 3, 4):
            for idx in range(0, max(len(part) - size + 1, 0)):
                tokens.append(part[idx : idx + size])
    return list(dict.fromkeys(tokens))


def _section_span(section: TemplateSection) -> int:
    return max(0, int(section.end_block_idx or 0) - int(section.start_block_idx or 0))


def append_material_attachments(
    doc,
    assignments: List[Dict[str, Any]],
    seen_file_fingerprints: Optional[Set[str]] = None,
) -> None:
    materials = flatten_materials(assignments)
    image_count = 0
    for material in materials[:8]:
        if not material.get("file_path"):
            continue
        if image_count == 0:
            doc.add_paragraph("素材原件：")
        if append_single_material_attachment(doc, material, seen_file_fingerprints):
            image_count += 1


def append_single_material_attachment(
    doc,
    material: Dict[str, Any],
    seen_file_fingerprints: Optional[Set[str]] = None,
) -> bool:
    file_path = str(material.get("file_path") or "").strip()
    if not file_path:
        return False
    path = _resolve_material_path(file_path)
    if not path.exists():
        doc.add_paragraph(f"[素材文件未找到: {file_path}]")
        return False
    fingerprint = _file_fingerprint(path)
    if fingerprint and seen_file_fingerprints is not None and fingerprint in seen_file_fingerprints:
        return False
    if path.suffix.lower() in {".png", ".jpg", ".jpeg", ".bmp", ".gif"}:
        doc.add_paragraph(str(material.get("name") or path.name))
        try:
            doc.add_picture(str(path), width=Inches(5.8))
            if fingerprint and seen_file_fingerprints is not None:
                seen_file_fingerprints.add(fingerprint)
            return True
        except Exception:
            doc.add_paragraph(f"[图片插入失败: {path}]")
            return False
    doc.add_paragraph(f"素材文件：{material.get('name') or path.name}（{path}）")
    return True


def _resolve_material_path(file_path: str) -> Path:
    path = Path(file_path)
    if path.is_absolute():
        return path
    candidates = [
        Path.cwd() / path,
        Path(__file__).resolve().parents[3] / path,
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[0]


def _file_fingerprint(path: Path) -> str:
    if not path.exists() or not path.is_file():
        return ""
    try:
        stat = path.stat()
        digest = hashlib.md5(path.read_bytes()).hexdigest()
        return f"{digest}:{stat.st_size}"
    except Exception:
        return str(path.resolve()).lower()
