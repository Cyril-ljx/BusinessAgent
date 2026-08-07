"""Render generated rich text, markdown tables, and simple HTML into DOCX."""

from __future__ import annotations

import base64
import re
from dataclasses import dataclass, field
from html.parser import HTMLParser
from io import BytesIO
from typing import Dict, Iterable, List, Optional, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


_HTML_TAG_RE = re.compile(
    r"<\s*(p|div|h[1-6]|table|thead|tbody|tr|td|th|img|ul|ol|li|strong|b|em|i|u|br)\b",
    re.IGNORECASE,
)


@dataclass
class HtmlNode:
    tag: Optional[str] = None
    attrs: Dict[str, str] = field(default_factory=dict)
    children: List[Union["HtmlNode", str]] = field(default_factory=list)


class _SimpleHtmlParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.root = HtmlNode(tag="root")
        self.stack: List[HtmlNode] = [self.root]

    def handle_starttag(self, tag: str, attrs: List[tuple[str, Optional[str]]]) -> None:
        node = HtmlNode(tag=tag.lower(), attrs={key.lower(): value or "" for key, value in attrs})
        self.stack[-1].children.append(node)
        if node.tag not in {"br", "img", "hr", "meta", "link", "input"}:
            self.stack.append(node)

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        for index in range(len(self.stack) - 1, 0, -1):
            if self.stack[index].tag == tag:
                del self.stack[index:]
                return

    def handle_data(self, data: str) -> None:
        if data:
            self.stack[-1].children.append(data)


def _clean_generated_line(line: str) -> str:
    text = line.strip()
    while text.startswith("#"):
        text = text[1:].strip()
    return _remove_json_leakage_from_render_line(text)


def _remove_json_leakage_from_render_line(text: str) -> str:
    """Last-resort guard: never render raw requirement JSON into DOCX."""
    if not text:
        return ""
    markers = (
        '[{',
        '{"',
        '"requirement"',
        '"quote"',
        '"condition"',
        '"severity"',
        '"level"',
        '"level":"P0"',
        '"level": "P0"',
    )
    if not any(marker in text for marker in markers):
        return text

    prefix = ""
    match = re.match(r"^(\s*[-*•]?\s*[^:：{]{1,40}[:：])\s*", text)
    if match:
        prefix = match.group(1).strip()

    cleaned = re.sub(r"\[\s*\{.*?\}\s*\]", "详见招标文件对应要求", text)
    cleaned = re.sub(r"\{\s*\"(?:requirement|quote|condition|severity|level|value|name)\".*?\}", "详见招标文件对应要求", cleaned)
    cleaned = re.sub(r"\{[^{}]*(?:\"quote\"|\"severity\"|\"level\"|\"condition\"|\"requirement\")[^{}]*\}", "详见招标文件对应要求", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()

    if any(marker in cleaned for marker in markers):
        return f"{prefix} 详见招标文件对应要求。" if prefix else ""
    return cleaned


def _append_generated_text(doc: Document, generated: str) -> None:
    if _looks_like_html(generated):
        _append_html(doc, generated)
        return
    _append_plain_or_markdown(doc, generated)


def _looks_like_html(value: str) -> bool:
    return bool(value and _HTML_TAG_RE.search(value))


def _append_plain_or_markdown(doc: Document, generated: str) -> None:
    table_buffer: List[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            if not _append_markdown_table(doc, table_buffer):
                for raw in table_buffer:
                    text = _clean_generated_line(raw)
                    if text:
                        doc.add_paragraph(text)
            table_buffer = []

    for line in generated.splitlines():
        if _is_markdown_table_line(line):
            table_buffer.append(line)
            continue
        flush_table()
        text = _clean_generated_line(line)
        if text:
            doc.add_paragraph(text)
    flush_table()


def _append_html(doc: Document, html: str) -> None:
    parser = _SimpleHtmlParser()
    parser.feed(html or "")
    inline_buffer: List[Union[HtmlNode, str]] = []

    def flush_inline() -> None:
        nonlocal inline_buffer
        if inline_buffer:
            paragraph = doc.add_paragraph()
            for child in inline_buffer:
                _render_inline(paragraph, child)
            inline_buffer = []

    for child in parser.root.children:
        if isinstance(child, str):
            if child.strip():
                inline_buffer.append(child)
            continue
        if _is_block_node(child):
            flush_inline()
            _render_block(doc, child)
        else:
            inline_buffer.append(child)
    flush_inline()


def _is_block_node(node: HtmlNode) -> bool:
    return node.tag in {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "table", "ul", "ol", "li"}


def _render_block(doc: Document, node: HtmlNode) -> None:
    tag = node.tag or ""
    if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        level = min(int(tag[1]), 4)
        paragraph = doc.add_heading(level=level)
        _apply_paragraph_style(paragraph, node.attrs.get("style", ""))
        _render_inline_children(paragraph, node.children, bold=True)
        return
    if tag in {"p", "div", "li"}:
        if not _node_has_renderable_content(node):
            return
        style = "List Bullet" if tag == "li" else None
        paragraph = doc.add_paragraph(style=style)
        _apply_paragraph_style(paragraph, node.attrs.get("style", ""))
        _render_inline_children(paragraph, node.children)
        return
    if tag in {"ul", "ol"}:
        list_style = "List Number" if tag == "ol" else "List Bullet"
        for item in _iter_child_nodes(node, {"li"}):
            paragraph = doc.add_paragraph(style=list_style)
            _render_inline_children(paragraph, item.children)
        return
    if tag == "table":
        _render_html_table(doc, node)


def _css_declarations(value: str) -> Dict[str, str]:
    declarations: Dict[str, str] = {}
    for item in str(value or "").split(";"):
        name, separator, raw_value = item.partition(":")
        if separator and name.strip() and raw_value.strip():
            declarations[name.strip().lower()] = raw_value.strip()
    return declarations


def _css_points(value: str) -> Optional[float]:
    match = re.fullmatch(r"(-?\d+(?:\.\d+)?)pt", str(value or "").strip(), re.IGNORECASE)
    return float(match.group(1)) if match else None


def _apply_paragraph_style(paragraph, style: str) -> None:
    declarations = _css_declarations(style)
    alignment = declarations.get("text-align", "").lower()
    alignment_value = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(alignment)
    if alignment_value is not None:
        paragraph.alignment = alignment_value
    for css_name, property_name in (
        ("text-indent", "first_line_indent"),
        ("margin-left", "left_indent"),
        ("margin-right", "right_indent"),
        ("margin-top", "space_before"),
        ("margin-bottom", "space_after"),
    ):
        points = _css_points(declarations.get(css_name, ""))
        if points is not None:
            setattr(paragraph.paragraph_format, property_name, Pt(points))
    line_height = declarations.get("line-height", "")
    points = _css_points(line_height)
    if points is not None:
        paragraph.paragraph_format.line_spacing = Pt(points)
    elif line_height:
        try:
            paragraph.paragraph_format.line_spacing = float(line_height)
        except ValueError:
            pass


def _render_inline_children(
    paragraph,
    children: Iterable[Union[HtmlNode, str]],
    *,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    font_size: Optional[float] = None,
    color: Optional[str] = None,
) -> None:
    for child in children:
        _render_inline(
            paragraph,
            child,
            bold=bold,
            italic=italic,
            underline=underline,
            font_size=font_size,
            color=color,
        )


def _render_inline(
    paragraph,
    item: Union[HtmlNode, str],
    *,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    font_size: Optional[float] = None,
    color: Optional[str] = None,
) -> None:
    if isinstance(item, str):
        text = item.replace("\xa0", " ")
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = bold or None
        run.italic = italic or None
        run.underline = underline or None
        if font_size is not None:
            run.font.size = Pt(font_size)
        if color and re.fullmatch(r"#[0-9a-fA-F]{6}", color):
            run.font.color.rgb = RGBColor.from_string(color[1:])
        return

    tag = item.tag or ""
    if tag == "br":
        paragraph.add_run().add_break()
        return
    if tag == "img":
        _add_data_image(paragraph, item.attrs.get("src", ""))
        return
    declarations = _css_declarations(item.attrs.get("style", ""))
    next_bold = bold or tag in {"strong", "b"} or declarations.get("font-weight", "").lower() in {"bold", "600", "700", "800", "900"}
    next_italic = italic or tag in {"em", "i"} or declarations.get("font-style", "").lower() == "italic"
    next_underline = underline or tag == "u" or "underline" in declarations.get("text-decoration", "").lower()
    next_font_size = _css_points(declarations.get("font-size", "")) or font_size
    next_color = declarations.get("color", "") if re.fullmatch(r"#[0-9a-fA-F]{6}", declarations.get("color", "")) else color
    if tag == "table":
        # Inline tables are rendered after the current paragraph by the block path;
        # contentEditable can create nested tables, so keep a text fallback here.
        text = _node_text(item).strip()
        if text:
            run = paragraph.add_run(text)
            run.bold = next_bold or None
            run.italic = next_italic or None
            run.underline = next_underline or None
        return
    _render_inline_children(
        paragraph,
        item.children,
        bold=next_bold,
        italic=next_italic,
        underline=next_underline,
        font_size=next_font_size,
        color=next_color,
    )


def _add_data_image(paragraph, src: str) -> None:
    if not src.startswith("data:image/") or ";base64," not in src:
        paragraph.add_run("[图片无法导出：仅支持本地插入图片]")
        return
    try:
        payload = src.split(",", 1)[1]
        image_stream = BytesIO(base64.b64decode(payload))
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(5.8))
    except Exception:
        paragraph.add_run("[图片导出失败]")


def _render_html_table(doc: Document, table_node: HtmlNode) -> bool:
    rows = []
    for row_node in _iter_descendant_nodes(table_node, {"tr"}):
        cells = [_node_text(cell).strip() for cell in _iter_child_nodes(row_node, {"td", "th"})]
        if cells:
            rows.append(cells)
    return append_rows_table(doc, rows)


def _iter_child_nodes(node: HtmlNode, tags: set[str]) -> Iterable[HtmlNode]:
    for child in node.children:
        if isinstance(child, HtmlNode) and child.tag in tags:
            yield child


def _iter_descendant_nodes(node: HtmlNode, tags: set[str]) -> Iterable[HtmlNode]:
    for child in node.children:
        if not isinstance(child, HtmlNode):
            continue
        if child.tag in tags:
            yield child
        yield from _iter_descendant_nodes(child, tags)


def _node_text(node: Union[HtmlNode, str]) -> str:
    if isinstance(node, str):
        return node.replace("\xa0", " ")
    if node.tag == "br":
        return "\n"
    if node.tag == "img":
        return ""
    return "".join(_node_text(child) for child in node.children)


def _node_has_renderable_content(node: HtmlNode) -> bool:
    if _node_text(node).strip():
        return True
    return any(isinstance(child, HtmlNode) and child.tag == "img" for child in node.children)


def _is_markdown_table_line(line: str) -> bool:
    text = line.strip()
    return text.startswith("|") and text.endswith("|") and text.count("|") >= 2


def _split_markdown_table_row(line: str) -> List[str]:
    text = _clean_generated_line(line).strip()
    if text.startswith("|"):
        text = text[1:]
    if text.endswith("|"):
        text = text[:-1]
    return [cell.strip() for cell in text.split("|")]


def _is_markdown_separator_row(cells: List[str]) -> bool:
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip() or "---") for cell in cells)


def _append_markdown_table(doc: Document, lines: List[str]) -> bool:
    rows: List[List[str]] = []
    for line in lines:
        cells = _split_markdown_table_row(line)
        if _is_markdown_separator_row(cells):
            continue
        if any(cell for cell in cells):
            rows.append(cells)
    return append_rows_table(doc, rows)


def append_rows_table(doc: Document, rows: List[List[str]]) -> bool:
    if not rows:
        return False
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    _mark_generated_table(table._tbl)
    for row_index, row in enumerate(rows):
        for col_index in range(col_count):
            table.cell(row_index, col_index).text = row[col_index] if col_index < len(row) else ""
    _fit_table_to_page(table)
    return True


def _fit_all_tables_to_page(doc: Document) -> None:
    """Final guard: keep generated tables within page width."""
    for table in doc.tables:
        if _is_generated_table(table._tbl):
            _fit_table_to_page(table)


def _mark_generated_table(tbl) -> None:
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    marker = OxmlElement("w:tblDescription")
    marker.set(qn("w:val"), "ngu-generated-table")
    tbl_pr.append(marker)


def _is_generated_table(tbl) -> bool:
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        return False
    return any(
        child.get(qn("w:val")) == "ngu-generated-table"
        for child in tbl_pr.findall(qn("w:tblDescription"))
    )


def _fit_table_to_page(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    for tag in ("w:tblInd", "w:tblLayout"):
        for child in list(tbl_pr.findall(qn(tag))):
            tbl_pr.remove(child)

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit")
    tbl_pr.append(layout)

    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is not None:
        tbl.remove(tbl_grid)

    for cell_width in list(tbl.findall(".//" + qn("w:tcW"))):
        parent = cell_width.getparent()
        if parent is not None:
            parent.remove(cell_width)


append_markdown_table = _append_markdown_table
is_markdown_table_line = _is_markdown_table_line
append_generated_text = _append_generated_text
fit_generated_tables_to_page = _fit_all_tables_to_page
