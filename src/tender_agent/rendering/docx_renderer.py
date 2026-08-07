"""DOCX renderer for bid documents."""

from pathlib import Path
from typing import Any, Dict, List, Optional, Set
import shutil
from copy import deepcopy
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from tender_agent.knowledge.tender_template_copier import copy_tender_template_by_node
from tender_agent.knowledge.section_copier import copy_section_from_master
from tender_agent.knowledge.docx_ooxml import (
    block_has_page_break as _block_has_page_break,
    block_text as _block_text,
)
from tender_agent.rendering.material_renderers import (
    append_material_attachments as _append_material_attachments,
    copy_template_supporting_certificates_as_subsections as _copy_template_supporting_certificates_as_subsections,
    render_structured_materials as _copy_structured_materials,
)
from tender_agent.rendering.docx_safety import (
    remove_raw_ooxml_text_paragraphs as _remove_raw_ooxml_text_paragraphs,
)
from tender_agent.rendering.generated_content import (
    append_generated_text as _append_generated_text,
    fit_generated_tables_to_page as _fit_all_tables_to_page,
)

from tender_agent.rendering.render_plan import (
    build_render_plan,
    build_todo_items,
    prepare_render_outline,
    resolve_node_render_decision,
)
from tender_agent.rendering.tender_template_renderers import (
    copy_pdf_tender_template_simple as _copy_pdf_tender_template_simple,
    copy_tender_template_simple as _copy_tender_template_simple,
)


def render_blank_bid(
    outline: List[Dict[str, Any]],
    title_info: Dict[str, str],
    master_template_path: str,
    output_path: str,
    company_name: str = "投标单位：",
    submission_date: Optional[str] = None,
    material_assignments: Optional[List[Dict[str, Any]]] = None,
    generated_sections: Optional[Dict[str, str]] = None,
    db_session: Any = None,
    tech_master_path: str = "data/knowledge/master/技术文件.docx",
    tender_doc_path: Optional[str] = None,
    company_id: str = "",
) -> str:
    """Render a bid DOCX from template and outline."""
    if submission_date is None:
        from datetime import datetime

        now = datetime.now()
        submission_date = f"{now.year}年{now.month:02d}月{now.day:02d}日"

    shutil.copy(master_template_path, output_path)
    doc = Document(output_path)
    material_assignments = material_assignments if material_assignments is not None else []

    render_outline = prepare_render_outline(outline)
    doc_type = _infer_doc_type(title_info, render_outline)
    header_references = _collect_header_references(doc)
    placeholders = {
        "{{ purchaser_name }}": title_info.get("purchaser", "采购方"),
        "{{ project_name }}": title_info.get("project_name", title_info.get("title", "项目名称")),
        "{{ doc_type }}": doc_type,
        "{{ company_name }}": _format_company_name(company_name, doc_type),
        "{{ submission_date }}": submission_date,
    }
    _replace_placeholders(doc, placeholders)
    _trim_master_template_tail(doc, list(placeholders.values()))
    _restore_header_references(doc, header_references)
    _ensure_page_break_at_end(doc)
    _insert_toc_page_content(doc, render_outline)
    _append_outline_chapters(
        doc,
        render_outline,
        material_assignments=material_assignments,
        generated_sections=generated_sections or {},
        db_session=db_session,
        tech_master_path=tech_master_path,
        tender_doc_path=tender_doc_path,
        company_id=company_id,
    )
    if os.getenv("INCLUDE_TODO_SUMMARY", "false").lower() in {"1", "true", "yes", "on"}:
        todo_plan = build_render_plan(render_outline, material_assignments or [], generated_sections or {})
        _append_todo_summary(doc, todo_plan)
    _enable_auto_update_fields(doc)
    _fit_all_tables_to_page(doc)
    _remove_raw_ooxml_text_paragraphs(doc)

    doc.save(output_path)
    return str(Path(output_path).absolute())


def _replace_placeholders(doc: Document, placeholders: Dict[str, str]) -> None:
    for para in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in para.text:
                _replace_in_paragraph(para, placeholder, value)


def _replace_in_paragraph(paragraph, placeholder: str, value: str) -> None:
    if placeholder not in paragraph.text:
        return
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, value)
            return
    full_text = paragraph.text
    new_text = full_text.replace(placeholder, value)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def _infer_doc_type(title_info: Dict[str, str], outline: List[Dict[str, Any]]) -> str:
    explicit = str(title_info.get("doc_type") or title_info.get("document_type") or "").strip()
    if explicit:
        return explicit
    names = _collect_render_outline_names(outline)
    compact = re.sub(r"\s+", "", "".join(names))
    if "报价文件" in compact or "报价书" in compact or "报价函" in compact or "报价人" in compact:
        return "报价文件"
    if "响应文件" in compact or "响应函" in compact:
        return "响应文件"
    if "投标文件" in compact or "投标函" in compact:
        return "投标文件"
    return "投标文件"


def _collect_render_outline_names(outline: List[Dict[str, Any]]) -> List[str]:
    names: List[str] = []
    for node in outline:
        name = str(node.get("name") or "")
        if name:
            names.append(name)
        names.extend(_collect_render_outline_names(node.get("children", []) or []))
    return names


def _format_company_name(company_name: str, doc_type: str) -> str:
    company = str(company_name or "").strip()
    if re.match(r"^(投标单位|投标人|报价人|供应商)[:：]", company):
        return company
    label = "报价人" if "报价" in str(doc_type or "") else "投标单位"
    return f"{label}：{company}" if company else f"{label}："


def _collect_header_references(doc: Document) -> List[Any]:
    references: List[Any] = []
    for sect_pr in doc.element.body.findall(".//" + qn("w:sectPr")):
        for header_ref in sect_pr.findall(qn("w:headerReference")):
            references.append(deepcopy(header_ref))
    return references


def _restore_header_references(doc: Document, header_references: List[Any]) -> None:
    if not header_references:
        return
    sect_pr = doc.element.body.sectPr
    if sect_pr is None:
        sect_pr = OxmlElement("w:sectPr")
        doc.element.body.append(sect_pr)

    existing = {
        (
            ref.get(qn("w:type")) or "default",
            ref.get(qn("r:id")) or "",
        )
        for ref in sect_pr.findall(qn("w:headerReference"))
    }
    insert_at = 0
    for header_ref in header_references:
        key = (
            header_ref.get(qn("w:type")) or "default",
            header_ref.get(qn("r:id")) or "",
        )
        if key in existing:
            continue
        sect_pr.insert(insert_at, deepcopy(header_ref))
        existing.add(key)
        insert_at += 1


def _ensure_page_break_at_end(doc: Document) -> None:
    body = doc.element.body
    content_blocks = [child for child in body.iterchildren() if not child.tag.endswith("}sectPr")]
    if content_blocks and _block_has_page_break(content_blocks[-1]):
        return

    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    run.append(page_break)
    paragraph.append(run)

    sect_pr = body.sectPr
    if sect_pr is not None:
        body.insert(body.index(sect_pr), paragraph)
    else:
        body.append(paragraph)


def _trim_master_template_tail(doc: Document, marker_values: List[str]) -> None:
    body = doc.element.body
    blocks = list(body.iterchildren())
    last_front_idx: Optional[int] = None
    for idx, block in enumerate(blocks):
        if not block.tag.endswith("}p"):
            continue
        text = _block_text(block)
        if any(value and value in text for value in marker_values):
            last_front_idx = idx

    if last_front_idx is None:
        return

    for block in blocks[last_front_idx + 1 :]:
        if block.tag.endswith("sectPr"):
            continue
        body.remove(block)


def _insert_toc_page_content(doc: Document, outline: List[Dict[str, Any]]) -> None:
    toc_title = doc.add_paragraph("目  录")
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if toc_title.runs:
        toc_title.runs[0].bold = True
        toc_title.runs[0].font.size = Pt(16)
    _append_toc_field(doc)


def _append_toc_field(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    # Keep the printed TOC focused on the authoritative bid outline. Grafted
    # knowledge-base headings may use Heading 4 for Word navigation, but they
    # should not crowd the formal directory page.
    instr.text = r'TOC \o "1-3" \h \z'
    run._r.append(instr)

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_separate)

    placeholder = OxmlElement("w:t")
    placeholder.text = "请在 Word 中更新目录"
    run._r.append(placeholder)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _flatten_outline_for_toc(outline: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []

    def walk(nodes: List[Dict[str, Any]]) -> None:
        for node in nodes:
            items.append(node)
            walk(node.get("children", []) or [])

    walk(outline)
    return items


def _enable_auto_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)
    else:
        update_fields.set(qn("w:val"), "true")


def _append_outline_chapters(
    doc: Document,
    outline: List[Dict[str, Any]],
    material_assignments: Optional[List[Dict[str, Any]]] = None,
    generated_sections: Optional[Dict[str, str]] = None,
    db_session: Any = None,
    tech_master_path: Optional[str] = None,
    tender_doc_path: Optional[str] = None,
    company_id: str = "",
) -> None:
    render_plan = build_render_plan(outline, material_assignments or [], generated_sections or {})
    render_context = {
        "certificate_categories": set(),
        "file_fingerprints": set(),
        "tech_sections": set(),
        "company_id": company_id,
        "rendered_template_assignments": [],
    }

    for top_node in outline:
        if _is_toc_node(top_node):
            continue
        _ensure_page_break_at_end(doc)
        template_package_copied = _add_chapter(
            doc,
            top_node,
            level=1,
            db_session=db_session,
            tech_master_path=tech_master_path,
            tender_doc_path=tender_doc_path,
            render_context=render_context,
            render_plan=render_plan,
        )
        if template_package_copied:
            continue
        _append_child_chapters(
            doc,
            top_node.get("children", []) or [],
            level=2,
            db_session=db_session,
            tech_master_path=tech_master_path,
            tender_doc_path=tender_doc_path,
            render_context=render_context,
            render_plan=render_plan,
        )


    _merge_rendered_template_assignments(
        material_assignments or [],
        render_context.get("rendered_template_assignments") or [],
    )

def _append_child_chapters(
    doc: Document,
    nodes: List[Dict[str, Any]],
    level: int,
    db_session: Any = None,
    tech_master_path: Optional[str] = None,
    tender_doc_path: Optional[str] = None,
    render_context: Optional[Dict[str, Any]] = None,
    render_plan: Optional[Dict[str, Any]] = None,
) -> None:
    visible_nodes = [node for node in nodes if not _is_toc_node(node)]
    for index, node in enumerate(visible_nodes):
        # Parent headings can share a page with their first child, but sibling
        # material points should start on separate pages.
        if index > 0:
            _ensure_page_break_at_end(doc)
        template_package_copied = _add_chapter(
            doc,
            node,
            level=level,
            db_session=db_session,
            tech_master_path=tech_master_path,
            tender_doc_path=tender_doc_path,
            render_context=render_context,
            render_plan=render_plan,
        )
        if template_package_copied:
            continue
        _append_child_chapters(
            doc,
            node.get("children", []) or [],
            level=level + 1,
            db_session=db_session,
            tech_master_path=tech_master_path,
            tender_doc_path=tender_doc_path,
            render_context=render_context,
            render_plan=render_plan,
        )



def _add_chapter(
    doc: Document,
    node: Dict[str, Any],
    level: int,
    db_session: Any = None,
    tech_master_path: Optional[str] = None,
    tender_doc_path: Optional[str] = None,
    render_context: Optional[Dict[str, Set[str]]] = None,
    render_plan: Optional[Dict[str, Any]] = None,
) -> bool:
    node_id = str(node.get("id", "")).strip()
    name = str(node.get("name", "")).strip()
    title_text = f"{node_id} {name}".strip() if node_id else name
    if render_plan is None:
        render_plan = build_render_plan([node], [], {})
    decision = resolve_node_render_decision(render_plan, node)
    assigned = decision["assignments"]

    style_name = f"Heading {level}" if level <= 3 else "Heading 3"
    try:
        heading_paragraph = doc.add_paragraph(title_text, style=style_name)
    except Exception:
        heading_paragraph = doc.add_paragraph(title_text)
        if heading_paragraph.runs:
            heading_paragraph.runs[0].bold = True
    if decision.get("compact_heading"):
        _compact_copied_template_heading(heading_paragraph)

    if node.get("children"):
        if tender_doc_path and decision.get("template_package_step") == "official_template_package":
            try:
                if copy_tender_template_by_node(node, tender_doc_path, doc):
                    _record_rendered_tender_template(render_context, node, "template_package")
                    _copy_template_supporting_certificates_as_subsections(
                        doc,
                        node,
                        decision.get("template_package_supporting_assignments") or [],
                        db_session,
                        level,
                        render_context,
                    )
                    return True
            except Exception as exc:
                doc.add_paragraph(f"[招标文件范本复制失败: {str(exc)[:120]}]")
        return False

    generated = decision["generated"]

    render_hook = node.get("render_hook") if isinstance(node.get("render_hook"), dict) else {}
    if (
        not generated
        and render_hook
        and str(render_hook.get("type") or "") == "tech_section"
        and str(render_hook.get("copy_mode") or "") == "docx_block"
    ):
        chapter_id = str(render_hook.get("chapter_id") or "").strip()
        if chapter_id and tech_master_path and db_session is not None:
            try:
                copy_section_from_master(
                    chapter_id=chapter_id,
                    master_path=tech_master_path,
                    target_doc=doc,
                    db_session=db_session,
                    heading_base_level=min(level + 1, 4),
                    parent_number=str(node.get("id") or ""),
                    max_heading_depth=0,
                    company_id=str((render_context or {}).get("company_id") or ""),
                )
                return
            except Exception as exc:
                doc.add_paragraph(f"[技术母版复制失败: {str(exc)[:120]}]")
                return
        doc.add_paragraph("[技术母版钩子缺少 chapter_id 或数据库连接，请人工检查。]")
        return

    # Highest priority: copy official forms/templates from the uploaded tender file.
    if tender_doc_path and decision.get("template_fallback_steps"):
        template_copied = False
        template_node = _node_with_template_evidence(node, assigned)
        template_copy_method = "template_leaf"
        for step in decision.get("template_fallback_steps") or []:
            try:
                if _try_template_fallback_step(step, template_node, tender_doc_path, doc):
                    template_copied = True
                    if step == "official_template" and template_node.get("anchor_start"):
                        template_copy_method = "located_section"
                    break
            except Exception as exc:
                if step == "official_template":
                    doc.add_paragraph(f"[招标文件范本复制失败: {str(exc)[:120]}]")
        if template_copied:
            _record_rendered_tender_template(render_context, node, template_copy_method)
            if decision.get("copy_generated_after_template"):
                _append_generated_text(doc, generated)
            if decision.get("append_supporting_materials_after_template", True):
                _copy_template_supporting_certificates_as_subsections(
                    doc,
                    node,
                    assigned,
                    db_session,
                    level,
                    render_context,
                )
            if decision.get("structured_material_phase") == "after_template":
                _copy_structured_materials(
                    doc,
                    node,
                    assigned,
                    db_session,
                    tech_master_path,
                    level,
                    render_context,
                    suppress_headings=True,
                )
            return

    if decision["has_tender_template_material"]:
        if generated:
            _append_generated_text(doc, generated)
        else:
            doc.add_paragraph(str(decision.get("template_required_placeholder") or ""))
        if decision.get("append_supporting_materials_after_template", True):
            _copy_template_supporting_certificates_as_subsections(
                doc,
                node,
                assigned,
                db_session,
                level,
                render_context,
            )
        return

    if decision.get("structured_material_phase") == "before_generated":
        if _copy_structured_materials(doc, node, assigned, db_session, tech_master_path, level, render_context):
            return

    if generated:
        _append_generated_text(doc, generated)
        if decision.get("structured_material_phase") == "after_generated":
            _copy_structured_materials(doc, node, assigned, db_session, tech_master_path, level, render_context)
        return

    if decision.get("structured_material_phase") == "fallback":
        if _copy_structured_materials(doc, node, assigned, db_session, tech_master_path, level, render_context):
            return

    if assigned and decision["is_manual_only_assignment"]:
        for line in decision.get("manual_placeholders") or []:
            if str(line or "").strip():
                doc.add_paragraph(str(line).strip())
        return

    if assigned:
        doc.add_paragraph("已匹配素材：")
        for line in decision.get("material_summary_lines") or []:
            doc.add_paragraph(str(line))
        seen_files = (render_context or {}).get("file_fingerprints") if render_context else None
        _append_material_attachments(doc, assigned, seen_files)
        doc.add_paragraph(str(decision.get("material_summary_placeholder") or ""))
        return

    doc.add_paragraph(str(decision.get("empty_placeholder") or ""))


def _try_template_fallback_step(step: str, node: Dict[str, Any], tender_doc_path: str, doc: Document) -> bool:
    if step == "official_template":
        return bool(copy_tender_template_by_node(node, tender_doc_path, doc))
    if step == "pdf_template":
        return bool(_copy_pdf_tender_template_simple(node, tender_doc_path, doc))
    if step == "docx_template":
        return bool(_copy_tender_template_simple(node, tender_doc_path, doc))
    return False


def _node_with_template_evidence(
    node: Dict[str, Any],
    assignments: List[Dict[str, Any]],
) -> Dict[str, Any]:
    enriched = dict(node)
    for assignment in assignments or []:
        for material in assignment.get("materials") or []:
            if not isinstance(material, dict) or str(material.get("source") or "") != "tender_template":
                continue
            for key in (
                "source_section_id",
                "source_anchor",
                "anchor_start",
                "anchor_end",
                "copy_method",
            ):
                value = material.get(key)
                if value:
                    enriched[key] = value
            return enriched
    return enriched


def _record_rendered_tender_template(
    render_context: Optional[Dict[str, Any]],
    node: Dict[str, Any],
    copy_method: str,
) -> None:
    if render_context is None:
        return
    records = render_context.setdefault("rendered_template_assignments", [])

    def append_record(item: Dict[str, Any], method: str) -> None:
        node_id = str(item.get("_source_id") or item.get("id") or "").strip()
        node_name = str(item.get("name") or "").strip()
        if not node_id and not node_name:
            return
        records.append(
            {
                "node_id": node_id,
                "node_name": node_name,
                "materials": [
                    {
                        "source": "tender_template",
                        "name": node_name or node_id,
                        "render_status": "copied",
                        "resolved_by": "renderer_fallback",
                        "copy_method": method,
                        "note": "渲染时已从招标原文复制范本",
                    }
                ],
            }
        )

    if copy_method == "template_package":
        for child in node.get("children") or []:
            if isinstance(child, dict):
                append_record(child, "template_package_child")
    else:
        append_record(node, copy_method)


def _merge_rendered_template_assignments(
    material_assignments: List[Dict[str, Any]],
    rendered_assignments: List[Dict[str, Any]],
) -> None:
    if not rendered_assignments:
        return

    def assignment_key(item: Dict[str, Any]) -> str:
        return str(item.get("node_id") or item.get("outline_node_id") or "").strip()

    by_id = {assignment_key(item): item for item in material_assignments if isinstance(item, dict) and assignment_key(item)}
    for rendered in rendered_assignments:
        if not isinstance(rendered, dict):
            continue
        node_id = assignment_key(rendered)
        if not node_id:
            continue
        rendered_material = (rendered.get("materials") or [{}])[0]
        existing = by_id.get(node_id)
        if existing is None:
            material_assignments.append(rendered)
            by_id[node_id] = rendered
            continue
        materials = existing.setdefault("materials", [])
        updated = False
        for material in materials:
            if isinstance(material, dict) and str(material.get("source") or "") == "tender_template":
                material.update({k: v for k, v in rendered_material.items() if v})
                updated = True
        if not updated:
            materials.append(rendered_material)
        existing["node_name"] = existing.get("node_name") or rendered.get("node_name")




def _is_toc_node(node: Dict[str, Any]) -> bool:
    name = str(node.get("name", "")).replace(" ", "")
    return name in {"目录", "投标文件目录", "响应文件目录"}


def _compact_copied_template_heading(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0
    fmt.keep_with_next = False




def _append_todo_summary(doc: Document, render_plan: Dict[str, Any]) -> None:
    _ensure_page_break_at_end(doc)
    doc.add_paragraph("待补充内容清单", style="Heading 1")
    doc.add_paragraph("以下章节仍需人工处理：")

    todo_items = build_todo_items(render_plan)
    if not todo_items:
        doc.add_paragraph("- 无")
        return

    for item in todo_items:
        p = doc.add_paragraph()
        p.add_run(f"- {item.get('id', '')} {item.get('name', '')}").bold = True
