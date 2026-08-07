"""Tender source template copy helpers for DOCX rendering."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from tender_agent.knowledge.docx_ooxml import (
    append_body_block as _append_body_block,
    fit_copied_tender_template_tables_to_page,
)
from tender_agent.knowledge.tender_template_copier import normalize_copied_tender_template_block
from tender_agent.rendering.generated_content import (
    append_markdown_table as _append_markdown_table,
    append_rows_table as _append_rows_table,
    is_markdown_table_line as _is_markdown_table_line,
)
from tender_agent.rendering.render_plan import (
    pdf_table_fallback_allowed as _pdf_table_fallback_allowed_for_node,
    pdf_text_template_fallback_allowed as _pdf_text_template_fallback_allowed_for_node,
    template_search_tokens as _template_search_tokens,
)

def _copy_pdf_tender_template_simple(node: Dict[str, Any], tender_pdf_path: str, target_doc: Document) -> bool:
    source_path = Path(tender_pdf_path)
    if not source_path.exists() or source_path.suffix.lower() != ".pdf":
        return False
    node_name = str(node.get("name", ""))
    if not (
        _pdf_table_fallback_allowed_for_node(node_name)
        or _pdf_text_template_fallback_allowed_for_node(node_name)
    ):
        return False
    tokens = _template_search_tokens(node_name) or [_norm_pdf_text(node_name)]
    tokens = [token for token in tokens if token]
    if not tokens:
        return False

    try:
        import fitz
    except Exception:
        return False

    try:
        pdf = fitz.open(str(source_path))
    except Exception:
        return False
    try:
        candidates = _pdf_template_heading_candidates(pdf, node_name, tokens)
        for candidate in candidates:
            if _pdf_table_fallback_allowed_for_node(node_name):
                rows = _pdf_first_table_after_heading(pdf, candidate, node_name)
                if rows and _append_pdf_rows_table(target_doc, rows):
                    return True
            if _pdf_text_template_fallback_allowed_for_node(node_name):
                lines = _pdf_text_template_lines_after_heading(pdf, candidate, node_name)
                if lines and _append_pdf_text_template(target_doc, lines):
                    return True
    finally:
        try:
            pdf.close()
        except Exception:
            pass
    return False


def _pdf_template_heading_candidates(pdf, node_name: str, tokens: List[str]) -> List[Dict[str, Any]]:
    target = _norm_pdf_text(node_name)
    candidates: List[Dict[str, Any]] = []
    for page_index in range(len(pdf)):
        page = pdf[page_index]
        for line in _pdf_text_lines(page):
            text = _norm_pdf_text(line.get("text", ""))
            if not text:
                continue
            if not _pdf_heading_matches_template(text, target, tokens):
                continue
            score = page_index * 10
            if target and (text == target or target in text):
                score += 1000
            if _looks_like_template_heading(text):
                score += 200
            candidates.append({"page_index": page_index, "y": float(line.get("y") or 0), "text": text, "score": score})
    # Prefer later exact/template-looking occurrences; directory entries usually appear earlier and have no table below.
    return sorted(candidates, key=lambda item: item.get("score", 0), reverse=True)


def _pdf_text_lines(page) -> List[Dict[str, Any]]:
    lines: List[Dict[str, Any]] = []
    try:
        data = page.get_text("dict") or {}
    except Exception:
        return lines
    for block in data.get("blocks", []) or []:
        for line in block.get("lines", []) or []:
            spans = line.get("spans", []) or []
            text = "".join(str(span.get("text") or "") for span in spans).strip()
            if not text:
                continue
            bbox = line.get("bbox") or [0, 0, 0, 0]
            lines.append({"text": text, "y": bbox[1] if len(bbox) > 1 else 0})
    return sorted(lines, key=lambda item: float(item.get("y") or 0))


def _pdf_heading_matches_template(text: str, target: str, tokens: List[str]) -> bool:
    text = _norm_pdf_text(text)
    target = _norm_pdf_text(target)
    normalized_tokens = [_norm_pdf_text(token) for token in tokens]
    if target and (text == target or target in text or text in target):
        return True
    return any(token and token in text for token in normalized_tokens)


def _pdf_first_table_after_heading(pdf, candidate: Dict[str, Any], node_name: str) -> List[List[str]]:
    start_page = int(candidate.get("page_index") or 0)
    start_y = float(candidate.get("y") or 0)
    target = _norm_pdf_text(node_name)
    for page_index in range(start_page, min(len(pdf), start_page + 3)):
        page = pdf[page_index]
        tables = _pdf_page_tables(page)
        if not tables:
            if page_index > start_page and _pdf_has_other_template_heading_before_y(page, target, 10**9):
                return []
            continue
        for table in tables:
            bbox = table.get("bbox") or [0, 0, 0, 0]
            table_y = float(bbox[1] if len(bbox) > 1 else 0)
            min_heading_y = min(start_y, table_y) if page_index == start_page else 0
            max_heading_y = max(start_y, table_y) if page_index == start_page else table_y
            if _pdf_has_other_template_heading_before_y(page, target, max_heading_y, min_y=min_heading_y):
                continue
            rows = _normalize_pdf_table_rows(table.get("rows") or [])
            if rows and _pdf_table_rows_match_node(rows, node_name):
                return rows
    return []


def _pdf_page_tables(page) -> List[Dict[str, Any]]:
    try:
        finder = page.find_tables()
        tables = getattr(finder, "tables", []) or []
    except Exception:
        return []
    result: List[Dict[str, Any]] = []
    for table in tables:
        try:
            rows = table.extract()
        except Exception:
            continue
        result.append({"bbox": list(getattr(table, "bbox", []) or []), "rows": rows})
    return sorted(result, key=lambda item: float((item.get("bbox") or [0, 0])[1] if len(item.get("bbox") or []) > 1 else 0))


def _pdf_has_other_template_heading_before_y(page, target: str, y_limit: float, min_y: float = 0) -> bool:
    for line in _pdf_text_lines(page):
        y = float(line.get("y") or 0)
        if y <= min_y + 2 or y >= y_limit:
            continue
        text = _norm_pdf_text(line.get("text", ""))
        if not text or text == target or target in text:
            continue
        if len(text) <= 35 and _looks_like_template_heading(text):
            return True
    return False


def _pdf_table_rows_match_node(rows: List[List[str]], node_name: str) -> bool:
    text = _norm_pdf_text("".join("".join(row) for row in rows))
    if not text:
        return False
    name = _norm_pdf_text(node_name)
    if not _pdf_table_fallback_allowed_for_node(name):
        return False
    # Review/scoring tables may mention a form name, but they are not the form itself.
    if any(
        token in text
        for token in (
            "评审点要求概况",
            "评审点具体描述",
            "审查内容",
            "评审内容",
            "评审因素",
            "评审标准",
            "评分标准",
            "评分细则",
            "分值构成",
            "商务部分",
            "技术部分",
            "报价得分",
            "商务分",
            "技术分",
            "价格分",
        )
    ):
        return False
    if "报价" in name:
        return any(token in text for token in ("响应报价", "总价", "单价", "采购项目名称", "服务名称", "货物名称"))
    if "政策" in name:
        return any(token in text for token in ("节能", "环保", "中小企业", "制造商", "认证证书"))
    if "偏离" in name or "响应表" in name:
        return any(token in text for token in ("偏离", "响应", "采购需求", "招标要求", "磋商文件要求"))
    name_tokens = [token for token in _template_search_tokens(name) if len(token) >= 3]
    return bool(name_tokens) and any(token in text for token in name_tokens)


def _pdf_text_template_lines_after_heading(pdf, candidate: Dict[str, Any], node_name: str) -> List[str]:
    start_page = int(candidate.get("page_index") or 0)
    start_y = float(candidate.get("y") or 0)
    target = _norm_pdf_text(node_name)
    lines: List[str] = []
    started = False
    for page_index in range(start_page, min(len(pdf), start_page + 4)):
        page = pdf[page_index]
        for line in _pdf_text_lines(page):
            y = float(line.get("y") or 0)
            if page_index == start_page and y + 2 < start_y:
                continue
            raw = str(line.get("text") or "").strip()
            text = _norm_pdf_text(raw)
            if not text:
                continue
            if not started:
                if not _pdf_heading_matches_template(text, target, _template_search_tokens(node_name)):
                    continue
                started = True
            elif _pdf_should_stop_text_template(text, target):
                return _trim_pdf_text_template_lines(lines, node_name)
            if _pdf_line_is_scoring_or_review(text):
                return []
            lines.append(_clean_pdf_extracted_line(raw))
    return _trim_pdf_text_template_lines(lines, node_name)


def _pdf_should_stop_text_template(text: str, target: str) -> bool:
    if not text or text == target or target in text:
        return False
    if _looks_like_inline_template_list_item(text):
        return False
    if _looks_like_numbered_list_item(text) and _looks_like_requirement_list_body_line(text):
        return False
    if re.match(r"^格式[一二三四五六七八九十0-9]+[:：]", text):
        return True
    if len(text) <= 35 and _looks_like_next_template_heading(text):
        return True
    return False


def _pdf_line_is_scoring_or_review(text: str) -> bool:
    return any(
        token in text
        for token in (
            "评审因素",
            "评审标准",
            "评分标准",
            "评分细则",
            "分值构成",
            "商务部分",
            "技术部分",
            "报价得分",
        )
    )


def _trim_pdf_text_template_lines(lines: List[str], node_name: str) -> List[str]:
    cleaned = [line for line in (_clean_pdf_extracted_line(line) for line in lines) if line]
    if len(cleaned) < 2:
        return []
    joined = _norm_pdf_text("".join(cleaned))
    name = _norm_pdf_text(node_name)
    if name and name not in joined:
        return []
    return cleaned[:80]


def _append_pdf_text_template(doc: Document, lines: List[str]) -> bool:
    paragraphs = _pdf_text_template_paragraphs(lines)
    if not paragraphs:
        return False
    for text in paragraphs:
        paragraph = doc.add_paragraph(text)
        if (
            len(_norm_text(text)) <= 25
            and _looks_like_template_heading(_norm_text(text))
            and not _looks_like_inline_template_list_item(text)
        ):
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True


def _pdf_text_template_paragraphs(lines: List[str]) -> List[str]:
    paragraphs: List[str] = []
    buffer: List[str] = []

    def flush() -> None:
        nonlocal buffer
        text = _merge_pdf_template_lines(buffer)
        if text:
            paragraphs.append(text)
        buffer = []

    for raw in lines or []:
        line = _clean_pdf_extracted_line(raw)
        if not line:
            flush()
            continue
        if _pdf_text_line_should_stand_alone(line):
            flush()
            paragraphs.append(line)
            continue
        if _pdf_text_line_starts_new_paragraph(line):
            flush()
        buffer.append(line)
        if _pdf_text_line_ends_paragraph(line):
            flush()
    flush()
    return paragraphs


def _merge_pdf_template_lines(lines: List[str]) -> str:
    result = ""
    for line in [str(item or "").strip() for item in lines if str(item or "").strip()]:
        if not result:
            result = line
            continue
        if _pdf_line_should_attach_without_space(result, line):
            result += line
        else:
            result += " " + line
    return result.strip()


def _pdf_text_line_should_stand_alone(line: str) -> bool:
    text = str(line or "").strip()
    compact = _norm_text(text)
    if not compact:
        return True
    if _looks_like_inline_template_list_item(text):
        return False
    if len(compact) <= 25 and _looks_like_template_heading(compact):
        return True
    if re.search(r"(?:报价人|投标人|供应商|法定代表人|委托代理人|日期|年\s*月\s*日|盖章|签字|签名)\s*[:：]?\s*$", text):
        return True
    return False


def _pdf_text_line_starts_new_paragraph(line: str) -> bool:
    text = str(line or "").strip()
    return bool(re.match(r"^(?:\d+[.)、]|[（(][一二三四五六七八九十0-9]+[）)])", text))


def _looks_like_inline_template_list_item(text: str) -> bool:
    value = str(text or "").strip()
    compact = _norm_text(value)
    if not compact:
        return False
    if not re.match(r"^(?:\d+[.)、]|[（(][一二三四五六七八九十0-9]+[）)])", compact):
        return False
    if not re.search(r"[；;]$", compact):
        return False
    return any(
        token in compact
        for token in (
            "响应函",
            "投标函",
            "身份证明",
            "授权委托书",
            "委托书",
            "报价表",
            "资格审查资料",
            "资格证明",
            "响应方案",
            "应答方案",
            "其他资料",
            "其他材料",
        )
    )


def _pdf_text_line_ends_paragraph(line: str) -> bool:
    text = str(line or "").strip()
    return bool(re.search(r"[。；;：:]$", text))


def _pdf_line_should_attach_without_space(previous: str, current: str) -> bool:
    previous = str(previous or "")
    current = str(current or "")
    if not previous or not current:
        return True
    if re.search(r"[（(￥¥]$", previous) or re.match(r"^[）),，。；;：:%％元]", current):
        return True
    if re.search(r"[一-鿿]$", previous) and re.match(r"^[一-鿿]", current):
        return True
    return False


def _clean_pdf_extracted_line(text: str) -> str:
    return _collapse_pdf_doubled_text(str(text or "").strip())


def _normalize_pdf_table_rows(raw_rows: List[List[Any]]) -> List[List[str]]:
    rows: List[List[str]] = []
    for row in raw_rows or []:
        cells = [str(cell or "").strip().replace("\n", " ") for cell in (row or [])]
        if any(cells):
            rows.append(cells)
    if not rows:
        return []
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    non_empty_cells = sum(1 for row in rows for cell in row if cell.strip())
    return rows if len(rows) >= 2 and width >= 2 and non_empty_cells >= 3 else []


def _append_pdf_rows_table(doc: Document, rows: List[List[str]]) -> bool:
    return _append_rows_table(doc, rows)


def _copy_tender_template_simple(node: Dict[str, Any], tender_doc_path: str, target_doc: Document) -> bool:
    source_path = Path(tender_doc_path)
    if not source_path.exists() or source_path.suffix.lower() != ".docx":
        return False
    source_doc = Document(str(source_path))
    blocks = list(source_doc.element.body.iterchildren())
    paragraph_by_element = {p._element: p for p in source_doc.paragraphs}

    tokens = _template_search_tokens(str(node.get("name", "")))
    if not tokens:
        return False

    candidates: List[Dict[str, int]] = []
    for idx, block in enumerate(blocks):
        if not block.tag.endswith("}p"):
            continue
        para = paragraph_by_element.get(block)
        text = _norm_text(para.text if para is not None else "")
        if not text:
            continue
        if any(token in text for token in tokens) and _looks_like_template_heading(text):
            end_idx = _template_copy_end_index(blocks, paragraph_by_element, idx)
            score = _template_candidate_score(blocks, paragraph_by_element, idx, end_idx)
            candidates.append({"start": idx, "end": end_idx, "score": score})
    if not candidates:
        return False

    best = max(candidates, key=lambda item: (int(item.get("score") or 0), int(item.get("start") or 0)))
    start_idx = int(best["start"])
    end_idx = int(best["end"])

    copied = 0
    markdown_table_buffer: List[str] = []

    def flush_markdown_table() -> None:
        nonlocal copied, markdown_table_buffer
        if markdown_table_buffer:
            if _append_markdown_table(target_doc, markdown_table_buffer):
                copied += 1
            markdown_table_buffer = []

    for block in blocks[start_idx:end_idx]:
        if block.tag.endswith("sectPr"):
            continue
        if block.tag.endswith("}p"):
            para = paragraph_by_element.get(block)
            text = para.text if para is not None else ""
            if _is_markdown_table_line(text):
                markdown_table_buffer.append(text)
                continue
            flush_markdown_table()
        else:
            flush_markdown_table()
        copied_block = deepcopy(block)
        normalize_copied_tender_template_block(copied_block)
        fit_copied_tender_template_tables_to_page(copied_block, target_doc)
        _append_body_block(target_doc, copied_block)
        copied += 1
    flush_markdown_table()
    return copied > 0


def _template_copy_end_index(blocks, paragraph_by_element, start_idx: int) -> int:
    end_idx = min(len(blocks), start_idx + 80)
    for idx in range(start_idx + 1, min(len(blocks), start_idx + 80)):
        block = blocks[idx]
        if block.tag.endswith("}p"):
            para = paragraph_by_element.get(block)
            text = _norm_text(para.text if para is not None else "")
            if idx > start_idx + 3 and _looks_like_response_cover_boundary(blocks, paragraph_by_element, idx):
                return idx
            if (
                idx > start_idx + 3
                and _looks_like_next_template_heading(text)
                and not _looks_like_numbered_list_item(text)
                and not _looks_like_requirement_list_body_line(text)
            ):
                return idx
    return end_idx


def _template_candidate_score(blocks, paragraph_by_element, start_idx: int, end_idx: int) -> int:
    texts: List[str] = []
    for block in blocks[start_idx:end_idx]:
        if not block.tag.endswith("}p"):
            continue
        para = paragraph_by_element.get(block)
        text = _norm_text(para.text if para is not None else "")
        if text:
            texts.append(text)
    joined = "".join(texts)
    score = min(len(joined), 2000) + max(0, end_idx - start_idx) * 5
    if any(token in joined for token in ("以下", "加盖公章", "签名", "日期", "复印件", "提供", "证明文件")):
        score += 1200
    if any(token in joined for token in ("目录", "索引表", "评审索引")) and len(joined) < 250:
        score -= 800
    # Later same-title hits are more likely to be real body sections than table-of-contents entries.
    score += min(start_idx, 500)
    return score


def _norm_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def _norm_pdf_text(text: str) -> str:
    return _norm_text(_collapse_pdf_doubled_text(text))


def _collapse_pdf_doubled_text(text: str) -> str:
    raw = str(text or "")
    compact = re.sub(r"\s+", "", raw)
    if len(compact) < 6 or len(compact) % 2:
        return raw
    pair_count = len(compact) // 2
    if pair_count and all(compact[idx] == compact[idx + 1] for idx in range(0, len(compact), 2)):
        collapsed = "".join(compact[idx] for idx in range(0, len(compact), 2))
        return collapsed
    return raw


def _looks_like_numbered_list_item(text: str) -> bool:
    value = _norm_text(text)
    return bool(re.match(r"^(\d{1,2}[\)）.．、]|[（(]\d{1,2}[）)])", value))



def _looks_like_requirement_list_body_line(text: str) -> bool:
    value = _norm_text(text).strip("。；;，,")
    if not value or len(value) > 90:
        return False
    body_tokens = (
        "营业执照",
        "复印件",
        "一般纳税人",
        "合并流水",
        "项目经验证明",
        "资质证明",
        "纳税信用",
        "征信报告",
        "企业报表",
        "完税",
        "税源地",
        "发票类型",
        "发票量",
        "社保缴纳",
        "员工社保",
        "采购需求",
        "响应方案",
        "应答方案",
        "服务方案",
        "项目团队",
        "专业能力",
        "服务能力",
        "管理制度",
        "奖惩制度",
        "方案介绍",
        "劳务服务",
        "服务经验",
        "合同协议",
    )
    if any(token in value for token in body_tokens):
        return True
    if "证明" in value and any(token in value for token in ("提供", "投标人", "企业", "公司", "项目", "税", "社保")):
        return True
    if "证书" in value and any(token in value for token in ("提供", "投标人", "企业", "公司", "资质", "信用", "纳税")):
        return True
    return False



def _looks_like_template_heading(text: str) -> bool:
    if not text or len(text) > 40:
        return False
    return any(text.endswith(suffix) or suffix in text for suffix in ("函", "表", "书", "声明", "证明", "承诺", "偏离表", "介绍", "一览"))


def _looks_like_next_template_heading(text: str) -> bool:
    if not _looks_like_template_heading(text):
        return False
    return bool(re.match(r"^([一二三四五六七八九十]+[、.．]|[（(]?[一二三四五六七八九十0-9]+[）).．、])", text)) or len(text) <= 18


def _looks_like_response_cover_boundary(blocks, paragraph_by_element, idx: int) -> bool:
    block = blocks[idx]
    para = paragraph_by_element.get(block)
    text = _norm_text(para.text if para is not None else "")
    if text not in {"投标文件", "响应文件", "投标响应文件", "报价文件", "商务投标文件", "技术投标文件"}:
        return False
    hits = 0
    for offset in range(1, 25):
        if idx + offset >= len(blocks):
            break
        next_block = blocks[idx + offset]
        if not next_block.tag.endswith("}p"):
            continue
        next_para = paragraph_by_element.get(next_block)
        next_text = _norm_text(next_para.text if next_para is not None else "")
        if any(token in next_text for token in ("商务技术部分", "商务部分", "技术部分", "项目名称", "投标单位", "供应商名称", "法定代表人")):
            hits += 1
    return hits >= 2

copy_pdf_tender_template_simple = _copy_pdf_tender_template_simple
copy_tender_template_simple = _copy_tender_template_simple
