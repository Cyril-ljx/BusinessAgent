from dotenv import load_dotenv
load_dotenv(encoding="utf-8-sig")
import asyncio
import hashlib
import html
import json
import os
import shutil
import subprocess
import sys
import tempfile
import time as _time
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional
import re

from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, Response
from loguru import logger
from pydantic import BaseModel, Field

from tender_agent.parsing import (
    parse_document,
    locate_explicit_composition_sections,
    locate_route_sections_llm,
    locate_route_sections_fallback,
    assemble_section_content,
)
from tender_agent.llm.gateway import llm_gateway
from tender_agent.understanding.graph import outline_review_graph
from tender_agent.understanding.composer import normalize_outline_numbering
from tender_agent.understanding.user_outline_parser import (
    outline_has_user_source,
    parse_user_outline_text,
    renumber_user_outline,
)
from tender_agent.understanding.material_mapper import map_materials
from tender_agent.understanding.rag_retriever import build_rag_contexts
from tender_agent.understanding.content_generator import generate_content
from tender_agent.understanding.compliance_checker import run_compliance_checks
from tender_agent.understanding.consistency_checker import check_consistency
from tender_agent.understanding.analysis_facts import build_tender_analysis_facts, strategy_prompt
from tender_agent.rendering.render_plan import build_render_decision_report
from tender_agent.utils.outline_renderer import render_outline_markdown
from tender_agent.observability.langsmith_tracing import trace_stage
from tender_agent.api.auth import get_current_user

# LangSmith env compatibility:
# some stacks read LANGSMITH_* while older configs use LANGCHAIN_*.
if os.getenv("LANGCHAIN_API_KEY") and not os.getenv("LANGSMITH_API_KEY"):
    os.environ["LANGSMITH_API_KEY"] = os.getenv("LANGCHAIN_API_KEY", "")
if os.getenv("LANGCHAIN_PROJECT") and not os.getenv("LANGSMITH_PROJECT"):
    os.environ["LANGSMITH_PROJECT"] = os.getenv("LANGCHAIN_PROJECT", "")
if os.getenv("LANGCHAIN_ENDPOINT") and not os.getenv("LANGSMITH_ENDPOINT"):
    os.environ["LANGSMITH_ENDPOINT"] = os.getenv("LANGCHAIN_ENDPOINT", "")
if (
    os.getenv("LANGCHAIN_TRACING_V2", "").lower() in {"1", "true", "yes", "on"}
    and not os.getenv("LANGSMITH_TRACING")
):
    os.environ["LANGSMITH_TRACING"] = "true"
if (
    os.getenv("LANGSMITH_TRACING", "").lower() in {"1", "true", "yes", "on"}
    and not os.getenv("LANGSMITH_TRACING_V2")
):
    os.environ["LANGSMITH_TRACING_V2"] = "true"

try:
    from langsmith import traceable as _langsmith_traceable
except ImportError:  # pragma: no cover - LangSmith is optional at runtime.
    _langsmith_traceable = None


def _langsmith_tracing_enabled() -> bool:
    return any(
        os.getenv(name, "").lower() in {"1", "true", "yes", "on"}
        for name in ("LANGSMITH_TRACING", "LANGSMITH_TRACING_V2", "LANGCHAIN_TRACING_V2")
    )


def _task_trace_inputs(inputs: Dict[str, Any]) -> Dict[str, Any]:
    file_path = inputs.get("file_path")
    return {
        "task_id": inputs.get("task_id") or inputs.get("project_id"),
        "phase": inputs.get("trace_phase") or "initial_generation",
        "file_name": Path(str(file_path)).name if file_path else "",
        "company_id": inputs.get("company_id"),
        "company_name": inputs.get("company_name"),
        "generate_mode": inputs.get("generate_mode"),
    }


def _task_trace_outputs(output: Any) -> Dict[str, Any]:
    return {"status": "done" if output is None else type(output).__name__}


def _trace_task(fn):
    if _langsmith_traceable is None:
        return fn
    return _langsmith_traceable(
        name="TenderGenerationTask",
        run_type="chain",
        process_inputs=_task_trace_inputs,
        process_outputs=_task_trace_outputs,
        enabled=_langsmith_tracing_enabled(),
    )(fn)


@trace_stage("locator")
async def _locate_sections_stage(parsed):
    structured_sections = locate_explicit_composition_sections(parsed)
    if structured_sections:
        return structured_sections
    return await asyncio.to_thread(
        locate_route_sections_llm,
        parsed,
        llm_gateway.call_structured,
    )


@trace_stage("material_mapper")
async def _map_materials_stage(state: Dict[str, Any]) -> Dict[str, Any]:
    return await map_materials(state)


@trace_stage("rag_retriever")
async def _build_rag_contexts_stage(state: Dict[str, Any]) -> Dict[str, Any]:
    return await build_rag_contexts(state)


@trace_stage("content_generator")
async def _generate_content_stage(state: Dict[str, Any]) -> Dict[str, Any]:
    return await generate_content(state)


@trace_stage("compliance")
def _run_compliance_stage(state: Dict[str, Any]) -> Dict[str, Any]:
    return run_compliance_checks(state)


@trace_stage("consistency")
def _check_consistency_stage(state: Dict[str, Any]) -> Dict[str, Any]:
    return check_consistency(state)

app = FastAPI(title="NGU BusinessAgent API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
async def _log_runtime_flags():
    logger.info(
        "LLM provider(default)={}, LANGSMITH_TRACING={}, LANGSMITH_PROJECT={}",
        os.getenv("DEFAULT_LLM_PROVIDER", ""),
        os.getenv("LANGSMITH_TRACING", "false"),
        os.getenv("LANGSMITH_PROJECT", os.getenv("LANGCHAIN_PROJECT", "")),
    )
    _ensure_companies_ready()

tasks: Dict[str, Dict[str, Any]] = {}
background_tasks: Dict[str, asyncio.Task] = {}

UPLOAD_DIR = Path("data/upload")
OUTPUT_DIR = Path("data/output")
TEMPLATE_PATH = Path("assets/templates/master_template.docx")
TECH_MASTER_PATH = Path("data/knowledge/master/技术文件.docx")
CERTS_DIR = Path("data/knowledge/certs")
HISTORY_DIR = Path("data/knowledge/history")
DEFAULT_COMPANY_ID = "demo-company"
DEFAULT_COMPANY_NAME = "示例科技服务有限公司"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
CERTS_DIR.mkdir(parents=True, exist_ok=True)
HISTORY_DIR.mkdir(parents=True, exist_ok=True)


def _normalize_company_id(raw: str) -> str:
    value = re.sub(r"[^a-zA-Z0-9_.-]+", "-", str(raw or "").strip()).strip("-_.")
    return value[:80] or f"company-{uuid.uuid4().hex[:8]}"


def _ensure_company_storage(company_id: str) -> None:
    base = _company_knowledge_dir(company_id)
    for subdir in ("certs", "master", "history"):
        (base / subdir).mkdir(parents=True, exist_ok=True)


def _ensure_companies_ready() -> None:
    try:
        from tender_agent.core.db import engine, SessionLocal
        from data.knowledge.models import Base, Company

        Base.metadata.create_all(bind=engine)
        db = SessionLocal()
        try:
            default = db.query(Company).filter(Company.id == DEFAULT_COMPANY_ID).first()
            if default is None:
                default = Company(
                    id=DEFAULT_COMPANY_ID,
                    name=DEFAULT_COMPANY_NAME,
                    is_default=True,
                    is_active=True,
                )
                db.add(default)
            else:
                default.name = default.name or DEFAULT_COMPANY_NAME
                default.is_active = True
                if not db.query(Company).filter(Company.is_default == True).first():
                    default.is_default = True
            db.commit()
            _ensure_company_storage(DEFAULT_COMPANY_ID)
        finally:
            db.close()
    except Exception as exc:
        logger.warning("[companies] ensure failed: {}", str(exc)[:160])


def _task_meta_path(project_id: str) -> Path:
    return OUTPUT_DIR / f"{project_id}.task.json"


def _company_knowledge_dir(company_id: str) -> Path:
    safe_id = re.sub(r'[^a-zA-Z0-9_.-]+', "_", str(company_id or DEFAULT_COMPANY_ID)).strip("_")
    return Path("data/knowledge/companies") / (safe_id or DEFAULT_COMPANY_ID)


def _company_tech_master_path(company_id: str) -> Path:
    company_path = _company_knowledge_dir(company_id) / "master" / "技术文件.docx"
    if company_path.exists():
        return company_path
    return _resolve_tech_master_path()


def _resolve_tech_master_path() -> Path:
    """Resolve technical master path with robust fallbacks."""
    candidates = [
        Path("data/knowledge/master/技术文件.docx"),
        Path("scripts/技术文件.docx"),
        TECH_MASTER_PATH,
    ]
    for c in candidates:
        if c.exists():
            return c

    master_dir = Path("data/knowledge/master")
    if master_dir.exists():
        docx_files = sorted(master_dir.glob("*.docx"))
        if docx_files:
            for f in docx_files:
                if "技术" in f.name:
                    return f
            return docx_files[0]

    return Path("data/knowledge/master/技术文件.docx")


def _save_task_meta(task: Dict[str, Any]) -> None:
    try:
        payload = {
            "id": task.get("id"),
            "creator_user_id": task.get("creator_user_id"),
            "filename": task.get("filename"),
            "company_id": task.get("company_id"),
            "company_name": task.get("company_name"),
            "generate_mode": task.get("generate_mode"),
            "auto_add_score": task.get("auto_add_score"),
            "status": task.get("status"),
            "progress": task.get("progress"),
            "message": task.get("message"),
            "error": task.get("error"),
            "created_at": task.get("created_at"),
            "source_file_path": task.get("source_file_path"),
            "original_source_file_path": task.get("original_source_file_path"),
            "step_times": task.get("step_times", {}),
            "step_summaries": task.get("step_summaries", {}),
            "agent_metrics": task.get("agent_metrics", {}),
            "docx_path": task.get("docx_path"),
            "current_node": task.get("current_node"),
        }
        _task_meta_path(str(task.get("id", ""))).write_text(
            json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except Exception as exc:
        logger.warning(f"save task meta failed: {exc}")


def _get_task_or_restore(project_id: str) -> Optional[Dict[str, Any]]:
    task = tasks.get(project_id)
    if task:
        return task

    meta_file = _task_meta_path(project_id)
    if meta_file.exists():
        try:
            meta = json.loads(meta_file.read_text(encoding="utf-8"))
            restored_result = None
            output_file = OUTPUT_DIR / f"{project_id}.json"
            if output_file.exists():
                try:
                    restored_result = json.loads(output_file.read_text(encoding="utf-8"))
                except Exception as exc:
                    logger.warning(f"restore output json failed {project_id}: {exc}")
            restored_meta = {
                "id": project_id,
                "creator_user_id": meta.get("creator_user_id"),
                "filename": meta.get("filename", project_id),
                "company_id": meta.get("company_id") or DEFAULT_COMPANY_ID,
                "company_name": meta.get("company_name", ""),
                "generate_mode": meta.get("generate_mode", "balanced"),
                "auto_add_score": bool(meta.get("auto_add_score", True)),
                "status": meta.get("status", TaskStatus.PENDING),
                "progress": int(meta.get("progress", 0)),
                "message": meta.get("message", "任务恢复中"),
                "error": meta.get("error"),
                "created_at": meta.get("created_at", datetime.now().isoformat()),
                "result": restored_result,
                "source_file_path": meta.get("source_file_path"),
                "original_source_file_path": meta.get("original_source_file_path"),
                "step_times": meta.get("step_times", {}),
                "step_summaries": meta.get("step_summaries", {}),
                "agent_metrics": meta.get("agent_metrics", {}),
                "docx_path": meta.get("docx_path"),
                "current_node": meta.get("current_node"),
            }
            if restored_meta["status"] in {
                TaskStatus.PENDING,
                TaskStatus.PARSING,
                TaskStatus.LOCATING,
                TaskStatus.COMPOSING,
            }:
                restored_meta["status"] = TaskStatus.FAILED
                restored_meta["error"] = "服务重启导致后台任务中断，请重新运行该项目"
                restored_meta["message"] = "任务在服务重启时中断"
                restored_meta["current_node"] = "interrupted"
                _save_task_meta(restored_meta)
            tasks[project_id] = restored_meta
            return restored_meta
        except Exception as exc:
            logger.warning(f"restore task meta failed {project_id}: {exc}")

    output_file = OUTPUT_DIR / f"{project_id}.json"
    if not output_file.exists():
        return None

    try:
        result = json.loads(output_file.read_text(encoding="utf-8"))
    except Exception as exc:
        logger.warning(f"恢复任务失败 {project_id}: {exc}")
        return None

    source_file_path = result.get("source_file_path")
    if not source_file_path:
        for suffix in (".docx", ".doc"):
            candidate = UPLOAD_DIR / f"{project_id}{suffix}"
            if candidate.exists():
                source_file_path = str(candidate)
                break

    restored = {
        "id": project_id,
        "creator_user_id": None,
        "filename": Path(source_file_path).name if source_file_path else project_id,
        "company_id": DEFAULT_COMPANY_ID,
        "company_name": "投标单位",
        "generate_mode": "balanced",
        "auto_add_score": True,
        "status": TaskStatus.DONE,
        "progress": 100,
        "message": "任务已从历史结果恢复",
        "error": None,
        "created_at": datetime.now().isoformat(),
        "result": result,
        "source_file_path": source_file_path,
        "original_source_file_path": _find_original_upload_path(project_id, Path(source_file_path).name if source_file_path else None),
        "current_node": "done",
    }

    docx_path = OUTPUT_DIR / f"{project_id}_blank_bid.docx"
    if docx_path.exists():
        restored["docx_path"] = str(docx_path)
    tasks[project_id] = restored
    _save_task_meta(restored)
    return restored


def _project_creator_user_id(project_id: str) -> Optional[str]:
    task = tasks.get(project_id)
    if isinstance(task, dict):
        return str(task.get("creator_user_id") or "").strip() or None
    meta_file = _task_meta_path(project_id)
    if not meta_file.exists():
        return None
    try:
        meta = json.loads(meta_file.read_text(encoding="utf-8"))
    except Exception as exc:
        logger.warning("read project owner failed {}: {}", project_id, str(exc)[:120])
        return None
    return str(meta.get("creator_user_id") or "").strip() or None


def _project_id_from_request_path(path: str) -> Optional[str]:
    prefix = "/api/projects/"
    if not path.startswith(prefix):
        return None
    candidate = path[len(prefix):].split("/", 1)[0].strip()
    if not candidate or candidate == "upload":
        return None
    try:
        return str(uuid.UUID(candidate))
    except (ValueError, TypeError, AttributeError):
        return None


@app.middleware("http")
async def _project_access_control(request: Request, call_next):
    path = request.url.path
    if path == "/api/auth/me" or path == "/api/projects" or path.startswith("/api/projects/"):
        try:
            user = await get_current_user(request)
        except HTTPException as exc:
            return JSONResponse(status_code=exc.status_code, content={"detail": exc.detail})
        request.state.current_user = user

        project_id = _project_id_from_request_path(path)
        if project_id and _project_creator_user_id(project_id) != user.user_id:
            return JSONResponse(status_code=404, content={"detail": "Project not found"})
    return await call_next(request)


@app.get("/api/auth/me", summary="Get current authenticated user")
async def auth_me(request: Request):
    user = request.state.current_user
    return {"user_id": user.user_id, "nick_name": user.nick_name}


def _ensure_task_result(project_id: str, task: Dict[str, Any]) -> Dict[str, Any]:
    result = task.get("result")
    if isinstance(result, dict) and result:
        return result

    output_file = OUTPUT_DIR / f"{project_id}.json"
    if not output_file.exists():
        raise HTTPException(status_code=404, detail="Project result file not found")

    try:
        result = json.loads(output_file.read_text(encoding="utf-8"))
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Project result file is invalid: {str(exc)[:120]}")
    task["result"] = result
    _save_task_meta(task)
    return result


def _safe_upload_name(filename: str) -> str:
    suffix = Path(filename or "").suffix.lower()
    stem = re.sub(r'[\\/:*?"<>|\s]+', "_", Path(filename or "upload").stem).strip("_")[:60] or "upload"
    return f"{stem}_{uuid.uuid4().hex[:10]}{suffix}"


def _heading_level_from_paragraph(para, allow_numbered_fallback: bool = True) -> int:
    style_name = (para.style.name or "").strip().lower()
    if style_name.startswith("toc") or "目录" in style_name or "目錄" in style_name:
        return 0
    style_match = re.search(r"(?:^|\s)heading\s*(\d+)\b", style_name)
    if not style_match:
        style_match = re.search(r"(?:标题|標題)\s*(\d+)", style_name)
    if style_match:
        level = int(style_match.group(1))
        return level if 1 <= level <= 9 else 0
    try:
        p_pr = para._element.pPr
        if p_pr is not None:
            outline_lvl = p_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl")
            if outline_lvl is not None:
                val = outline_lvl.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                if val is not None and str(val).isdigit() and 0 <= int(val) <= 8:
                    return int(val) + 1
    except Exception:
        pass
    if not allow_numbered_fallback:
        return 0
    text = para.text.strip()
    numbered = re.match(r"^(\d+(?:\.\d+)*)(?:[.、]\s*|\s+)", text)
    if numbered and len(text) <= 80:
        return min(numbered.group(1).count(".") + 1, 4)
    return 0


def _clean_section_title(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^\d+(?:\.\d+)*[.、\s]+", "", text)
    return text.strip() or "未命名章节"


def _scan_template_sections_from_docx(
    file_path: Path,
    category: str,
    replace_current: bool = False,
    company_id: str = DEFAULT_COMPANY_ID,
    scope: str = "company",
) -> int:
    from docx import Document
    from docx.text.paragraph import Paragraph
    from tender_agent.core.db import get_db_session
    from data.knowledge.models import TemplateSection

    db = next(get_db_session())
    try:
        if replace_current:
            query = db.query(TemplateSection)
            if category == "technical_master":
                query = query.filter((TemplateSection.category == category) | (TemplateSection.category.is_(None)))
            else:
                query = query.filter(TemplateSection.category == category)
            query = query.filter(TemplateSection.company_id == company_id, TemplateSection.scope == scope)
            query.delete(synchronize_session=False)

        doc = Document(str(file_path))
        paragraph_blocks = []
        for block_idx, child in enumerate(doc.element.body.iterchildren()):
            if not child.tag.endswith("}p"):
                continue
            para = Paragraph(child, doc)
            paragraph_blocks.append((block_idx, para, para.text.strip()))

        explicit_heading_count = sum(
            1
            for _, para, text in paragraph_blocks
            if text and _heading_level_from_paragraph(para, allow_numbered_fallback=False) > 0
        )
        allow_numbered_fallback = explicit_heading_count < 3
        logger.info(
            "[knowledge] scan tech master headings explicit={} numbered_fallback={}",
            explicit_heading_count,
            allow_numbered_fallback,
        )

        finalized = []
        stack = []
        counters: Dict[int, int] = {}
        body_blocks = list(doc.element.body.iterchildren())
        last_block_idx = max(len(body_blocks) - 1, 0)
        for block_idx, para, text in paragraph_blocks:
            level = _heading_level_from_paragraph(para, allow_numbered_fallback=allow_numbered_fallback)
            if level <= 0 or not text or text in {"目录", "目 录"}:
                continue

            while stack and stack[-1]["level"] >= level:
                closing = stack.pop()
                closing["end_block_idx"] = block_idx
                finalized.append(closing)

            counters[level] = counters.get(level, 0) + 1
            for key in list(counters.keys()):
                if key > level:
                    del counters[key]
            for key in range(1, level):
                counters.setdefault(key, 1)
            chapter_id = ".".join(str(counters[key]) for key in sorted(counters) if key <= level)
            title = _clean_section_title(text)
            parent_path = stack[-1]["full_path"] if stack else ""
            full_path = f"{parent_path}/{title}" if parent_path else title
            stack.append(
                {
                    "chapter_id": chapter_id,
                    "title": title,
                    "full_path": full_path,
                    "level": level,
                    "start_block_idx": block_idx + 1,
                    "end_block_idx": None,
                }
            )

        while stack:
            closing = stack.pop()
            closing["end_block_idx"] = last_block_idx + 1
            finalized.append(closing)

        finalized.sort(key=lambda item: item["start_block_idx"])
        now = datetime.now()
        for sec in finalized:
            section_start = max(int(sec["start_block_idx"] or 0) - 1, 0)
            section_end = min(int(sec["end_block_idx"] or 0), len(body_blocks))
            section_blocks = body_blocks[section_start:section_end]
            raster_image_count = sum(
                len(block.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"))
                for block in section_blocks
            )
            vector_graphic_count = sum(_block_has_unsupported_word_graphic(block) for block in section_blocks)
            table_count = sum(str(block.tag).endswith("}tbl") for block in section_blocks)
            db.add(
                TemplateSection(
                    chapter_id=sec["chapter_id"],
                    title=sec["title"],
                    full_path=sec["full_path"],
                    level=sec["level"],
                    company_id=company_id,
                    scope=scope,
                    start_block_idx=sec["start_block_idx"],
                    end_block_idx=sec["end_block_idx"],
                    category=category,
                    has_image=bool(raster_image_count or vector_graphic_count),
                    has_table=bool(table_count),
                    image_count=raster_image_count,
                    table_count=table_count,
                    metadata_info={
                        "source_file_path": str(file_path),
                        "imported_at": now.isoformat(),
                        "vector_graphic_count": vector_graphic_count,
                    },
                    is_current=True,
                )
            )
        db.commit()
        return len(finalized)
    finally:
        db.close()


class TaskStatus(str):
    PENDING = "pending"
    PARSING = "parsing"
    LOCATING = "locating"
    COMPOSING = "composing"  # V11: LLM compose / graph processing
    OUTLINE_REVIEW = "outline_review"
    DONE = "done"
    FAILED = "failed"
    CANCELLED = "cancelled"


def _outline_available(status: str) -> bool:
    return status in {TaskStatus.OUTLINE_REVIEW, TaskStatus.DONE}


class TaskStatusResponse(BaseModel):
    status: str
    progress: int
    message: str
    error: Optional[str] = None
    current_node: Optional[str] = None
    step_times: Optional[Dict[str, float]] = None
    step_summaries: Optional[Dict[str, Dict[str, Any]]] = None
    agent_metrics: Optional[Dict[str, Dict[str, Any]]] = None


class OutlineResponse(BaseModel):
    company_id: Optional[str] = None
    company_name: Optional[str] = None
    title_info: Dict[str, Any]
    outline: list
    tender_requirements: Dict[str, Any] = {}
    tender_requirements_stats: Dict[str, Any] = {}
    analysis_facts: Dict[str, Any] = {}
    analysis_facts_summary: Dict[str, int] = {}
    analysis_recommendations: Dict[str, Any] = {}
    material_assignments: list = []
    render_decisions: list = []
    retrieval_summary: Dict[str, Any] = {}
    generated_sections: Dict[str, str] = {}
    project_facts: Dict[str, Any] = {}
    compliance_report: Dict[str, Any] = {}
    consistency_report: Dict[str, Any] = {}
    workflow_stage: str = ""
    source_file_path: Optional[str] = None
    block_index: list[Dict[str, Any]] = []
    warnings: list[str]
    stats: Dict[str, Any]


class ProjectListItem(BaseModel):
    id: str
    filename: str
    project_name: str = ""
    purchaser: str = ""
    company_id: str = ""
    company_name: str = ""
    status: str
    progress: int = 0
    message: str = ""
    error: Optional[str] = None
    created_at: Optional[str] = None
    updated_at: Optional[str] = None
    has_result: bool = False
    has_docx: bool = False
    outline_count: int = 0
    generated_section_count: int = 0
    compliance_issue_count: int = 0
    consistency_conflict_count: int = 0


class SourceLocateQuery(BaseModel):
    id: str
    text: str


class SourceLocateRequest(BaseModel):
    anchors: list[str] = []
    query_items: list[SourceLocateQuery] = Field(default_factory=list)


class OutlineFromTextRequest(BaseModel):
    text: str
    remap: bool = False


class SourceAskRequest(BaseModel):
    question: str
    top_k: int = 8


class SectionDraftRequest(BaseModel):
    content: str = ""
    remove_material_sources: list[str] = Field(default_factory=list)


class TenderTemplatePreviewRequest(BaseModel):
    name: str = ""
    anchor_start: str = ""
    anchor_end: str = ""
    copy_method: str = ""


class AppendOutlineChildrenRequest(BaseModel):
    parent_id: str
    sections: list[Dict[str, Any]] = []


class AppendTechSectionsRequest(BaseModel):
    parent_id: str
    section_ids: list[str] = []
    company_id: str = DEFAULT_COMPANY_ID


class TechSectionBulkDeleteRequest(BaseModel):
    section_ids: list[str] = Field(default_factory=list)


class OutlineNodeHookRequest(BaseModel):
    action: str


class SourceAskCitation(BaseModel):
    id: str
    title: str = ""
    quote: str = ""
    anchor: Optional[str] = None
    page_no: Optional[int] = None
    preview_page_no: Optional[int] = None
    score: float = 0.0


class SourceAskResponse(BaseModel):
    answer: str
    citations: list[SourceAskCitation] = Field(default_factory=list)
    confidence: str = "low"
    used_llm: bool = False


class _SourceAskLLMResponse(BaseModel):
    answer: str
    citation_ids: list[str] = Field(default_factory=list)
    confidence: str = "low"


class AnalysisRecommendationsRequest(BaseModel):
    force: bool = False


class _AnalysisAdviceItem(BaseModel):
    fact_id: str = ""
    kind: str = ""
    suggestions: list[str] = Field(default_factory=list)


class _AnalysisAdviceResponse(BaseModel):
    items: list[_AnalysisAdviceItem] = Field(default_factory=list)


class CompanyResponse(BaseModel):
    id: str
    name: str
    is_default: bool = False
    is_active: bool = True


class CompanyCreateRequest(BaseModel):
    name: str
    id: Optional[str] = None
    is_default: bool = False


class CompanyUpdateRequest(BaseModel):
    name: Optional[str] = None
    is_default: Optional[bool] = None
    is_active: Optional[bool] = None


def convert_doc_to_docx(doc_path: Path) -> Path:
    output_dir = doc_path.parent
    new_path = output_dir / f"{doc_path.stem}.docx"

    # Prefer Windows + Word COM conversion when available.
    if sys.platform == "win32":
        try:
            import pythoncom
            import win32com.client
            pythoncom.CoInitialize()
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(str(doc_path.resolve()))
            try:
                doc.SaveAs2(str(new_path.resolve()), FileFormat=16)  # 16 = docx
            finally:
                doc.Close(False)
                word.Quit()
                pythoncom.CoUninitialize()
            if new_path.exists():
                return new_path
        except Exception as e:
            logger.warning(f"Word COM 转换失败: {e}")

    # Fallback to LibreOffice conversion.
    try:
        cmd = [
            "libreoffice", "--headless",
            "--convert-to", "docx",
            "--outdir", str(output_dir),
            str(doc_path),
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        if result.returncode == 0 and new_path.exists():
            return new_path
    except Exception as e:
        logger.warning(f"LibreOffice 转换失败: {e}")

    raise RuntimeError("无法将 .doc 转换为 .docx，请先另存为 .docx 后上传。")


def _source_preview_pdf_path(project_id: str) -> Path:
    return OUTPUT_DIR / f"{project_id}_source_preview.pdf"


def _source_preview_meta_path(project_id: str) -> Path:
    return OUTPUT_DIR / f"{project_id}_source_preview.source.txt"


def _find_libreoffice_executable() -> str:
    executable = shutil.which("libreoffice") or shutil.which("soffice")
    if executable:
        return executable
    if sys.platform == "win32":
        for candidate in (
            Path(os.environ.get("PROGRAMFILES", r"C:\Program Files")) / "LibreOffice/program/soffice.exe",
            Path(os.environ.get("PROGRAMFILES(X86)", r"C:\Program Files (x86)")) / "LibreOffice/program/soffice.exe",
        ):
            if candidate.exists():
                return str(candidate)
    raise FileNotFoundError("LibreOffice executable not found")


def _find_original_upload_path(project_id: str, filename: Optional[str] = None) -> Optional[str]:
    suffixes: list[str] = []
    if filename:
        suffix = Path(filename).suffix.lower()
        if suffix:
            suffixes.append(suffix)
    for suffix in (".doc", ".pdf", ".docx"):
        if suffix not in suffixes:
            suffixes.append(suffix)
    for suffix in suffixes:
        candidate = UPLOAD_DIR / f"{project_id}{suffix}"
        if candidate.exists():
            return str(candidate)
    return None


def _source_path_for_preview(
    project_id: str,
    result: Optional[Dict[str, Any]] = None,
    task: Optional[Dict[str, Any]] = None,
) -> Optional[str]:
    result = result or {}
    task = task or {}
    for value in (
        task.get("original_source_file_path"),
        result.get("original_source_file_path"),
    ):
        if value and Path(value).exists():
            return str(value)
    original = _find_original_upload_path(project_id, task.get("filename"))
    if original:
        return original
    for value in (result.get("source_file_path"), task.get("source_file_path")):
        if value and Path(value).exists():
            return str(value)
    return None


def _ensure_source_preview_pdf(project_id: str, source_path: Path) -> Path:
    source_path = Path(source_path)
    if not source_path.exists():
        raise FileNotFoundError(str(source_path))
    if source_path.suffix.lower() == ".pdf":
        return source_path

    preview_path = _source_preview_pdf_path(project_id)
    preview_meta_path = _source_preview_meta_path(project_id)
    source_identity = str(source_path.resolve())
    preview_source_identity = ""
    if preview_meta_path.exists():
        try:
            preview_source_identity = preview_meta_path.read_text(encoding="utf-8").strip()
        except Exception:
            preview_source_identity = ""
    if (
        preview_path.exists()
        and preview_path.stat().st_mtime >= source_path.stat().st_mtime
        and preview_source_identity == source_identity
    ):
        return preview_path

    if sys.platform == "win32":
        try:
            import pythoncom
            import win32com.client

            pythoncom.CoInitialize()
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(str(source_path.resolve()))
            try:
                doc.SaveAs(str(preview_path.resolve()), FileFormat=17)  # 17 = PDF
            finally:
                doc.Close(False)
                word.Quit()
                pythoncom.CoUninitialize()
            if preview_path.exists():
                preview_meta_path.write_text(source_identity, encoding="utf-8")
                return preview_path
        except Exception as exc:
            logger.warning("Word COM 转 PDF 失败: {}", exc)

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            libreoffice = _find_libreoffice_executable()
            profile_uri = (Path(tmpdir) / "libreoffice-profile").resolve().as_uri()
            cmd = [
                libreoffice,
                f"-env:UserInstallation={profile_uri}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                tmpdir,
                str(source_path),
            ]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
            generated = Path(tmpdir) / f"{source_path.stem}.pdf"
            if result.returncode == 0 and generated.exists():
                shutil.copy2(generated, preview_path)
                preview_meta_path.write_text(source_identity, encoding="utf-8")
                return preview_path
    except Exception as exc:
        logger.warning("LibreOffice 转 PDF 失败: {}", exc)

    raise RuntimeError("无法生成原文件 PDF 预览，请确认本机安装了 Word 或 LibreOffice。")


def _normalize_match_text(text: str) -> str:
    compact = str(text or "").lower()
    compact = re.sub(r"[\s\|\-—_·•●`~!@#$%^&*()（）\[\]【】{}<>《》“”\"'‘’、,，。；;：:？?！!]+", "", compact)
    return compact


def _candidate_match_fragments(text: str) -> list[str]:
    raw = str(text or "")
    normalized = _normalize_match_text(raw)
    candidates: list[str] = []
    for token in re.split(r"[\n\r\|：:，,。；;、\s]+", raw):
        compact = _normalize_match_text(token)
        if len(compact) >= 8:
            candidates.append(compact[:80])
    if len(normalized) >= 12:
        candidates.extend([normalized[:120], normalized[:80], normalized[:50], normalized[:30]])
    deduped: list[str] = []
    seen: set[str] = set()
    for item in sorted(candidates, key=len, reverse=True):
        if not item or item in seen:
            continue
        seen.add(item)
        deduped.append(item)
    return deduped


def _load_preview_page_texts(preview_pdf: Path) -> list[dict[str, Any]]:
    try:
        import fitz
    except Exception as exc:
        raise RuntimeError(f"PDF 预览依赖 PyMuPDF 不可用: {exc}") from exc

    pages: list[dict[str, Any]] = []
    pdf = fitz.open(str(preview_pdf))
    try:
        for page_no, page in enumerate(pdf, start=1):
            text = page.get_text("text") or ""
            pages.append(
                {
                    "page_no": page_no,
                    "text": text,
                    "normalized": _normalize_match_text(text),
                }
            )
    finally:
        pdf.close()
    return pages


def _locate_text_page(text: str, pages: list[dict[str, Any]], start_page: int = 1) -> int | None:
    if not pages:
        return None
    candidates = _candidate_match_fragments(text)
    if not candidates:
        return None

    start_idx = max(0, min(len(pages) - 1, start_page - 1))
    search_order = list(range(start_idx, len(pages))) + list(range(0, start_idx))

    for fragment in candidates:
        for page_idx in search_order:
            if fragment in (pages[page_idx].get("normalized") or ""):
                return int(pages[page_idx]["page_no"])

    best_page: int | None = None
    best_score = 0
    for page in pages:
        normalized = page.get("normalized") or ""
        score = sum(1 for fragment in candidates[:6] if fragment and fragment in normalized)
        if score > best_score:
            best_score = score
            best_page = int(page["page_no"])
    return best_page if best_score > 0 else None



def _source_ask_compact_text(value: Any, limit: int = 1200) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    return text[:limit]


def _source_ask_expand_question(question: str) -> str:
    text = str(question or "")
    compact = re.sub(r"\s+", "", text)
    expansions: list[str] = []
    if any(term in compact for term in ("投标保证金", "响应保证金", "保证金")):
        expansions.extend(["投标保证金", "响应保证金", "保证金", "缴纳", "交纳", "提交", "到账", "截止", "截止时间", "递交截止", "退还"])
    if any(term in compact for term in ("截止", "时间", "什么时候", "何时", "日期")):
        expansions.extend(["截止", "截止时间", "递交截止", "提交截止", "报价截止", "开标时间", "时间", "日期"])
    if any(term in compact for term in ("目录", "组成", "构成", "清单", "提交哪些", "需要哪些")):
        expansions.extend(["目录", "组成", "构成", "清单", "投标文件", "响应文件", "应提交", "包括", "格式"])
    if not expansions:
        return text
    return f"{text} {' '.join(dict.fromkeys(expansions))}"


def _source_ask_terms(question: str) -> set[str]:
    expanded = _source_ask_expand_question(question)
    compact = re.sub(r"\s+", "", str(expanded or "").lower())
    terms = {token.lower() for token in re.findall(r"[A-Za-z0-9_\-]{2,}|[一-鿿]{2,}", expanded or "")}
    chinese = "".join(re.findall(r"[一-鿿]", compact))
    for size in (2, 3, 4):
        for idx in range(0, max(0, len(chinese) - size + 1)):
            terms.add(chinese[idx : idx + size])
    return {term for term in terms if len(term) >= 2}

def _source_ask_snippet_score(question: str, snippet: Dict[str, Any]) -> float:
    terms = _source_ask_terms(question)
    if not terms:
        return 0.0
    text = f"{snippet.get('title') or ''} {snippet.get('text') or ''}".lower()
    compact_text = re.sub(r"\s+", "", text)
    score = 0.0
    for term in terms:
        count = compact_text.count(term.lower())
        if count:
            score += min(count, 4) * (2.0 if len(term) >= 3 else 1.0)
            if term in re.sub(r"\s+", "", str(snippet.get("title") or "").lower()):
                score += 3.0
    return score


def _source_ask_add_snippet(
    snippets: list[Dict[str, Any]],
    seen: set[str],
    *,
    source: str,
    title: str,
    text_value: Any,
    anchor: Optional[str] = None,
    page_no: Optional[int] = None,
    preview_page_no: Optional[int] = None,
) -> None:
    text_value = _source_ask_compact_text(text_value)
    if len(text_value) < 8:
        return
    key = re.sub(r"\s+", "", f"{anchor or ''}:{text_value[:120]}")
    if key in seen:
        return
    seen.add(key)
    snippets.append(
        {
            "source": source,
            "title": _source_ask_compact_text(title, 120),
            "text": text_value,
            "anchor": anchor,
            "page_no": page_no,
            "preview_page_no": preview_page_no,
        }
    )


def _source_ask_collect_snippets(result: Dict[str, Any]) -> list[Dict[str, Any]]:
    snippets: list[Dict[str, Any]] = []
    seen: set[str] = set()
    # Source Q&A must search the original tender text only. Do not add generated
    # outline, extracted checklist summaries, or other system-composed facts here.
    for item in result.get("block_index") or []:
        if not isinstance(item, dict):
            continue
        _source_ask_add_snippet(
            snippets,
            seen,
            source="block_index",
            title=str(item.get("title") or item.get("kind") or item.get("anchor") or "原文块"),
            text_value=item.get("text"),
            anchor=str(item.get("anchor") or "") or None,
            page_no=item.get("page_no") if isinstance(item.get("page_no"), int) else None,
            preview_page_no=item.get("preview_page_no") if isinstance(item.get("preview_page_no"), int) else None,
        )

    for section in result.get("located_sections") or []:
        if not isinstance(section, dict):
            continue
        _source_ask_add_snippet(
            snippets,
            seen,
            source="located_sections",
            title=str(section.get("title") or section.get("section_title") or section.get("section_id") or "相关章节"),
            text_value=section.get("content") or section.get("text") or section.get("quote"),
            anchor=str(section.get("anchor_start") or "") or None,
            page_no=section.get("page_no") if isinstance(section.get("page_no"), int) else None,
            preview_page_no=section.get("preview_page_no") if isinstance(section.get("preview_page_no"), int) else None,
        )
    return snippets


def _source_ask_rank_snippets(question: str, result: Dict[str, Any], top_k: int = 8) -> list[Dict[str, Any]]:
    snippets = _source_ask_collect_snippets(result)
    expanded_question = _source_ask_expand_question(question)
    compact_question = re.sub(r"\s+", "", str(question or ""))
    for snippet in snippets:
        score = _source_ask_snippet_score(expanded_question, snippet)
        asks_deadline = any(term in compact_question for term in ("截止", "什么时候", "何时", "时间", "日期"))
        snippet_text = str(snippet.get("text") or "")
        compact_snippet = re.sub(r"\s+", "", snippet_text)
        if asks_deadline and re.search(r"\d{4}年|\d{1,2}月\d{1,2}日|\d{1,2}[:：]\d{2}|截止|前到账|前交纳|前缴纳", compact_snippet):
            score += 18.0
        if asks_deadline and any(term in compact_snippet for term in ("不予退还", "退还情形", "将不予退还")):
            score -= 24.0
        snippet["score"] = score
    ranked = [item for item in sorted(snippets, key=lambda item: item.get("score", 0), reverse=True) if item.get("score", 0) > 0]
    if not ranked:
        ranked = snippets[: max(1, top_k)]
    return ranked[: max(1, min(top_k, 18))]

def _source_ask_citations_from_snippets(snippets: list[Dict[str, Any]]) -> list[SourceAskCitation]:
    citations: list[SourceAskCitation] = []
    for idx, item in enumerate(snippets, 1):
        citations.append(
            SourceAskCitation(
                id=f"S{idx}",
                title=str(item.get("title") or ""),
                quote=_source_ask_compact_text(item.get("text"), 320),
                anchor=item.get("anchor"),
                page_no=item.get("page_no") if isinstance(item.get("page_no"), int) else None,
                preview_page_no=item.get("preview_page_no") if isinstance(item.get("preview_page_no"), int) else None,
                score=float(item.get("score") or 0.0),
            )
        )
    return citations


def _source_ask_resolve_citation_pages(project_id: str, result: Dict[str, Any], citations: list[SourceAskCitation]) -> None:
    anchors = [str(item.anchor or "").strip() for item in citations if str(item.anchor or "").strip()]
    locations: dict[str, Dict[str, Any]] = {}
    if anchors:
        try:
            locations = _resolve_anchor_locations(project_id, result, anchors)
        except Exception as exc:
            logger.debug("source ask citation anchor resolve failed for {}: {}", project_id, str(exc)[:160])
    for item in citations:
        anchor = str(item.anchor or "").strip()
        if not anchor:
            continue
        loc = locations.get(anchor) or {}
        if item.preview_page_no is None and isinstance(loc.get("preview_page_no"), int):
            item.preview_page_no = loc.get("preview_page_no")
        if item.page_no is None and isinstance(loc.get("page_no"), int):
            item.page_no = loc.get("page_no")

    unresolved = [
        item
        for item in citations
        if item.preview_page_no is None and len(_normalize_match_text(item.quote)) >= 8
    ]
    if not unresolved:
        return
    try:
        query_locations = _resolve_query_locations(
            project_id,
            result,
            [SourceLocateQuery(id=item.id, text=item.quote) for item in unresolved],
        )
    except Exception as exc:
        logger.debug("source ask citation text resolve failed for {}: {}", project_id, str(exc)[:160])
        return
    for item in unresolved:
        page = (query_locations.get(item.id) or {}).get("preview_page_no")
        if isinstance(page, int) and page > 0:
            item.preview_page_no = page


def _source_ask_fallback_answer(question: str, citations: list[SourceAskCitation]) -> str:
    if not citations:
        return "没有在当前项目的原文索引中检索到相关内容。"
    lines = ["根据当前检索到的原文片段："]
    for item in citations[:3]:
        prefix = f"[{item.id}]"
        title = f"{item.title}：" if item.title else ""
        lines.append(f"{prefix} {title}{item.quote}")
    lines.append("如需确认，请点击引用跳转原文页核对。")
    return "\n".join(lines)


def _source_ask_prompt(question: str, citations: list[SourceAskCitation]) -> str:
    source_text = "\n\n".join(
        f"[{item.id}] 标题：{item.title or '原文片段'}\n锚点：{item.anchor or '-'} 页码：{item.page_no or item.preview_page_no or '-'}\n证据：{item.quote}"
        for item in citations
    )
    return f"""你是招标文件原文问答助手。请只根据【原文证据】回答用户问题。

规则：
1. 不允许使用外部知识或猜测。
2. 证据可能包含两类：原文片段，以及系统从原文抽取出的结构化事实。两类都可以使用。
3. 如果证据不足以回答，请明确说“当前原文证据未说明”，不要硬猜。
4. 回答要简洁、面向业务人员；如果用户问时间/金额/地点，请优先给出明确值并说明来源；如果用户问目录/清单，请用编号列表。
5. citation_ids 只能填写下面给出的片段编号，如 S1、S2。

【用户问题】
{question}

【原文证据】
{source_text}
"""

def _resolve_anchor_locations(project_id: str, result: Dict[str, Any], anchors: list[str]) -> dict[str, Dict[str, Any]]:
    block_index = result.get("block_index") or []
    block_lookup = {
        str(item.get("anchor") or ""): item
        for item in block_index
        if str(item.get("anchor") or "").strip()
    }
    requested = []
    for anchor in anchors or []:
        key = str(anchor or "").strip()
        if key and key in block_lookup:
            requested.append(key)
    if not requested:
        return {}

    source_file_path = _source_path_for_preview(project_id, result)
    if not source_file_path:
        return {}
    preview_pdf = _ensure_source_preview_pdf(project_id, Path(source_file_path))
    pages = _load_preview_page_texts(preview_pdf)

    def anchor_sort_key(anchor: str) -> tuple[int, str]:
        match = re.search(r"(\d+)", anchor)
        return (int(match.group(1)) if match else 10**9, anchor)

    resolved: dict[str, Dict[str, Any]] = {}
    last_preview_page = 1
    ordered_anchors = sorted(dict.fromkeys(requested), key=anchor_sort_key)
    for anchor in ordered_anchors:
        block = block_lookup.get(anchor) or {}
        display_page_no = block.get("page_no")
        display_page_no_end = block.get("page_no_end")
        if display_page_no_end is None:
            display_page_no_end = display_page_no

        preview_page_no = _locate_text_page(str(block.get("text") or ""), pages, start_page=last_preview_page)
        preview_page_no_end = preview_page_no
        if isinstance(preview_page_no, int) and preview_page_no > 0:
            last_preview_page = preview_page_no

        resolved[anchor] = {
            "anchor": anchor,
            # page_no is the page label extracted from the source document. It is
            # kept for display only; PDF iframe jumps must use preview_page_no.
            "page_no": display_page_no,
            "page_no_end": display_page_no_end,
            "preview_page_no": preview_page_no,
            "preview_page_no_end": preview_page_no_end,
            "text": block.get("text"),
            "kind": block.get("kind"),
        }

    # Headings or short titles are often too brief to match reliably in the PDF text layer.
    # When neighboring anchors resolve to the same preview page, inherit that page as a safe fallback.
    for idx, anchor in enumerate(ordered_anchors):
        item = resolved.get(anchor) or {}
        if isinstance(item.get("preview_page_no"), int):
            continue

        prev_page = None
        next_page = None
        if idx > 0:
            prev_page = resolved.get(ordered_anchors[idx - 1], {}).get("preview_page_no")
        if idx + 1 < len(ordered_anchors):
            next_page = resolved.get(ordered_anchors[idx + 1], {}).get("preview_page_no")

        inferred_page = None
        if isinstance(prev_page, int) and isinstance(next_page, int):
            inferred_page = prev_page if prev_page == next_page else prev_page
        elif isinstance(prev_page, int):
            inferred_page = prev_page
        elif isinstance(next_page, int):
            inferred_page = next_page

        if isinstance(inferred_page, int) and inferred_page > 0:
            item["preview_page_no"] = inferred_page
            item["preview_page_no_end"] = inferred_page

    return resolved


def _resolve_query_locations(
    project_id: str,
    result: Dict[str, Any],
    query_items: list[SourceLocateQuery],
) -> dict[str, Dict[str, Any]]:
    """Locate fact quotes in the preview PDF, whose page order drives the iframe."""
    if not query_items:
        return {}
    source_file_path = _source_path_for_preview(project_id, result)
    if not source_file_path:
        return {}
    preview_pdf = _ensure_source_preview_pdf(project_id, Path(source_file_path))
    pages = _load_preview_page_texts(preview_pdf)
    locations: dict[str, Dict[str, Any]] = {}
    for query in query_items[:180]:
        query_id = str(query.id or "").strip()
        text = str(query.text or "").strip()
        if not query_id or len(_normalize_match_text(text)) < 8:
            continue
        preview_page_no = _locate_text_page(text, pages, start_page=1)
        if isinstance(preview_page_no, int) and preview_page_no > 0:
            locations[query_id] = {"preview_page_no": preview_page_no}
    return locations


@_trace_task
async def process_tender_task(
    task_id: str,
    file_path: Path,
    company_id: str,
    company_name: str,
    generate_mode: str = "balanced",
):
    try:
        task = tasks[task_id]
        step_times: Dict[str, float] = {}
        step_summaries: Dict[str, Dict[str, Any]] = {}
        agent_metrics: Dict[str, Dict[str, Any]] = {}
        llm_gateway.reset_usage()

        def _record_agent_metric(agent_id: str, duration_sec: float) -> None:
            agent_metrics[agent_id] = {
                "duration_sec": round(float(duration_sec or 0), 3),
            }
            task["agent_metrics"] = dict(agent_metrics)

        # ============== Step 1: 解析文档 ==============
        task["original_source_file_path"] = str(file_path)
        # Convert legacy .doc files to .docx before parsing.
        if file_path.suffix.lower() == ".doc":
            task["message"] = "正在转换 .doc 格式"
            task["current_node"] = "parser"
            _save_task_meta(task)
            logger.info(f"Task {task_id} converting .doc to .docx")
            file_path = convert_doc_to_docx(file_path)
        task["source_file_path"] = str(file_path)

        task["status"] = TaskStatus.PARSING
        task["progress"] = 10
        task["message"] = "正在解析招标文件"
        task["current_node"] = "parser"
        _save_task_meta(task)
        logger.info(f"Task {task_id} start parsing: {file_path}")

        if task.get("cancelled"):
            raise asyncio.CancelledError()
        t_parse = _time.time()
        parsed = parse_document(str(file_path))
        step_times["parsing"] = _time.time() - t_parse
        task["step_times"] = dict(step_times)
        _save_task_meta(task)
        step_summaries["parser"] = {
            "inputs": {"file_path": str(file_path)},
            "outputs": {
                "flat_sections": len(parsed.flat_sections),
                "full_text_chars": len(parsed.full_text),
            },
            "duration_sec": round(step_times["parsing"], 3),
        }
        task["step_summaries"] = dict(step_summaries)
        _record_agent_metric("parser", step_times["parsing"])
        _save_task_meta(task)
        logger.info(
            f"Task {task_id} parsed, {len(parsed.flat_sections)} sections, "
            f"{len(parsed.full_text)} chars"
        )

        if task.get("cancelled"):
            raise asyncio.CancelledError()

        # ============== Step 2: Locate relevant tender sections ==============
        task["status"] = TaskStatus.LOCATING
        task["progress"] = 25
        task["message"] = "正在智能定位关键章节"
        task["current_node"] = "locator"
        _save_task_meta(task)

        # title info can be generated later in graph title node.
        title_info_from_rules = {}

        t_loc = _time.time()
        located_sections = await _locate_sections_stage(parsed)



        logger.info("[V12] skip pre-title extraction future; title will be produced in graph")

        if not located_sections:
            located_sections = locate_route_sections_fallback(parsed)
            if not located_sections:
                raise RuntimeError(
                    "未能在文档中定位到与投标目录相关的章节。"
                    "请确认上传的是有效的招标/磋商/比选文件。"
                )

        # Load located section content and measure locating stage time.
        sections_with_content = assemble_section_content(parsed, located_sections)
        step_times["locating"] = _time.time() - t_loc
        task["step_times"] = dict(step_times)
        _save_task_meta(task)
        step_summaries["locator"] = {
            "inputs": {"candidate_sections": len(located_sections)},
            "outputs": {"located_sections": len(sections_with_content)},
            "duration_sec": round(step_times["locating"], 3),
        }
        task["step_summaries"] = dict(step_summaries)
        _record_agent_metric("locator", step_times["locating"])
        _save_task_meta(task)
        logger.info(
            f"Task {task_id} located {len(sections_with_content)} relevant sections"
        )

        if task.get("cancelled"):
            raise asyncio.CancelledError()
        # ============== Step 3: Prepare graph input ==============
        task["status"] = TaskStatus.COMPOSING
        task["progress"] = 50
        task["message"] = "正在根据招标要求生成投标目录"
        task["current_node"] = "composer"
        _save_task_meta(task)

        # V12: 标题信息由 graph 中 title 节点生成
        initial_state = {
            "project_id": task_id,
            "file_name": file_path.name,
            "source_file_path": str(file_path),
            "company_id": company_id,
            "company_name": company_name,
            "head_text": parsed.full_text[:2000],
            "located_sections": sections_with_content,
            "requirement_source_sections": _requirement_source_sections(parsed),
            "block_index": parsed.block_index,
            "warnings": [],
            "title_info": title_info_from_rules if title_info_from_rules else {},
        }

        if task.get("cancelled"):
            raise asyncio.CancelledError()
        # ============== Step 4: Run LangGraph ==============
        task["status"] = TaskStatus.COMPOSING
        task["progress"] = 80
        task["message"] = "正在抽取招标要求并生成目录草稿"
        task["current_node"] = "requirement_extractor"
        _save_task_meta(task)

        t_graph = _time.time()
        task["current_node"] = "title"
        final_state = dict(initial_state)
        graph_nodes = [
            "title",
            "requirement_extractor",
            "composer",
        ]
        summary_key_map = {
            "title": "title",
            "composer": "composer",
            "requirement_extractor": "requirements",
        }
        node_started_at: Dict[str, float] = {"title": _time.time()}
        async for update in outline_review_graph.astream(initial_state, stream_mode="updates"):
            if not isinstance(update, dict):
                continue
            for node_name, payload in update.items():
                if isinstance(payload, dict):
                    final_state.update(payload)
                if node_name in graph_nodes:
                    # update arrives when this node finishes in LangGraph
                    now = _time.time()
                    summary_key = summary_key_map[node_name]
                    started = node_started_at.get(node_name, now)
                    duration_sec = max(0.0, now - started)
                    step_summaries.setdefault(summary_key, {})
                    step_summaries[summary_key]["duration_sec"] = round(duration_sec, 3)
                    _record_agent_metric(summary_key, duration_sec)
                    task["step_summaries"] = dict(step_summaries)
                    idx = graph_nodes.index(node_name)
                    if idx + 1 < len(graph_nodes):
                        next_node = graph_nodes[idx + 1]
                        node_started_at.setdefault(next_node, now)
                        task["current_node"] = next_node
                        task["progress"] = min(98, 80 + (idx + 1) * 3)
                    else:
                        task["current_node"] = "composer"
                        task["progress"] = 99
                    _save_task_meta(task)
        step_times["graph_total"] = _time.time() - t_graph
        task["step_times"] = dict(step_times)
        _save_task_meta(task)

        # ============== Step 5: Finalize result ==============
        title_info = final_state.get("title_info", {}) or {}
        outline = final_state.get("final_outline") or final_state.get("outline") or []
        merged_outline = final_state.get("merged_outline", []) or []
        material_assignments: list[dict] = []
        generated_sections: dict[str, str] = {}
        compliance_report: dict[str, Any] = {}
        consistency_report: dict[str, Any] = {}
        warnings = final_state.get("warnings", []) or []
        step_summaries["composer"] = {
            "inputs": {"located_sections": len(sections_with_content)},
            "outputs": {"outline_nodes": len(outline)},
            "duration_sec": step_summaries.get("composer", {}).get("duration_sec"),
        }
        tender_requirements = final_state.get("tender_requirements", {}) or {}
        tender_requirements_stats = final_state.get("tender_requirements_stats", {}) or {}
        step_summaries["requirements"] = {
            "inputs": {"located_sections": len(sections_with_content)},
            "outputs": {
                "technical_requirements": len(tender_requirements.get("technical_requirements", []) or []),
                "invalidation": len(tender_requirements.get("invalidation", []) or []),
                "material_checklist": len(tender_requirements.get("material_checklist", []) or []),
            },
            "duration_sec": step_summaries.get("requirements", {}).get("duration_sec") or tender_requirements_stats.get("elapsed_sec"),
        }
        step_summaries["material_mapper"] = {
            "inputs": {"outline_nodes": len(outline)},
            "outputs": {"assigned_nodes": 0, "skipped": True},
            "duration_sec": 0,
        }
        step_summaries["rag"] = {
            "inputs": {"assigned_nodes": 0},
            "outputs": {
                "context_nodes": 0,
                "fact_nodes": 0,
                "high_confidence_nodes": 0,
            },
            "duration_sec": 0,
        }
        step_summaries["content"] = {
            "inputs": {"context_nodes": 0},
            "outputs": {"generated_sections": 0},
            "duration_sec": 0,
        }
        step_summaries["compliance"] = {
            "inputs": {"outline_nodes": len(outline)},
            "outputs": {"issues": 0, "skipped": True},
            "duration_sec": 0,
        }
        step_summaries["consistency"] = {
            "inputs": {"outline_nodes": len(outline)},
            "outputs": {"conflicts": 0, "skipped": True},
            "duration_sec": 0,
        }

        render_decisions = build_render_decision_report(outline, material_assignments, generated_sections)

        output_data = {
            "company_id": company_id or DEFAULT_COMPANY_ID,
            "company_name": company_name,
            "title_info": title_info,
            "outline": outline,
            "outline_flat_raw": merged_outline,
            "tender_requirements": tender_requirements,
            "tender_requirements_stats": tender_requirements_stats,
            "material_assignments": material_assignments,
            "render_decisions": render_decisions,
            "rag_contexts": {},
            "retrieval_summary": {},
            "generated_sections": generated_sections,
            "project_facts": {},
            "compliance_report": compliance_report,
            "consistency_report": consistency_report,
            "workflow_stage": "outline_review",
            "source_file_path": str(file_path),
            "original_source_file_path": task.get("original_source_file_path"),
            "block_index": parsed.block_index,
            "located_sections": sections_with_content,
            "warnings": warnings,
            "stats": {
                "located_sections": [
                    {
                        "id": s["section_id"],
                        "title": s["title"],
                        "relevance": s["relevance"],
                        "anchor_start": s.get("anchor_start"),
                        "anchor_end": s.get("anchor_end"),
                    }
                    for s in sections_with_content
                ],
                "section_count": len(sections_with_content),
                "total_extracted": _count_outline_nodes(outline),
                "generated_section_count": len(generated_sections),
                "compliance_issue_count": len(compliance_report.get("issues", [])),
                "consistency_conflict_count": len(consistency_report.get("conflicts", [])),
                "rule_version": compliance_report.get("rule_version", "v1"),
            },
        }

        output_file = OUTPUT_DIR / f"{task_id}.json"
        output_file.write_text(
            json.dumps(output_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        md_content = render_outline_markdown(
            outline, title=title_info.get("title", file_path.stem)
        )
        md_file = OUTPUT_DIR / f"{task_id}.md"
        md_file.write_text(md_content, encoding="utf-8")

        task["result"] = output_data
        task["step_times"] = step_times
        task["step_summaries"] = step_summaries
        task["agent_metrics"] = dict(agent_metrics)
        task["status"] = TaskStatus.OUTLINE_REVIEW
        task["progress"] = 100
        task["message"] = "Outline draft ready; confirm before material mapping"
        task["current_node"] = "outline_review"
        _save_task_meta(task)
        logger.info(f"Task {task_id} completed")
    except asyncio.CancelledError:
        logger.info(f"Task {task_id} cancelled by user")
        task["status"] = TaskStatus.CANCELLED
        task["error"] = "任务已被用户取消"
        task["message"] = "任务已取消"
        task["current_node"] = "cancelled"
        _save_task_meta(task)
    except Exception as e:
        logger.error(f"Task {task_id} failed: {str(e)}", exc_info=True)
        task["status"] = TaskStatus.FAILED
        task["error"] = str(e)
        task["message"] = f"处理失败: {str(e)[:100]}"
        task["current_node"] = "failed"
        _save_task_meta(task)


# ============== API 接口 ==============

@app.get("/api/health")
async def health():
    return {"status": "ok"}


def _count_outline_nodes(nodes: list) -> int:
    total = 0
    for node in nodes or []:
        total += 1
        total += _count_outline_nodes(node.get("children") or [])
    return total


def _requirement_source_sections(parsed) -> list[dict[str, Any]]:
    """Expose full parsed sections to high-recall fact extraction only."""
    rows: list[dict[str, Any]] = []
    for index, section in enumerate(parsed.flat_sections or []):
        start_idx = getattr(section, "start_item_idx", None)
        end_idx = getattr(section, "end_item_idx", None)
        rows.append(
            {
                "section_id": str(getattr(section, "id", "") or f"full-{index + 1}"),
                "title": str(getattr(section, "title", "") or ""),
                "relevance": "full_document_requirement_source",
                "content": str(getattr(section, "content", "") or ""),
                "anchor_start": f"p{start_idx}" if isinstance(start_idx, int) else None,
                "anchor_end": f"p{end_idx}" if isinstance(end_idx, int) else None,
                "page_no": getattr(section, "page_no", None),
            }
        )
    return rows


def _refresh_render_decisions(result: Dict[str, Any]) -> list[dict]:
    """Refresh renderer decision audit data for the current result payload."""
    if not isinstance(result, dict):
        return []
    try:
        decisions = build_render_decision_report(
            result.get("outline", []) or [],
            result.get("material_assignments", []) or [],
            result.get("generated_sections", {}) or {},
        )
    except Exception as exc:
        logger.warning("[render-plan] failed to build render decisions: {}", str(exc)[:160])
        decisions = []
    result["render_decisions"] = decisions
    return decisions


def _has_index_table_source(nodes: list) -> bool:
    for node in nodes or []:
        if not isinstance(node, dict):
            continue
        if str(node.get("source")) == "index_table":
            return True
        if _has_index_table_source(node.get("children") or []):
            return True
    return False


def _preserve_existing_index_outline_if_source_lost(old_outline: list, incoming_outline: list) -> list:
    """Avoid stale clients overwriting authoritative index-table outlines without source tags."""
    if not _has_index_table_source(old_outline):
        return incoming_outline
    if _has_index_table_source(incoming_outline):
        return incoming_outline
    old_names = {
        str(node.get("name") or "").strip()
        for node in old_outline or []
        if isinstance(node, dict) and str(node.get("name") or "").strip()
    }
    incoming_names = {
        str(node.get("name") or "").strip()
        for node in incoming_outline or []
        if isinstance(node, dict) and str(node.get("name") or "").strip()
    }
    if old_names and not old_names.issubset(incoming_names):
        logger.warning(
            "[outline] incoming outline lost index_table source and root groups; preserving existing authoritative outline"
        )
        return old_outline
    return incoming_outline


def _project_item_from_files(project_id: str, meta: Dict[str, Any], result: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    result_dict = result if isinstance(result, dict) else {}
    title_info = result_dict.get("title_info", {}) or {}
    stats = result_dict.get("stats", {}) or {}
    output_file = OUTPUT_DIR / f"{project_id}.json"
    task_file = _task_meta_path(project_id)
    docx_file = OUTPUT_DIR / f"{project_id}_blank_bid.docx"
    updated_candidates = [p.stat().st_mtime for p in (task_file, output_file, docx_file) if p.exists()]
    updated_at = None
    if updated_candidates:
        updated_at = datetime.fromtimestamp(max(updated_candidates)).isoformat()

    outline = result_dict.get("outline", []) or []
    generated_sections = result_dict.get("generated_sections", {}) or {}
    if not isinstance(generated_sections, dict):
        generated_sections = {}
    compliance_report = result_dict.get("compliance_report", {}) or {}
    consistency_report = result_dict.get("consistency_report", {}) or {}
    return {
        "id": project_id,
        "filename": meta.get("filename") or project_id,
        "project_name": title_info.get("project_name") or title_info.get("title") or meta.get("filename") or project_id,
        "purchaser": title_info.get("purchaser") or "",
        "company_id": meta.get("company_id") or DEFAULT_COMPANY_ID,
        "company_name": meta.get("company_name") or "",
        "status": meta.get("status") or (TaskStatus.DONE if result_dict else TaskStatus.PENDING),
        "progress": int(meta.get("progress") or (100 if result_dict else 0)),
        "message": meta.get("message") or "",
        "error": meta.get("error"),
        "created_at": meta.get("created_at"),
        "updated_at": updated_at,
        "has_result": output_file.exists(),
        "has_docx": docx_file.exists(),
        "outline_count": _count_outline_nodes(outline),
        "generated_section_count": int(stats.get("generated_section_count") or len(generated_sections)),
        "compliance_issue_count": int(
            stats.get("compliance_issue_count")
            or len((compliance_report if isinstance(compliance_report, dict) else {}).get("issues", []) or [])
        ),
        "consistency_conflict_count": int(
            stats.get("consistency_conflict_count")
            or len((consistency_report if isinstance(consistency_report, dict) else {}).get("conflicts", []) or [])
        ),
    }


@app.get("/api/projects", response_model=list[ProjectListItem], summary="List generated projects")
async def list_projects(request: Request, limit: int = 100):
    current_user_id = request.state.current_user.user_id
    project_ids = set()
    for path in OUTPUT_DIR.glob("*.task.json"):
        project_ids.add(path.name.replace(".task.json", ""))
    for path in OUTPUT_DIR.glob("*.json"):
        if not path.name.endswith(".task.json"):
            project_ids.add(path.stem)

    items = []
    for project_id in project_ids:
        meta: Dict[str, Any] = {"id": project_id}
        result: Optional[Dict[str, Any]] = None
        meta_file = _task_meta_path(project_id)
        output_file = OUTPUT_DIR / f"{project_id}.json"
        if meta_file.exists():
            try:
                meta.update(json.loads(meta_file.read_text(encoding="utf-8")))
            except Exception as exc:
                logger.warning(f"read project meta failed {project_id}: {exc}")
        if str(meta.get("creator_user_id") or "").strip() != current_user_id:
            continue
        if output_file.exists():
            try:
                result = json.loads(output_file.read_text(encoding="utf-8"))
            except Exception as exc:
                logger.warning(f"read project result failed {project_id}: {exc}")
        items.append(_project_item_from_files(project_id, meta, result))

    items.sort(key=lambda item: item.get("updated_at") or item.get("created_at") or "", reverse=True)
    return items[: max(1, min(limit, 500))]


@app.delete("/api/projects/{project_id}", summary="Delete project and generated files")
async def delete_project(project_id: str):
    try:
        normalized_id = str(uuid.UUID(str(project_id)))
    except (ValueError, TypeError, AttributeError) as exc:
        raise HTTPException(status_code=400, detail="项目 ID 格式无效") from exc
    if normalized_id != str(project_id).lower():
        raise HTTPException(status_code=400, detail="项目 ID 格式无效")

    project_files = [
        OUTPUT_DIR / f"{normalized_id}.task.json",
        OUTPUT_DIR / f"{normalized_id}.json",
        OUTPUT_DIR / f"{normalized_id}.md",
        OUTPUT_DIR / f"{normalized_id}_blank_bid.docx",
        _source_preview_pdf_path(normalized_id),
        _source_preview_meta_path(normalized_id),
        *(UPLOAD_DIR / f"{normalized_id}{suffix}" for suffix in (".doc", ".docx", ".pdf")),
    ]
    if normalized_id not in tasks and not any(path.exists() for path in project_files):
        raise HTTPException(status_code=404, detail="项目不存在")

    running_task = background_tasks.pop(normalized_id, None)
    if running_task and not running_task.done():
        running_task.cancel()
        try:
            await running_task
        except asyncio.CancelledError:
            pass

    tasks.pop(normalized_id, None)
    deleted_files = 0
    for path in project_files:
        try:
            if path.exists() and path.is_file():
                path.unlink()
                deleted_files += 1
        except OSError as exc:
            logger.warning("delete project file failed project={} path={} error={}", normalized_id, path, exc)
            raise HTTPException(status_code=500, detail=f"项目文件删除失败：{path.name}") from exc

    return {"ok": True, "project_id": normalized_id, "deleted_files": deleted_files}


@app.get("/api/companies", response_model=list[CompanyResponse], summary="List companies")
async def list_companies(include_inactive: bool = False):
    from tender_agent.core.db import get_db_session
    from data.knowledge.models import Company

    _ensure_companies_ready()
    db = next(get_db_session())
    try:
        query = db.query(Company)
        if not include_inactive:
            query = query.filter(Company.is_active == True)
        rows = query.order_by(Company.is_default.desc(), Company.name.asc()).all()
        if not rows:
            return [
                {
                    "id": DEFAULT_COMPANY_ID,
                    "name": DEFAULT_COMPANY_NAME,
                    "is_default": True,
                    "is_active": True,
                }
            ]
        return [
            {
                "id": row.id,
                "name": row.name,
                "is_default": bool(row.is_default),
                "is_active": bool(row.is_active),
            }
            for row in rows
        ]
    except Exception as exc:
        logger.warning("[companies] list failed, fallback to default: {}", str(exc)[:120])
        return [
            {
                "id": DEFAULT_COMPANY_ID,
                "name": DEFAULT_COMPANY_NAME,
                "is_default": True,
                "is_active": True,
            }
        ]
    finally:
        db.close()


@app.post("/api/companies", response_model=CompanyResponse, summary="Create company")
async def create_company(payload: CompanyCreateRequest):
    from tender_agent.core.db import get_db_session
    from data.knowledge.models import Company

    _ensure_companies_ready()
    name = payload.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="公司名称不能为空")
    company_id = _normalize_company_id(payload.id or f"company-{uuid.uuid4().hex[:8]}")
    db = next(get_db_session())
    try:
        existing = db.query(Company).filter(Company.id == company_id).first()
        if existing is not None:
            raise HTTPException(status_code=409, detail="公司ID已存在，请换一个ID")
        if payload.is_default:
            db.query(Company).update({Company.is_default: False}, synchronize_session=False)
        row = Company(id=company_id, name=name, is_default=bool(payload.is_default), is_active=True)
        db.add(row)
        db.commit()
        _ensure_company_storage(company_id)
        return {
            "id": row.id,
            "name": row.name,
            "is_default": bool(row.is_default),
            "is_active": bool(row.is_active),
        }
    finally:
        db.close()


@app.put("/api/companies/{company_id}", response_model=CompanyResponse, summary="Update company")
async def update_company(company_id: str, payload: CompanyUpdateRequest):
    from tender_agent.core.db import get_db_session
    from data.knowledge.models import Company

    _ensure_companies_ready()
    db = next(get_db_session())
    try:
        row = db.query(Company).filter(Company.id == company_id).first()
        if row is None:
            raise HTTPException(status_code=404, detail="公司不存在")
        if payload.name is not None:
            name = payload.name.strip()
            if not name:
                raise HTTPException(status_code=400, detail="公司名称不能为空")
            row.name = name
        if payload.is_active is not None:
            if company_id == DEFAULT_COMPANY_ID and payload.is_active is False:
                raise HTTPException(status_code=400, detail="默认广东公司不能停用")
            row.is_active = bool(payload.is_active)
        if payload.is_default is not None:
            if payload.is_default:
                if not row.is_active:
                    raise HTTPException(status_code=400, detail="停用公司不能设为默认")
                db.query(Company).update({Company.is_default: False}, synchronize_session=False)
                row.is_default = True
            elif row.is_default:
                raise HTTPException(status_code=400, detail="请先把其他公司设为默认")
        db.commit()
        _ensure_company_storage(row.id)
        return {
            "id": row.id,
            "name": row.name,
            "is_default": bool(row.is_default),
            "is_active": bool(row.is_active),
        }
    finally:
        db.close()


@app.delete("/api/companies/{company_id}", summary="Deactivate company")
async def deactivate_company(company_id: str):
    from tender_agent.core.db import get_db_session
    from data.knowledge.models import Company

    if company_id == DEFAULT_COMPANY_ID:
        raise HTTPException(status_code=400, detail="默认广东公司不能停用")
    _ensure_companies_ready()
    db = next(get_db_session())
    try:
        row = db.query(Company).filter(Company.id == company_id).first()
        if row is None:
            raise HTTPException(status_code=404, detail="公司不存在")
        if row.is_default:
            raise HTTPException(status_code=400, detail="默认公司不能停用，请先切换默认公司")
        row.is_active = False
        db.commit()
        return {"ok": True}
    finally:
        db.close()


@app.post("/api/projects/upload")
async def upload_file(
    request: Request,
    file: UploadFile = File(...),
    company_id: str = Form(DEFAULT_COMPANY_ID),
    company_name: str = Form(...),
    generate_mode: str = Form("balanced"),
    auto_add_score: str = Form("1"),
):
    if not file.filename.lower().endswith((".docx", ".doc", ".pdf")):
        raise HTTPException(status_code=400, detail="Only .docx/.doc/.pdf are supported")

    task_id = str(uuid.uuid4())
    ext = Path(file.filename).suffix.lower()
    save_path = UPLOAD_DIR / f"{task_id}{ext}"
    content = await file.read()
    save_path.write_bytes(content)

    tasks[task_id] = {
        "id": task_id,
        "creator_user_id": request.state.current_user.user_id,
        "filename": file.filename,
        "company_id": company_id or DEFAULT_COMPANY_ID,
        "company_name": company_name,
        "generate_mode": generate_mode,
        "auto_add_score": auto_add_score == "1",
        "status": TaskStatus.PENDING,
        "progress": 0,
        "message": "任务已创建，等待处理",
        "error": None,
        "created_at": datetime.now().isoformat(),
        "result": None,
        "source_file_path": str(save_path),
        "original_source_file_path": str(save_path),
        "current_node": "parser",
    }
    _save_task_meta(tasks[task_id])

    background_task = asyncio.create_task(
        process_tender_task(
            task_id=task_id,
            file_path=save_path,
            company_id=company_id or DEFAULT_COMPANY_ID,
            company_name=company_name,
            generate_mode=generate_mode,
        )
    )
    background_tasks[task_id] = background_task
    background_task.add_done_callback(
        lambda completed, project_id=task_id: background_tasks.pop(project_id, None)
    )
    return {"project_id": task_id}


@app.get("/api/projects/{project_id}/status", response_model=TaskStatusResponse)
async def get_task_status(project_id: str):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    return {
        "status": task["status"],
        "progress": task["progress"],
        "message": task["message"],
        "error": task["error"],
        "current_node": task.get("current_node"),
        "step_times": task.get("step_times", {}),
        "step_summaries": task.get("step_summaries", {}),
        "agent_metrics": task.get("agent_metrics", {}),
    }

@app.post("/api/projects/{project_id}/cancel", summary="Cancel task")
async def cancel_project(project_id: str):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    # 已结束任务无需取消
    if task["status"] in [TaskStatus.OUTLINE_REVIEW, TaskStatus.DONE, TaskStatus.FAILED, TaskStatus.CANCELLED]:
        return {"message": "任务已结束，无需取消"}
    # Mark task as cancelled.
    task["cancelled"] = True
    task["status"] = TaskStatus.CANCELLED
    task["message"] = "任务已取消"
    _save_task_meta(task)
    running_task = background_tasks.get(project_id)
    if running_task and not running_task.done():
        running_task.cancel()
    return {"message": "任务取消成功"}


@app.get("/api/projects/{project_id}/outline", response_model=OutlineResponse)
async def get_outline(project_id: str):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if not _outline_available(task["status"]):
        raise HTTPException(status_code=400, detail="Task not completed")
    result = _ensure_task_result(project_id, task)
    if outline_has_user_source(result.get("outline", [])):
        result["outline"] = renumber_user_outline(result.get("outline", []))
    else:
        result["outline"] = normalize_outline_numbering(result.get("outline", []))
    analysis_facts = build_tender_analysis_facts(result.get("tender_requirements", {}))
    return {
        "company_id": result.get("company_id") or task.get("company_id") or DEFAULT_COMPANY_ID,
        "company_name": result.get("company_name") or task.get("company_name"),
        "title_info": result.get("title_info", {}),
        "outline": result["outline"],
        "tender_requirements": result.get("tender_requirements", {}),
        "tender_requirements_stats": result.get("tender_requirements_stats", {}),
        "analysis_facts": analysis_facts.get("groups", {}),
        "analysis_facts_summary": analysis_facts.get("summary", {}),
        "analysis_recommendations": result.get("analysis_recommendations", {}),
        "material_assignments": result.get("material_assignments", []),
        "render_decisions": _refresh_render_decisions(result),
        "retrieval_summary": result.get("retrieval_summary", {}),
        "generated_sections": result.get("generated_sections", {}),
        "compliance_report": result.get("compliance_report", {}),
        "consistency_report": result.get("consistency_report", {}),
        "workflow_stage": result.get("workflow_stage") or ("material_ready" if task["status"] == TaskStatus.DONE else "outline_review"),
        "source_file_path": result.get("source_file_path"),
        "original_source_file_path": result.get("original_source_file_path"),
        "block_index": result.get("block_index", []),
        "warnings": result.get("warnings", []),
        "stats": result.get("stats", {}),
    }


@app.get("/api/projects/{project_id}/analysis-facts", summary="Get traceable tender analysis facts")
async def get_analysis_facts(project_id: str):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if not _outline_available(task["status"]):
        raise HTTPException(status_code=400, detail="Task not completed")

    result = _ensure_task_result(project_id, task)
    facts = build_tender_analysis_facts(result.get("tender_requirements", {}))
    return {
        "facts": facts.get("groups", {}),
        "summary": facts.get("summary", {}),
        "recommendations": result.get("analysis_recommendations", {}),
    }


@app.post("/api/projects/{project_id}/analysis-recommendations", summary="Generate scoring and risk recommendations from analysis facts")
async def generate_analysis_recommendations(project_id: str, body: AnalysisRecommendationsRequest):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if not _outline_available(task["status"]):
        raise HTTPException(status_code=400, detail="Task not completed")

    result = _ensure_task_result(project_id, task)
    existing = result.get("analysis_recommendations")
    if isinstance(existing, dict) and existing.get("items") and not body.force:
        return existing

    facts = build_tender_analysis_facts(result.get("tender_requirements", {}))
    groups = facts.get("groups", {}) if isinstance(facts, dict) else {}
    scoring = groups.get("scoring", []) if isinstance(groups, dict) else []
    risks = groups.get("invalidation", []) if isinstance(groups, dict) else []
    if not scoring and not risks:
        return {"items": [], "generated_at": None, "used_llm": False}

    valid_fact_groups = {
        str(item.get("id")): "scoring"
        for item in scoring
        if isinstance(item, dict) and item.get("id")
    }
    valid_fact_groups.update({
        str(item.get("id")): "invalidation"
        for item in risks
        if isinstance(item, dict) and item.get("id")
    })
    try:
        llm_result: _AnalysisAdviceResponse = await llm_gateway.async_call_structured(
            strategy_prompt(facts),
            _AnalysisAdviceResponse,
            max_tokens=2600,
        )
    except Exception as exc:
        logger.warning("analysis recommendations LLM failed for project {}: {}", project_id, exc)
        raise HTTPException(status_code=502, detail="分析建议生成失败，请稍后重试") from exc

    items = []
    for item in llm_result.items or []:
        fact_id = str(item.fact_id or "")
        expected_group = valid_fact_groups.get(fact_id)
        suggestions = [str(value).strip() for value in item.suggestions if str(value).strip()][:4]
        if not expected_group or not suggestions:
            continue
        items.append({
            "fact_id": fact_id,
            "kind": expected_group,
            "suggestions": suggestions,
        })

    payload = {
        "items": items,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "used_llm": True,
    }
    result["analysis_recommendations"] = payload
    output_file = OUTPUT_DIR / f"{project_id}.json"
    output_file.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    task["result"] = result
    _save_task_meta(task)
    return payload


def _normalize_outline_key(value: Any) -> str:
    text = str(value or "")
    text = re.sub(r"\s+", "", text)
    text = re.sub(r"^\d+(?:\.\d+)*[.、\s]*", "", text)
    return text


def _flatten_outline_nodes(outline: list[dict]) -> list[dict]:
    items: list[dict] = []

    def walk(nodes: list[dict]) -> None:
        for node in nodes or []:
            items.append(node)
            children = node.get("children") or []
            if children:
                walk(children)

    walk(outline)
    return items


def _rekey_generated_sections(
    old_outline: list[dict],
    new_outline: list[dict],
    generated_sections: dict[str, str],
) -> dict[str, str]:
    """Keep existing drafts after outline renumber/reorder by matching names."""
    if not isinstance(generated_sections, dict) or not generated_sections:
        return {}

    old_by_id = {str(node.get("id", "")): node for node in _flatten_outline_nodes(old_outline)}
    old_name_to_content: dict[str, str] = {}
    kept: dict[str, str] = {}

    for old_id, content in generated_sections.items():
        text = str(content or "").strip()
        if not text:
            continue
        old_node = old_by_id.get(str(old_id))
        if old_node:
            old_name_to_content[_normalize_outline_key(old_node.get("name"))] = text

    for node in _flatten_outline_nodes(new_outline):
        children = node.get("children") or []
        if children:
            continue
        node_id = str(node.get("id", ""))
        name_key = _normalize_outline_key(node.get("name"))
        if name_key in old_name_to_content:
            kept[node_id] = old_name_to_content[name_key]
            continue
        if node_id in generated_sections and str(generated_sections[node_id] or "").strip():
            kept[node_id] = str(generated_sections[node_id]).strip()
    return kept


def _rekey_material_assignments(
    old_outline: list[dict],
    new_outline: list[dict],
    assignments: list[dict],
) -> list[dict]:
    """Keep material assignments after outline renumber/reorder by matching names."""
    if not isinstance(assignments, list) or not assignments:
        return []

    old_by_id = {str(node.get("id", "")): node for node in _flatten_outline_nodes(old_outline)}
    assignment_by_name: dict[str, dict] = {}

    for assignment in assignments:
        if not isinstance(assignment, dict):
            continue
        old_id = str(assignment.get("node_id", "") or assignment.get("outline_node_id", ""))
        old_node = old_by_id.get(old_id)
        name = old_node.get("name") if old_node else assignment.get("node_name") or assignment.get("name")
        name_key = _normalize_outline_key(name)
        materials = assignment.get("materials") or []
        if name_key and isinstance(materials, list) and materials:
            assignment_by_name[name_key] = assignment

    kept: list[dict] = []
    for node in _flatten_outline_nodes(new_outline):
        if node.get("children") or []:
            continue
        name_key = _normalize_outline_key(node.get("name"))
        assignment = assignment_by_name.get(name_key)
        if not assignment:
            continue
        copied = dict(assignment)
        copied["node_id"] = str(node.get("id", ""))
        copied["node_name"] = str(node.get("name", ""))
        kept.append(copied)
    return kept


def _merge_generated_sections_for_outline(
    outline: list[dict],
    existing: dict[str, str],
    regenerated: dict[str, str],
) -> dict[str, str]:
    """Merge regenerated content without overwriting existing user-visible drafts."""
    existing = existing if isinstance(existing, dict) else {}
    regenerated = regenerated if isinstance(regenerated, dict) else {}
    merged: dict[str, str] = {}
    valid_ids = {
        str(node.get("id", ""))
        for node in _flatten_outline_nodes(outline)
        if not (node.get("children") or [])
    }

    for node_id, content in regenerated.items():
        key = str(node_id)
        if key in valid_ids and str(content or "").strip():
            merged[key] = str(content).strip()

    for node_id, content in existing.items():
        key = str(node_id)
        if key in valid_ids and str(content or "").strip():
            merged[key] = str(content).strip()

    return merged


@app.put("/api/projects/{project_id}/outline")
async def update_outline(project_id: str, body: dict):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if not _outline_available(task["status"]):
        raise HTTPException(status_code=400, detail="Task not completed")
    output_file = OUTPUT_DIR / f"{project_id}.json"
    if not output_file.exists():
        raise HTTPException(status_code=404, detail="Project result file not found")

    import json as _json
    data = _json.loads(output_file.read_text(encoding="utf-8"))
    old_outline = data.get("outline", []) or []
    incoming_outline = _preserve_existing_index_outline_if_source_lost(
        old_outline,
        body.get("outline", []),
    )
    if outline_has_user_source(incoming_outline):
        outline = renumber_user_outline(incoming_outline)
    else:
        outline = normalize_outline_numbering(incoming_outline)
    data["generated_sections"] = _rekey_generated_sections(
        old_outline,
        outline,
        data.get("generated_sections", {}) or {},
    )
    data["material_assignments"] = _rekey_material_assignments(
        old_outline,
        outline,
        data.get("material_assignments", []) or [],
    )
    data["outline"] = outline
    data["material_assignments"] = []
    data["rag_contexts"] = {}
    data["retrieval_summary"] = {}
    data["generated_sections"] = {}
    data["compliance_report"] = {}
    data["consistency_report"] = {}
    data["workflow_stage"] = "outline_review"
    _refresh_render_decisions(data)
    output_file.write_text(_json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    # 兼容任务已恢复但 result 尚未装载到内存的场景
    if not isinstance(task.get("result"), dict):
        task["result"] = data
    else:
        task["result"]["outline"] = outline
        task["result"]["generated_sections"] = data["generated_sections"]
        task["result"]["material_assignments"] = data["material_assignments"]
        task["result"]["render_decisions"] = data.get("render_decisions", [])
        task["result"]["workflow_stage"] = "outline_review"

    task["status"] = TaskStatus.OUTLINE_REVIEW
    task["message"] = "Outline updated; confirm before material mapping"
    task["current_node"] = "outline_review"
    _save_task_meta(task)
    return {"ok": True}


@app.put("/api/projects/{project_id}/sections/{node_id}")
async def update_generated_section(project_id: str, node_id: str, body: SectionDraftRequest):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if not _outline_available(task["status"]):
        raise HTTPException(status_code=400, detail="Task not completed")

    output_file = OUTPUT_DIR / f"{project_id}.json"
    if not output_file.exists():
        raise HTTPException(status_code=404, detail="Project result file not found")

    data = json.loads(output_file.read_text(encoding="utf-8"))
    generated_sections = data.get("generated_sections") or {}
    if not isinstance(generated_sections, dict):
        generated_sections = {}
    generated_sections[str(node_id)] = str(body.content or "")
    data["generated_sections"] = generated_sections

    removed_sources = {str(source or "").strip() for source in body.remove_material_sources if str(source or "").strip()}
    material_assignments = data.get("material_assignments") or []
    if removed_sources and isinstance(material_assignments, list):
        updated_assignments: list[dict] = []
        for assignment in material_assignments:
            if not isinstance(assignment, dict) or str(assignment.get("node_id") or "") != str(node_id):
                updated_assignments.append(assignment)
                continue
            materials = assignment.get("materials") if isinstance(assignment.get("materials"), list) else []
            kept_materials = [
                material
                for material in materials
                if not isinstance(material, dict) or str(material.get("source") or "").strip() not in removed_sources
            ]
            if kept_materials:
                updated_assignments.append({**assignment, "materials": kept_materials})
        material_assignments = updated_assignments
        data["material_assignments"] = material_assignments

    stats = data.get("stats") if isinstance(data.get("stats"), dict) else {}
    stats["generated_section_count"] = len([v for v in generated_sections.values() if str(v or "").strip()])
    data["stats"] = stats
    _refresh_render_decisions(data)
    output_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    if not isinstance(task.get("result"), dict):
        task["result"] = data
    else:
        task["result"]["generated_sections"] = generated_sections
        task["result"]["material_assignments"] = material_assignments
        task["result"]["render_decisions"] = data.get("render_decisions", [])
        task["result"]["stats"] = stats
    _save_task_meta(task)
    return {
        "ok": True,
        "node_id": str(node_id),
        "generated_section_count": stats["generated_section_count"],
        "removed_material_sources": sorted(removed_sources),
        "render_decisions": data.get("render_decisions", []),
    }


@app.post("/api/projects/{project_id}/tender-template-preview")
async def preview_tender_template(project_id: str, body: TenderTemplatePreviewRequest):
    from tender_agent.knowledge.tender_template_preview import tender_template_html
    from tender_agent.knowledge.tender_template_copier import resolve_tender_template_span_by_node

    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if not _outline_available(task["status"]):
        raise HTTPException(status_code=400, detail="Task not completed")
    result = _ensure_task_result(project_id, task)
    source_path = _source_path_for_preview(project_id, result, task)
    if not source_path:
        return {"content_html": "", "format_source": "plain_text"}
    anchor_start = body.anchor_start
    anchor_end = body.anchor_end
    resolved = resolve_tender_template_span_by_node(
        {
            "name": body.name,
            "anchor_start": anchor_start,
            "anchor_end": anchor_end,
            "copy_method": body.copy_method,
        },
        source_path,
    )
    if resolved:
        anchor_start = resolved["anchor_start"]
        anchor_end = resolved["anchor_end"]
    return {
        "content_html": tender_template_html(source_path, anchor_start, anchor_end),
        "format_source": "docx_anchor" if anchor_start and anchor_end else "plain_text",
        **(resolved or {}),
    }


def _append_outline_children(nodes: list, parent_id: str, sections: list[Dict[str, Any]]) -> tuple[bool, list[dict]]:
    parent_id = str(parent_id or "").strip()
    inserted: list[dict] = []
    if not parent_id or not sections:
        return False, inserted

    def next_child_id(parent: dict, offset: int) -> str:
        base = str(parent.get("id") or "").strip()
        children = parent.get("children") if isinstance(parent.get("children"), list) else []
        max_suffix = 0
        prefix = f"{base}."
        for child in children:
            raw_id = str(child.get("id") or "")
            if raw_id.startswith(prefix):
                suffix = raw_id[len(prefix):].split(".", 1)[0]
                if suffix.isdigit():
                    max_suffix = max(max_suffix, int(suffix))
        return f"{base}.{max_suffix + offset}"

    def walk(items: list) -> bool:
        for node in items or []:
            if str(node.get("id") or "") == parent_id:
                children = node.get("children")
                if not isinstance(children, list):
                    children = []
                    node["children"] = children

                def append_section(parent: dict, section: Dict[str, Any]) -> None:
                    title = str(section.get("title") or section.get("name") or "").strip()
                    if not title:
                        return
                    parent_children = parent.get("children")
                    if not isinstance(parent_children, list):
                        parent_children = []
                        parent["children"] = parent_children
                    render_hook = section.get("render_hook") if isinstance(section.get("render_hook"), dict) else {}
                    section_id = str(render_hook.get("section_id") or "").strip()
                    if section_id and any(
                        str((child.get("render_hook") or {}).get("section_id") or "").strip() == section_id
                        for child in parent_children
                        if isinstance(child, dict)
                    ):
                        return
                    child = {
                        "id": next_child_id(parent, 1),
                        "name": title,
                        "required": bool(section.get("required", node.get("required", False))),
                        "has_template": bool(section.get("has_template", False)),
                        "_draft_content": str(section.get("content") or section.get("html") or "").strip(),
                        "children": [],
                    }
                    if render_hook:
                        child["render_hook"] = render_hook
                    if section.get("source"):
                        child["source"] = section.get("source")
                    parent_children.append(child)
                    inserted.append(child)
                    for nested in section.get("children") or []:
                        if isinstance(nested, dict):
                            append_section(child, nested)

                for section in sections:
                    append_section(node, section)
                return True
            if walk(node.get("children") if isinstance(node.get("children"), list) else []):
                return True
        return False

    return walk(nodes), inserted


def _collect_outline_ids(nodes: list) -> set[str]:
    ids: set[str] = set()
    for node in nodes or []:
        if not isinstance(node, dict):
            continue
        node_id = str(node.get("id") or "").strip()
        if node_id:
            ids.add(node_id)
        ids.update(_collect_outline_ids(node.get("children") if isinstance(node.get("children"), list) else []))
    return ids


def _remove_outline_node(nodes: list, node_id: str) -> tuple[bool, set[str]]:
    node_id = str(node_id or "").strip()
    if not node_id:
        return False, set()
    for index, node in enumerate(list(nodes or [])):
        if not isinstance(node, dict):
            continue
        if str(node.get("id") or "").strip() == node_id:
            removed = nodes.pop(index)
            return True, _collect_outline_ids([removed])
        children = node.get("children")
        if isinstance(children, list):
            ok, removed_ids = _remove_outline_node(children, node_id)
            if ok:
                return True, removed_ids
    return False, set()


def _find_outline_node(nodes: list, node_id: str) -> Optional[dict]:
    node_id = str(node_id or "").strip()
    if not node_id:
        return None
    for node in nodes or []:
        if not isinstance(node, dict):
            continue
        if str(node.get("id") or "").strip() == node_id:
            return node
        found = _find_outline_node(node.get("children") if isinstance(node.get("children"), list) else [], node_id)
        if found is not None:
            return found
    return None


@app.post("/api/projects/{project_id}/outline/children")
async def append_outline_children(project_id: str, body: AppendOutlineChildrenRequest):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if not _outline_available(task["status"]):
        raise HTTPException(status_code=400, detail="Task not completed")

    output_file = OUTPUT_DIR / f"{project_id}.json"
    if not output_file.exists():
        raise HTTPException(status_code=404, detail="Project result file not found")

    data = json.loads(output_file.read_text(encoding="utf-8"))
    outline = data.get("outline") if isinstance(data.get("outline"), list) else []
    ok, inserted = _append_outline_children(outline, body.parent_id, body.sections)
    if not ok:
        raise HTTPException(status_code=404, detail="Parent outline node not found")
    if not inserted:
        return {
            "ok": True,
            "outline": outline,
            "inserted": [],
            "generated_sections": data.get("generated_sections") or {},
            "render_decisions": data.get("render_decisions") or [],
            "stats": data.get("stats") or {},
        }

    generated_sections = data.get("generated_sections") or {}
    if not isinstance(generated_sections, dict):
        generated_sections = {}
    for child in inserted:
        draft_content = str(child.pop("_draft_content", "") or "").strip()
        generated_sections[str(child.get("id"))] = draft_content

    data["outline"] = outline
    data["generated_sections"] = generated_sections
    stats = data.get("stats") if isinstance(data.get("stats"), dict) else {}
    stats["outline_count"] = _count_outline_nodes(outline)
    stats["generated_section_count"] = len([v for v in generated_sections.values() if str(v or "").strip()])
    data["stats"] = stats
    _refresh_render_decisions(data)
    output_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    if not isinstance(task.get("result"), dict):
        task["result"] = data
    else:
        task["result"]["outline"] = outline
        task["result"]["generated_sections"] = generated_sections
        task["result"]["render_decisions"] = data.get("render_decisions", [])
        task["result"]["stats"] = stats
    _save_task_meta(task)
    return {
        "ok": True,
        "outline": outline,
        "inserted": inserted,
        "generated_sections": generated_sections,
        "render_decisions": data.get("render_decisions", []),
        "stats": stats,
    }


@app.delete("/api/projects/{project_id}/outline/nodes/{node_id}")
async def delete_outline_node(project_id: str, node_id: str):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if not _outline_available(task["status"]):
        raise HTTPException(status_code=400, detail="Task not completed")

    output_file = OUTPUT_DIR / f"{project_id}.json"
    if not output_file.exists():
        raise HTTPException(status_code=404, detail="Project result file not found")

    data = json.loads(output_file.read_text(encoding="utf-8"))
    outline = data.get("outline") if isinstance(data.get("outline"), list) else []
    ok, removed_ids = _remove_outline_node(outline, node_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Outline node not found")
    if not outline:
        raise HTTPException(status_code=400, detail="Cannot delete all outline nodes")

    generated_sections = data.get("generated_sections") or {}
    if isinstance(generated_sections, dict):
        for removed_id in removed_ids:
            generated_sections.pop(removed_id, None)
    else:
        generated_sections = {}

    assignments = []
    for assignment in data.get("material_assignments") or []:
        if not isinstance(assignment, dict):
            continue
        assignment_node_id = str(assignment.get("node_id") or assignment.get("outline_node_id") or "")
        if assignment_node_id not in removed_ids:
            assignments.append(assignment)

    data["outline"] = outline
    data["generated_sections"] = generated_sections
    data["material_assignments"] = assignments
    stats = data.get("stats") if isinstance(data.get("stats"), dict) else {}
    stats["outline_count"] = _count_outline_nodes(outline)
    stats["generated_section_count"] = len([v for v in generated_sections.values() if str(v or "").strip()])
    data["stats"] = stats
    _refresh_render_decisions(data)
    output_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    if not isinstance(task.get("result"), dict):
        task["result"] = data
    else:
        task["result"]["outline"] = outline
        task["result"]["generated_sections"] = generated_sections
        task["result"]["material_assignments"] = assignments
        task["result"]["render_decisions"] = data.get("render_decisions", [])
        task["result"]["stats"] = stats
    _save_task_meta(task)
    return {
        "ok": True,
        "outline": outline,
        "removed_ids": sorted(removed_ids),
        "generated_sections": generated_sections,
        "material_assignments": assignments,
        "render_decisions": data.get("render_decisions", []),
        "stats": stats,
    }


@app.put("/api/projects/{project_id}/outline/nodes/{node_id}/render-hook")
async def update_outline_node_render_hook(project_id: str, node_id: str, body: OutlineNodeHookRequest):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if not _outline_available(task["status"]):
        raise HTTPException(status_code=400, detail="Task not completed")

    output_file = OUTPUT_DIR / f"{project_id}.json"
    if not output_file.exists():
        raise HTTPException(status_code=404, detail="Project result file not found")

    action = str(body.action or "").strip()
    if action not in {"restore_master", "unlink_master"}:
        raise HTTPException(status_code=400, detail="Unsupported render hook action")

    data = json.loads(output_file.read_text(encoding="utf-8"))
    outline = data.get("outline") if isinstance(data.get("outline"), list) else []
    node = _find_outline_node(outline, node_id)
    if node is None:
        raise HTTPException(status_code=404, detail="Outline node not found")

    render_hook = node.get("render_hook") if isinstance(node.get("render_hook"), dict) else {}
    if str(render_hook.get("type") or "") != "tech_section":
        raise HTTPException(status_code=400, detail="Outline node is not a tech master hook")

    generated_sections = data.get("generated_sections") or {}
    if not isinstance(generated_sections, dict):
        generated_sections = {}

    if action == "restore_master":
        generated_sections.pop(str(node_id), None)
    elif action == "unlink_master":
        node.pop("render_hook", None)
        if str(node.get("source") or "") == "tech_master_hook":
            node.pop("source", None)

    data["outline"] = outline
    data["generated_sections"] = generated_sections
    stats = data.get("stats") if isinstance(data.get("stats"), dict) else {}
    stats["outline_count"] = _count_outline_nodes(outline)
    stats["generated_section_count"] = len([v for v in generated_sections.values() if str(v or "").strip()])
    data["stats"] = stats
    _refresh_render_decisions(data)
    output_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    if not isinstance(task.get("result"), dict):
        task["result"] = data
    else:
        task["result"]["outline"] = outline
        task["result"]["generated_sections"] = generated_sections
        task["result"]["render_decisions"] = data.get("render_decisions", [])
        task["result"]["stats"] = stats
    _save_task_meta(task)
    return {
        "ok": True,
        "outline": outline,
        "generated_sections": generated_sections,
        "render_decisions": data.get("render_decisions", []),
        "stats": stats,
    }


def _chapter_parent_id(chapter_id: str) -> str:
    value = str(chapter_id or "").strip()
    return value.rsplit(".", 1)[0] if "." in value else ""


def _tech_section_payload_tree(row: Any, rows_by_parent: Dict[str, list], company_id: str, required: bool) -> Dict[str, Any]:
    chapter_id = str(getattr(row, "chapter_id", "") or "")
    return {
        "title": str(getattr(row, "title", "") or getattr(row, "full_path", "") or chapter_id or "未命名章节"),
        "content": "",
        "required": required,
        "source": "tech_master_hook",
        "render_hook": {
            "type": "tech_section",
            "section_id": str(getattr(row, "id", "") or ""),
            "chapter_id": chapter_id,
            "copy_mode": "docx_block",
        },
        "children": [
            _tech_section_payload_tree(child, rows_by_parent, company_id, required)
            for child in rows_by_parent.get(chapter_id, [])
        ],
    }


@app.post("/api/projects/{project_id}/outline/tech-children")
async def append_tech_sections_as_outline_children(project_id: str, body: AppendTechSectionsRequest):
    from sqlalchemy import or_
    from sqlalchemy.orm import load_only
    from tender_agent.core.db import get_db_session
    from data.knowledge.models import TemplateSection

    section_ids = [str(item).strip() for item in body.section_ids or [] if str(item).strip()]
    if not section_ids:
        raise HTTPException(status_code=400, detail="No tech sections selected")

    db = next(get_db_session())
    try:
        all_rows = (
            db.query(TemplateSection)
            .options(
                load_only(
                    TemplateSection.id,
                    TemplateSection.company_id,
                    TemplateSection.scope,
                    TemplateSection.chapter_id,
                    TemplateSection.title,
                    TemplateSection.full_path,
                    TemplateSection.level,
                    TemplateSection.start_block_idx,
                )
            )
            .filter(
                TemplateSection.is_current == True,
                TemplateSection.deleted_at.is_(None),
                or_(TemplateSection.company_id == body.company_id, TemplateSection.scope == "shared"),
            )
            .order_by(TemplateSection.start_block_idx.asc())
            .all()
        )
        selected = [row for row in all_rows if str(row.id) in set(section_ids)]
    finally:
        db.close()

    if not selected:
        raise HTTPException(status_code=404, detail="Selected tech sections not found")

    selected_payloads: list[Dict[str, Any]] = []
    emitted_chapters: set[str] = set()
    for root in selected:
        root_chapter = str(getattr(root, "chapter_id", "") or "")
        descendants = [
            row for row in all_rows
            if root_chapter and str(getattr(row, "chapter_id", "") or "").startswith(f"{root_chapter}.")
        ]
        source_rows = descendants if descendants else [root]
        rows_by_parent: Dict[str, list] = {}
        for row in source_rows:
            chapter = str(getattr(row, "chapter_id", "") or "")
            rows_by_parent.setdefault(_chapter_parent_id(chapter), []).append(row)

        root_rows = (
            rows_by_parent.get(root_chapter, [])
            if descendants
            else [root]
        )
        for row in root_rows:
            chapter = str(getattr(row, "chapter_id", "") or "")
            if chapter in emitted_chapters:
                continue
            emitted_chapters.add(chapter)
            selected_payloads.append(_tech_section_payload_tree(row, rows_by_parent, body.company_id, required=True))

    return await append_outline_children(
        project_id,
        AppendOutlineChildrenRequest(parent_id=body.parent_id, sections=selected_payloads),
    )


@app.post("/api/projects/{project_id}/outline/from-text")
async def rebuild_outline_from_text(project_id: str, body: OutlineFromTextRequest):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if not _outline_available(task["status"]):
        raise HTTPException(status_code=400, detail="Task not completed")

    text = str(body.text or "").strip()
    if len(text) < 8:
        raise HTTPException(status_code=400, detail="Outline text is too short")

    output_file = OUTPUT_DIR / f"{project_id}.json"
    if not output_file.exists():
        raise HTTPException(status_code=404, detail="Project result file not found")

    try:
        outline, notes = await parse_user_outline_text(text)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        logger.warning("user outline parse failed for project {}: {}", project_id, str(exc)[:240])
        raise HTTPException(status_code=500, detail=f"Outline parse failed: {exc}") from exc

    data = json.loads(output_file.read_text(encoding="utf-8"))
    old_outline = data.get("outline", []) or []
    data["generated_sections"] = _rekey_generated_sections(
        old_outline,
        outline,
        data.get("generated_sections", {}) or {},
    )
    data["material_assignments"] = _rekey_material_assignments(
        old_outline,
        outline,
        data.get("material_assignments", []) or [],
    )
    data["outline"] = outline
    data["material_assignments"] = []
    data["rag_contexts"] = {}
    data["retrieval_summary"] = {}
    data["generated_sections"] = {}
    data["compliance_report"] = {}
    data["consistency_report"] = {}
    data["workflow_stage"] = "outline_review"
    data["outline_source"] = {
        "kind": "user_provided_outline",
        "updated_at": datetime.now().isoformat(),
        "text_preview": text[:1000],
        "notes": notes,
    }
    _refresh_render_decisions(data)
    output_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    task["result"] = data
    task["status"] = TaskStatus.OUTLINE_REVIEW
    task["message"] = "Outline updated from pasted text; confirm before material mapping"
    task["current_node"] = "outline_review"
    _save_task_meta(task)

    remap_result: dict[str, Any] | None = None
    remap_error = ""
    if body.remap:
        try:
            remap_result = await remap_materials(project_id)
            refreshed = _ensure_task_result(project_id, task)
            outline = refreshed.get("outline") or outline
        except Exception as exc:
            remap_error = str(exc)
            logger.warning("remap after user outline failed for project {}: {}", project_id, remap_error[:240])

    return {
        "ok": True,
        "outline": outline,
        "notes": notes,
        "remap": remap_result,
        "remap_error": remap_error,
    }


@app.post("/api/projects/{project_id}/render")
async def render_docx(project_id: str):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if task["status"] != TaskStatus.DONE:
        raise HTTPException(status_code=400, detail="Task not completed")

    if not TEMPLATE_PATH.exists():
        return {"download_url": f"/api/projects/{project_id}/download"}

    from tender_agent.rendering import render_blank_bid
    from tender_agent.core.db import get_db_session

    output_path = OUTPUT_DIR / f"{project_id}_blank_bid.docx"
    result = _ensure_task_result(project_id, task)

    tech_master_path = _company_tech_master_path(str(task.get("company_id") or DEFAULT_COMPANY_ID))
    logger.info("Resolved tech master path: {}", tech_master_path)

    db = next(get_db_session())
    try:
        tender_doc_path = result.get("source_file_path") or task.get("source_file_path")
        if not tender_doc_path:
            fallback_docx = UPLOAD_DIR / f"{project_id}.docx"
            if fallback_docx.exists():
                tender_doc_path = str(fallback_docx)
        material_assignments = result.setdefault("material_assignments", [])
        render_blank_bid(
            outline=result["outline"],
            title_info=result["title_info"],
            master_template_path=str(TEMPLATE_PATH),
            output_path=str(output_path),
            company_name=task["company_name"],
            material_assignments=material_assignments,
            generated_sections=result.get("generated_sections", {}),
            db_session=db,
            tech_master_path=str(tech_master_path),
            tender_doc_path=tender_doc_path,
            company_id=str(task.get("company_id") or DEFAULT_COMPANY_ID),
        )
        result["material_assignments"] = result.get("material_assignments", [])
        try:
            from tender_agent.understanding.compliance_checker import run_compliance_checks

            refreshed_compliance = run_compliance_checks(result)
            result["compliance_report"] = refreshed_compliance.get("compliance_report", refreshed_compliance)
        except Exception as exc:
            logger.warning("[render] refresh compliance report failed: {}", str(exc)[:160])
        _refresh_render_decisions(result)
        output_file = OUTPUT_DIR / f"{project_id}.json"
        if output_file.exists():
            data = json.loads(output_file.read_text(encoding="utf-8"))
            data["material_assignments"] = result["material_assignments"]
            data["render_decisions"] = result.get("render_decisions", [])
            data["compliance_report"] = result.get("compliance_report", data.get("compliance_report", {}))
            output_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
            result.update(data)
        task["result"] = result
    finally:
        db.close()
    task["docx_path"] = str(output_path)
    _save_task_meta(task)
    return {"download_url": f"/api/projects/{project_id}/download"}


@app.get("/api/projects/{project_id}/download")
async def download_file(project_id: str):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")

    docx_path = task.get("docx_path")
    if docx_path and Path(docx_path).exists():
        title = task.get("result", {}).get("title_info", {}).get("title", "投标书")
        return FileResponse(
            path=docx_path,
            filename=f"{title}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    md_file = OUTPUT_DIR / f"{project_id}.md"
    if md_file.exists():
        title = task.get("result", {}).get("title_info", {}).get("title", "投标书")
        return FileResponse(
            path=md_file,
            filename=f"{title}.md",
            media_type="text/markdown",
        )

    raise HTTPException(status_code=404, detail="File not found")



@app.get("/api/projects/{project_id}/source", summary="Download source file")
async def download_source_file(project_id: str):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")

    result = task.get("result") or {}
    source_file_path = _source_path_for_preview(project_id, result, task)
    if not source_file_path:
        raise HTTPException(status_code=404, detail="Source file not found")
    path = Path(source_file_path)
    if not path.exists():
        raise HTTPException(status_code=404, detail="Source file not found")

    suffix = path.suffix.lower()
    media = "application/octet-stream"
    if suffix == ".docx":
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif suffix == ".doc":
        media = "application/msword"
    elif suffix == ".pdf":
        media = "application/pdf"

    return FileResponse(path=str(path), filename=path.name, media_type=media)


@app.get("/api/projects/{project_id}/source-preview", summary="Preview source file as PDF")
async def preview_source_file(project_id: str):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")

    result = _ensure_task_result(project_id, task)
    source_file_path = _source_path_for_preview(project_id, result, task)
    if not source_file_path:
        raise HTTPException(status_code=404, detail="Source file not found")
    path = Path(source_file_path)
    if not path.exists():
        raise HTTPException(status_code=404, detail="Source file not found")

    try:
        preview_path = _ensure_source_preview_pdf(project_id, path)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Source preview unavailable: {exc}") from exc

    return FileResponse(
        path=str(preview_path),
        filename=preview_path.name,
        media_type="application/pdf",
        content_disposition_type="inline",
    )



@app.post("/api/projects/{project_id}/source-ask", response_model=SourceAskResponse, summary="Ask questions about the source tender document")
async def ask_source_document(project_id: str, body: SourceAskRequest):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if not _outline_available(task["status"]):
        raise HTTPException(status_code=400, detail="Task not completed")
    question = str(body.question or "").strip()
    if not question:
        raise HTTPException(status_code=400, detail="Question is required")

    result = _ensure_task_result(project_id, task)
    snippets = _source_ask_rank_snippets(question, result, top_k=body.top_k)
    citations = _source_ask_citations_from_snippets(snippets)
    _source_ask_resolve_citation_pages(project_id, result, citations)
    if not citations:
        return SourceAskResponse(answer="没有在当前项目的原文索引中检索到相关内容。", citations=[], confidence="low", used_llm=False)

    try:
        llm_result: _SourceAskLLMResponse = await llm_gateway.async_call_structured(
            _source_ask_prompt(question, citations),
            _SourceAskLLMResponse,
            max_tokens=1200,
        )
        allowed = {item.id for item in citations}
        selected_ids = [item for item in (llm_result.citation_ids or []) if item in allowed]
        selected = [item for item in citations if item.id in selected_ids] or citations[:3]
        answer = str(llm_result.answer or "").strip() or _source_ask_fallback_answer(question, selected)
        confidence = str(llm_result.confidence or "low").lower()
        if confidence not in {"low", "medium", "high"}:
            confidence = "low"
        return SourceAskResponse(answer=answer, citations=selected, confidence=confidence, used_llm=True)
    except Exception as exc:
        logger.warning("source ask LLM failed for project {}: {}", project_id, exc)
        return SourceAskResponse(
            answer=_source_ask_fallback_answer(question, citations),
            citations=citations[:5],
            confidence="low",
            used_llm=False,
        )

@app.post("/api/projects/{project_id}/source-locate", summary="Locate source anchors to preview pages")
async def locate_source_anchor_pages(project_id: str, body: SourceLocateRequest):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")

    result = _ensure_task_result(project_id, task)
    source_file_path = _source_path_for_preview(project_id, result, task)
    if not source_file_path:
        raise HTTPException(status_code=404, detail="Source file not found")
    try:
        locations = _resolve_anchor_locations(project_id, result, body.anchors or [])
        query_locations = _resolve_query_locations(project_id, result, body.query_items or [])
        preview_path = _ensure_source_preview_pdf(project_id, Path(source_file_path))
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Source file not found")
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Source locate unavailable: {exc}") from exc

    return {
        "preview_url": f"/api/projects/{project_id}/source-preview",
        "locations": locations,
        "query_locations": query_locations,
        "preview_file": Path(preview_path).name,
    }

# ============== 模板管理 API ==============

TEMPLATES_DIR = Path("data/templates")
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)


@app.get("/api/templates", summary="List templates")
async def list_templates():
    templates = []
    for f in TEMPLATES_DIR.glob("*.json"):
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            meta = data.get("meta", {})
            outline = data.get("outline", [])
            templates.append({
                "id": f.stem,
                "name": meta.get("name", f.stem),
                "category": meta.get("category", "未分类"),
                "chapter_count": len(outline),
                "created_at": meta.get("created_at", f.stat().st_mtime),
            })
        except Exception as e:
            logger.warning(f"Read template failed {f.name}: {e}")
    return templates


@app.post("/api/templates", summary="Save template")
async def save_template(body: dict):
    name = body.get("name", "").strip()
    outline = body.get("outline", [])
    category = body.get("category", "未分类")
    if not name:
        raise HTTPException(status_code=400, detail="Template name cannot be empty")
    template_id = name.replace(" ", "_").replace("/", "_")
    data = {
        "meta": {
            "name": name,
            "category": category,
            "created_at": datetime.now().isoformat(),
        },
        "outline": outline,
    }
    (TEMPLATES_DIR / f"{template_id}.json").write_text(
        json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return {"ok": True, "id": template_id}


@app.delete("/api/templates/{template_id}", summary="Delete template")
async def delete_template(template_id: str):
    path = TEMPLATES_DIR / f"{template_id}.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail="模板不存在")
    path.unlink()
    return {"ok": True}


@app.get("/api/templates/{template_id}", summary="Get template detail")
async def get_template(template_id: str):
    path = TEMPLATES_DIR / f"{template_id}.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail="模板不存在")
    return json.loads(path.read_text(encoding="utf-8"))


@app.post("/api/projects/{project_id}/apply-template", summary="Apply template")
async def apply_template(project_id: str, body: dict):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    template_id = body.get("template_id", "").strip()
    path = TEMPLATES_DIR / f"{template_id}.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail="模板不存在")
    template_data = json.loads(path.read_text(encoding="utf-8"))
    outline = normalize_outline_numbering(template_data.get("outline", []))
    result = _ensure_task_result(project_id, task)
    old_outline = result.get("outline", []) or []
    result["generated_sections"] = _rekey_generated_sections(
        old_outline,
        outline,
        result.get("generated_sections", {}) or {},
    )
    result["material_assignments"] = _rekey_material_assignments(
        old_outline,
        outline,
        result.get("material_assignments", []) or [],
    )
    result["outline"] = outline
    _refresh_render_decisions(result)
    output_file = OUTPUT_DIR / f"{project_id}.json"
    if output_file.exists():
        data = json.loads(output_file.read_text(encoding="utf-8"))
        data["generated_sections"] = result["generated_sections"]
        data["material_assignments"] = result["material_assignments"]
        data["outline"] = outline
        data["render_decisions"] = result.get("render_decisions", [])
        output_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    task["result"] = result
    _save_task_meta(task)
    return {"ok": True, "outline": outline}

@app.put("/api/projects/{project_id}/material-assignments", summary="Update material assignments")
async def update_material_assignments(project_id: str, body: dict):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if task["status"] != TaskStatus.DONE:
        raise HTTPException(status_code=400, detail="Task not completed")

    assignments = body.get("material_assignments", [])
    if not isinstance(assignments, list):
        raise HTTPException(status_code=400, detail="material_assignments 必须是 list")

    output_file = OUTPUT_DIR / f"{project_id}.json"
    if not output_file.exists():
        raise HTTPException(status_code=404, detail="Project result file not found")
    import json as _json
    data = _json.loads(output_file.read_text(encoding="utf-8"))
    data["material_assignments"] = assignments
    _refresh_render_decisions(data)
    output_file.write_text(_json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    if not isinstance(task.get("result"), dict):
        task["result"] = data
    else:
        task["result"]["material_assignments"] = assignments
    _save_task_meta(task)
    return {"ok": True}


@app.put("/api/projects/{project_id}/generated-sections", summary="Update generated section drafts")
async def update_generated_sections(project_id: str, body: dict):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if task["status"] != TaskStatus.DONE:
        raise HTTPException(status_code=400, detail="Task not completed")

    updates = body.get("generated_sections", {})
    if not isinstance(updates, dict):
        raise HTTPException(status_code=400, detail="generated_sections 必须是 object")

    result = _ensure_task_result(project_id, task)
    generated_sections = result.get("generated_sections", {}) or {}
    if not isinstance(generated_sections, dict):
        generated_sections = {}

    for node_id, content in updates.items():
        key = str(node_id)
        text = str(content or "").strip()
        if text:
            generated_sections[key] = text
        else:
            generated_sections.pop(key, None)

    result["generated_sections"] = generated_sections
    _refresh_render_decisions(result)
    output_file = OUTPUT_DIR / f"{project_id}.json"
    if output_file.exists():
        data = json.loads(output_file.read_text(encoding="utf-8"))
        data["generated_sections"] = generated_sections
        data["render_decisions"] = result.get("render_decisions", [])
        output_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    task["result"] = result
    _save_task_meta(task)
    return {"ok": True, "generated_sections_count": len(generated_sections)}


@_trace_task
async def _run_material_remap(project_id: str, trace_phase: str = "material_remap") -> None:
    task = _get_task_or_restore(project_id)
    if not task:
        return

    try:
        result = _ensure_task_result(project_id, task)
        outline = result.get("outline") or []
        existing_generated_sections = result.get("generated_sections", {}) or {}
        company_id = str(task.get("company_id") or result.get("company_id") or DEFAULT_COMPANY_ID)
        company_name = str(task.get("company_name") or result.get("company_name") or "")
        state = {
            "final_outline": outline,
            "outline": outline,
            "company_id": company_id,
            "company_name": company_name,
            "title_info": result.get("title_info") or {},
            "located_sections": result.get("located_sections") or (result.get("stats") or {}).get("located_sections") or [],
            "block_index": result.get("block_index") or [],
            "tender_requirements": result.get("tender_requirements") or {},
            "stats": result.get("stats") or {},
            "existing_material_assignments": result.get("material_assignments") or [],
            "source_file_path": result.get("source_file_path") or "",
        }

        task.update(progress=82, message="正在匹配知识库素材", current_node="material_mapper", error=None)
        _save_task_meta(task)
        mapped = await _map_materials_stage(state)
        assignments = mapped.get("material_assignments", []) or []
        warnings = mapped.get("warnings", []) or []
        state.update(mapped)

        task.update(progress=88, message="正在整理素材上下文", current_node="rag_retriever")
        _save_task_meta(task)
        rag_payload = await _build_rag_contexts_stage(state)
        state.update(rag_payload)

        task.update(progress=92, message="正在生成章节内容", current_node="content_generator")
        _save_task_meta(task)
        content_payload = await _generate_content_stage(state)
        state.update(content_payload)

        task.update(progress=97, message="正在核验章节内容", current_node="compliance_checker")
        _save_task_meta(task)
        compliance_payload = _run_compliance_stage(state)
        state.update(compliance_payload)
        consistency_payload = _check_consistency_stage(state)

        warnings.extend(rag_payload.get("warnings", []) or [])
        warnings.extend(content_payload.get("warnings", []) or [])

        result["company_id"] = company_id
        result["company_name"] = company_name
        result["material_assignments"] = assignments
        result["rag_contexts"] = rag_payload.get("rag_contexts", {})
        result["retrieval_summary"] = rag_payload.get("retrieval_summary", {})
        result["generated_sections"] = _merge_generated_sections_for_outline(
            outline,
            existing_generated_sections,
            content_payload.get("generated_sections", {}) or {},
        )
        result["project_facts"] = content_payload.get("project_facts", {})
        result["compliance_report"] = compliance_payload.get("compliance_report", {})
        result["consistency_report"] = consistency_payload.get("consistency_report", {})
        result["workflow_stage"] = "material_ready"
        _refresh_render_decisions(result)
        if warnings:
            result["warnings"] = (result.get("warnings") or []) + warnings

        task["result"] = result
        task["status"] = TaskStatus.DONE
        task["progress"] = 100
        task["message"] = "素材匹配与章节生成完成"
        task["current_node"] = "done"
        _save_task_meta(task)

        output_file = OUTPUT_DIR / f"{project_id}.json"
        if output_file.exists():
            data = json.loads(output_file.read_text(encoding="utf-8"))
            for key in (
                "company_id",
                "company_name",
                "material_assignments",
                "rag_contexts",
                "retrieval_summary",
                "generated_sections",
                "project_facts",
                "compliance_report",
                "consistency_report",
                "workflow_stage",
                "render_decisions",
            ):
                data[key] = result.get(key)
            if warnings:
                data["warnings"] = (data.get("warnings") or []) + warnings
            output_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    except asyncio.CancelledError:
        task.update(
            status=TaskStatus.OUTLINE_REVIEW,
            progress=80,
            message="素材匹配已取消，可重新确认目录",
            current_node="remap_cancelled",
        )
        _save_task_meta(task)
        raise
    except Exception as exc:
        logger.exception(f"Material remap failed for project {project_id}: {exc}")
        task.update(
            status=TaskStatus.OUTLINE_REVIEW,
            progress=80,
            message="素材匹配失败，可重新确认目录后重试",
            error=str(exc),
            current_node="remap_failed",
        )
        _save_task_meta(task)


@app.post("/api/projects/{project_id}/remap-materials", summary="Re-run material mapping")
async def remap_materials(project_id: str):
    task = _get_task_or_restore(project_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if not _outline_available(task["status"]):
        raise HTTPException(status_code=400, detail="Task not completed")

    result = _ensure_task_result(project_id, task)
    if not (result.get("outline") or []):
        raise HTTPException(status_code=400, detail="Outline is empty")

    running_task = background_tasks.get(project_id)
    if running_task and not running_task.done():
        return {"ok": True, "accepted": True, "already_running": True}

    task.update(progress=80, message="准备匹配素材", current_node="material_mapper", error=None)
    _save_task_meta(task)
    background_task = asyncio.create_task(_run_material_remap(project_id, trace_phase="material_remap"))
    background_tasks[project_id] = background_task
    background_task.add_done_callback(
        lambda completed, current_project_id=project_id: background_tasks.pop(current_project_id, None)
    )
    return {"ok": True, "accepted": True, "already_running": False}


@app.post("/api/knowledge/certificates/import-docx", summary="Import certificate images from DOCX headings")
async def import_certificate_docx(
    file: UploadFile = File(...),
    company_id: str = Form(DEFAULT_COMPANY_ID),
    scope: str = Form("company"),
):
    try:
        from tender_agent.knowledge.certificate_importer import import_certificate_images_from_docx
    except ModuleNotFoundError as exc:
        if exc.name == "docx":
            raise HTTPException(status_code=500, detail="当前 Python 环境缺少 python-docx，请先安装到 business-agent 环境") from exc
        raise

    if not (file.filename or "").lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="证书合集批量导入仅支持 .docx")
    target_dir = Path("data/knowledge/shared/imports") if scope == "shared" else _company_knowledge_dir(company_id) / "imports"
    target_dir.mkdir(parents=True, exist_ok=True)
    save_path = target_dir / _safe_upload_name(file.filename or "certificates.docx")
    content = await file.read()
    save_path.write_bytes(content)
    result = import_certificate_images_from_docx(
        save_path,
        company_id=company_id or DEFAULT_COMPANY_ID,
        scope="shared" if scope == "shared" else "company",
    )
    result["file_path"] = str(save_path)
    return result

@app.get("/api/knowledge/summary", summary="Knowledge base summary")
async def knowledge_summary(company_id: str = DEFAULT_COMPANY_ID):
    from sqlalchemy import func
    from sqlalchemy import or_
    from sqlalchemy import text
    from tender_agent.core.db import get_db_session
    from data.knowledge.models import Certificate, TemplateSection

    db = next(get_db_session())
    try:
        cert_total = db.query(func.count(Certificate.id)).filter(
            Certificate.is_current == True,
            Certificate.deleted_at.is_(None),
            or_(Certificate.company_id == company_id, Certificate.scope == "shared"),
        ).scalar() or 0
        tech_total = db.query(func.count(TemplateSection.id)).filter(
            TemplateSection.is_current == True,
            TemplateSection.deleted_at.is_(None),
            or_(TemplateSection.company_id == company_id, TemplateSection.scope == "shared"),
        ).scalar() or 0
        category_rows = (
            db.query(Certificate.category, func.count(Certificate.id))
            .filter(
                Certificate.is_current == True,
                Certificate.deleted_at.is_(None),
                or_(Certificate.company_id == company_id, Certificate.scope == "shared"),
            )
            .group_by(Certificate.category)
            .order_by(func.count(Certificate.id).desc())
            .all()
        )
        vector_available = False
        vector_note = "pgvector extension/table not detected"
        try:
            ext = db.execute(text("select 1 from pg_extension where extname='vector'")).first()
            vector_available = bool(ext)
            if vector_available:
                vector_note = "pgvector extension enabled"
        except Exception as exc:
            vector_note = f"pgvector check failed: {str(exc)[:80]}"
        return {
            "certificates": int(cert_total),
            "tech_sections": int(tech_total),
            "categories": [{"name": name, "count": int(count)} for name, count in category_rows if name],
            "tech_master_exists": _company_tech_master_path(company_id).exists(),
            "tech_master_path": str(_company_tech_master_path(company_id)),
            "vector_available": vector_available,
            "vector_note": vector_note,
        }
    finally:
        db.close()


@app.get("/api/knowledge/certificates", summary="List certificates")
def list_certificates(
    category: str | None = None,
    q: str | None = None,
    limit: int = 200,
    company_id: str = DEFAULT_COMPANY_ID,
):
    from sqlalchemy import or_
    from tender_agent.core.db import get_db_session
    from data.knowledge.models import Certificate

    db = next(get_db_session())
    try:
        query = db.query(Certificate).filter(
            Certificate.is_current == True,
            Certificate.deleted_at.is_(None),
            or_(Certificate.company_id == company_id, Certificate.scope == "shared"),
        )
        if category:
            query = query.filter(Certificate.category == category)
        if q:
            pattern = f"%{q.strip()}%"
            query = query.filter(
                or_(
                    Certificate.category.ilike(pattern),
                    Certificate.subcategory.ilike(pattern),
                    Certificate.name.ilike(pattern),
                    Certificate.cert_number.ilike(pattern),
                    Certificate.issuer.ilike(pattern),
                )
            )
        rows = query.order_by(Certificate.updated_at.desc()).limit(max(1, min(limit, 1000))).all()
        return [
            {
                "id": str(r.id),
                "category": r.category,
                "subcategory": r.subcategory,
                "name": r.name,
                "cert_number": r.cert_number,
                "issuer": r.issuer,
                "file_path": r.file_path,
                "file_type": r.file_type,
                "updated_at": r.updated_at.isoformat() if r.updated_at else None,
                "expire_date": r.expire_date.isoformat() if r.expire_date else None,
            }
            for r in rows
        ]
    finally:
        db.close()


@app.get("/api/knowledge/certificates/{certificate_id}/file", summary="Preview certificate file")
async def preview_certificate_file(certificate_id: str):
    from tender_agent.core.db import get_db_session
    from data.knowledge.models import Certificate

    db = next(get_db_session())
    try:
        row = db.query(Certificate).filter(
            Certificate.id == certificate_id,
            Certificate.is_current == True,
            Certificate.deleted_at.is_(None),
        ).first()
        if not row or not row.file_path:
            raise HTTPException(status_code=404, detail="Certificate file not found")
        path = Path(str(row.file_path))
        if not path.is_absolute():
            path = Path.cwd() / path
        if not path.exists() or not path.is_file():
            raise HTTPException(status_code=404, detail="Certificate file not found")
        suffix = path.suffix.lower()
        media = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".webp": "image/webp",
            ".gif": "image/gif",
            ".pdf": "application/pdf",
        }.get(suffix, "application/octet-stream")
        return FileResponse(
            path=str(path),
            filename=path.name,
            media_type=media,
            content_disposition_type="inline",
        )
    finally:
        db.close()


def _template_section_master_path(row: Any, company_id: str = DEFAULT_COMPANY_ID) -> Path:
    meta = getattr(row, "metadata_info", None) or {}
    source_path = str(meta.get("source_file_path") or "").strip()
    if source_path:
        path = Path(source_path)
        if not path.is_absolute():
            path = Path.cwd() / path
        if path.exists():
            return path
    return _company_tech_master_path(str(getattr(row, "company_id", "") or company_id or DEFAULT_COMPANY_ID))


def _template_section_image_url(row: Any, image_index: int, company_id: str) -> str:
    return f"/api/knowledge/tech-sections/{getattr(row, 'id', '')}/images/{image_index}?company_id={company_id}"


def _block_has_unsupported_word_graphic(block: Any) -> bool:
    """Return True when HTML preview cannot represent a Word vector drawing block."""
    from docx.oxml.ns import qn

    if block.findall(".//" + qn("a:blip")):
        return False
    vector_tags = {
        "wsp",
        "wgp",
        "shape",
        "group",
        "rect",
        "roundrect",
        "oval",
        "line",
        "polyline",
        "curve",
        "arc",
        "txbxContent",
        "relIds",
    }
    return any(str(element.tag).rsplit("}", 1)[-1] in vector_tags for element in block.iter())


def _template_section_preview_cache_id(row: Any, master_path: Path) -> str:
    stat = master_path.stat()
    identity = "|".join(
        [
            "numbering-v2",
            str(master_path.resolve()),
            str(stat.st_size),
            str(stat.st_mtime_ns),
            str(getattr(row, "start_block_idx", "")),
            str(getattr(row, "end_block_idx", "")),
        ]
    )
    digest = hashlib.sha256(identity.encode("utf-8", errors="surrogatepass")).hexdigest()[:16]
    section_id = re.sub(r"[^A-Za-z0-9_-]+", "_", str(getattr(row, "id", "") or "section"))[:64]
    return f"tech_section_{section_id}_{digest}"


def _ensure_template_section_preview_pdf(row: Any, company_id: str = DEFAULT_COMPANY_ID) -> Path:
    """Render one technical section with Word/LibreOffice while preserving original shapes."""
    from docx import Document
    from tender_agent.knowledge.tender_template_preview import freeze_numbering_labels

    master_path = _template_section_master_path(row, company_id)
    if not master_path.exists():
        raise FileNotFoundError(str(master_path))
    if getattr(row, "start_block_idx", None) is None or getattr(row, "end_block_idx", None) is None:
        raise RuntimeError("Tech section block range is missing")

    cache_id = _template_section_preview_cache_id(row, master_path)
    preview_path = _source_preview_pdf_path(cache_id)
    if preview_path.exists():
        return preview_path

    section_key = re.sub(r"[^A-Za-z0-9_-]+", "_", str(getattr(row, "id", "") or "section"))[:64]
    for stale in OUTPUT_DIR.glob(f"tech_section_{section_key}_*_source_preview.*"):
        if stale != preview_path:
            stale.unlink(missing_ok=True)

    with tempfile.TemporaryDirectory() as tmpdir:
        section_docx = Path(tmpdir) / "section.docx"
        shutil.copy2(master_path, section_docx)
        doc = Document(str(section_docx))
        body = doc.element.body
        blocks = list(body.iterchildren())
        start = max(int(row.start_block_idx or 0) - 1, 0)
        end = min(int(row.end_block_idx or 0), len(blocks))
        if start >= end:
            raise RuntimeError("Tech section block range is empty")
        freeze_numbering_labels(doc, blocks[start:end])
        for index, block in enumerate(blocks):
            if not (start <= index < end) and not block.tag.endswith("}sectPr"):
                body.remove(block)
        doc.save(str(section_docx))
        return _ensure_source_preview_pdf(cache_id, section_docx)


def _inline_docx_html(element: Any, row: Any, company_id: str, image_counter: Dict[str, int], related_parts: Dict[str, Any]) -> str:
    from docx.oxml.ns import qn
    from tender_agent.knowledge.docx_ooxml import is_inside_non_body_text

    parts: list[str] = []
    for child in element.iter():
        tag = str(child.tag).rsplit("}", 1)[-1]
        if tag == "t" and child.text and not is_inside_non_body_text(child):
            parts.append(html.escape(child.text))
        elif tag == "tab":
            parts.append("&nbsp;&nbsp;")
        elif tag == "br":
            parts.append("<br />")
        elif tag == "blip":
            rid = child.get(qn("r:embed")) or child.get(qn("r:link"))
            if rid and rid in related_parts:
                idx = image_counter["value"]
                image_counter["value"] += 1
                url = _template_section_image_url(row, idx, company_id)
                parts.append(f'<img src="{html.escape(url)}" alt="{html.escape(str(getattr(row, "title", "") or ""))}" />')
    return "".join(parts).strip()


def _table_block_html(block: Any, row: Any, company_id: str, image_counter: Dict[str, int], related_parts: Dict[str, Any]) -> str:
    from docx.oxml.ns import qn

    rows: list[str] = []
    for tr in block.findall(".//" + qn("w:tr")):
        cells: list[str] = []
        for tc in tr.findall("./" + qn("w:tc")):
            cell_html = _inline_docx_html(tc, row, company_id, image_counter, related_parts)
            cells.append(f"<td>{cell_html or '&nbsp;'}</td>")
        if cells:
            rows.append(f"<tr>{''.join(cells)}</tr>")
    return f"<table><tbody>{''.join(rows)}</tbody></table>" if rows else ""


def _template_section_content_html(row: Any, company_id: str = DEFAULT_COMPANY_ID) -> str:
    from docx import Document

    master_path = _template_section_master_path(row, company_id)
    if not master_path.exists() or getattr(row, "start_block_idx", None) is None or getattr(row, "end_block_idx", None) is None:
        return ""

    doc = Document(str(master_path))
    blocks = list(doc.element.body.iterchildren())
    start = max(int(row.start_block_idx or 0) - 1, 0)
    end = min(int(row.end_block_idx or 0), len(blocks))
    image_counter = {"value": 0}
    html_blocks: list[str] = []
    related_parts = doc.part.related_parts
    for block in blocks[start:end]:
        tag = str(block.tag).rsplit("}", 1)[-1]
        if tag == "p":
            body = _inline_docx_html(block, row, company_id, image_counter, related_parts)
            if body:
                html_blocks.append(f"<p>{body}</p>")
        elif tag == "tbl":
            table_html = _table_block_html(block, row, company_id, image_counter, related_parts)
            if table_html:
                html_blocks.append(table_html)
    return "".join(html_blocks)


def _extract_template_section_preview(row: Any, company_id: str = DEFAULT_COMPANY_ID, max_chars: int = 6000) -> Dict[str, Any]:
    from tender_agent.knowledge.section_copier import extract_section_text_from_master

    master_path = _template_section_master_path(row, company_id)
    text = ""
    try:
        text = extract_section_text_from_master(row, str(master_path), max_chars=max_chars)
    except Exception as exc:
        logger.warning(
            "[knowledge] extract tech preview failed section_id={} error={}",
            getattr(row, "id", ""),
            str(exc)[:160],
        )

    char_count = len(re.sub(r"\s+", "", text))
    image_count = 0
    vector_graphic_count = 0
    table_count = 0
    image_urls: list[Dict[str, Any]] = []
    try:
        from docx import Document
        from docx.oxml.ns import qn

        if master_path.exists() and getattr(row, "start_block_idx", None) is not None and getattr(row, "end_block_idx", None) is not None:
            doc = Document(str(master_path))
            blocks = list(doc.element.body.iterchildren())
            start = max(int(row.start_block_idx or 0) - 1, 0)
            end = min(int(row.end_block_idx or 0), len(blocks))
            for block in blocks[start:end]:
                tag = str(block.tag).rsplit("}", 1)[-1]
                if tag == "tbl":
                    table_count += 1
                if _block_has_unsupported_word_graphic(block):
                    vector_graphic_count += 1
                for blip in block.findall(".//" + qn("a:blip")):
                    rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
                    if not rid or rid not in doc.part.related_parts:
                        continue
                    idx = image_count
                    image_count += 1
                    if idx < 12:
                        image_urls.append(
                            {
                                "index": idx,
                                "url": f"/api/knowledge/tech-sections/{getattr(row, 'id', '')}/images/{idx}?company_id={company_id}",
                            }
                        )
    except Exception as exc:
        logger.debug("[knowledge] tech preview stats skipped: {}", str(exc)[:120])

    content_html = ""
    try:
        content_html = _template_section_content_html(row, company_id)
    except Exception as exc:
        logger.warning(
            "[knowledge] render tech preview html failed section_id={} error={}",
            getattr(row, "id", ""),
            str(exc)[:160],
        )

    paragraphs = [part.strip() for part in re.split(r"\n+", text) if part.strip()]
    raster_image_count = max(int(getattr(row, "image_count", None) or 0), image_count)
    rendered_preview_url = ""
    if vector_graphic_count:
        rendered_preview_url = (
            f"/api/knowledge/tech-sections/{getattr(row, 'id', '')}/preview"
            f"?company_id={company_id}&render=pdf"
        )
    return {
        "id": str(getattr(row, "id", "")),
        "chapter_id": getattr(row, "chapter_id", None),
        "title": getattr(row, "title", None),
        "full_path": getattr(row, "full_path", None),
        "level": int(getattr(row, "level", None) or 0),
        "span": max(0, int(getattr(row, "end_block_idx", None) or 0) - int(getattr(row, "start_block_idx", None) or 0)),
        "master_path": str(master_path),
        "content": text[:max_chars],
        "content_preview": text[:1200],
        "content_html": content_html,
        "paragraphs": paragraphs[:12],
        "char_count": int(getattr(row, "char_count", None) or char_count or 0),
        "image_count": raster_image_count,
        "vector_graphic_count": vector_graphic_count,
        "visual_count": raster_image_count + vector_graphic_count,
        "table_count": int(getattr(row, "table_count", None) or table_count or 0),
        "image_urls": image_urls,
        "rendered_preview_url": rendered_preview_url,
    }


def _template_section_image_blob(row: Any, image_index: int, company_id: str = DEFAULT_COMPANY_ID) -> tuple[bytes, str] | None:
    from docx import Document
    from docx.oxml.ns import qn

    master_path = _template_section_master_path(row, company_id)
    if not master_path.exists() or getattr(row, "start_block_idx", None) is None or getattr(row, "end_block_idx", None) is None:
        return None

    doc = Document(str(master_path))
    blocks = list(doc.element.body.iterchildren())
    start = max(int(row.start_block_idx or 0) - 1, 0)
    end = min(int(row.end_block_idx or 0), len(blocks))
    seen = 0
    for block in blocks[start:end]:
        for blip in block.findall(".//" + qn("a:blip")):
            rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
            image_part = doc.part.related_parts.get(rid) if rid else None
            if image_part is None:
                continue
            if seen == image_index:
                content_type = getattr(image_part, "content_type", None) or "image/png"
                return bytes(image_part.blob), str(content_type)
            seen += 1
    return None


@app.get("/api/knowledge/tech-sections", summary="List tech sections")
def list_tech_sections(q: str | None = None, limit: int = 500, company_id: str = DEFAULT_COMPANY_ID):
    from sqlalchemy import or_
    from tender_agent.core.db import get_db_session
    from data.knowledge.models import TemplateSection

    db = next(get_db_session())
    try:
        query = (
            db.query(TemplateSection)
            .filter(
                TemplateSection.is_current == True,
                TemplateSection.deleted_at.is_(None),
                or_(TemplateSection.company_id == company_id, TemplateSection.scope == "shared"),
            )
        )
        if q:
            pattern = f"%{q.strip()}%"
            query = query.filter(
                or_(
                    TemplateSection.chapter_id.ilike(pattern),
                    TemplateSection.title.ilike(pattern),
                    TemplateSection.full_path.ilike(pattern),
                    TemplateSection.category.ilike(pattern),
                )
            )
        rows = query.order_by(TemplateSection.start_block_idx.asc()).limit(max(1, min(limit, 2000))).all()
        return [
            {
                "id": str(r.id),
                "chapter_id": r.chapter_id,
                "title": r.title,
                "full_path": r.full_path,
                "level": r.level,
                "category": r.category,
                "span": max(0, int(r.end_block_idx or 0) - int(r.start_block_idx or 0)),
                "char_count": r.char_count,
                "image_count": r.image_count,
                "vector_graphic_count": int((r.metadata_info or {}).get("vector_graphic_count") or 0),
                "visual_count": int(r.image_count or 0)
                + int((r.metadata_info or {}).get("vector_graphic_count") or 0),
                "table_count": r.table_count,
            }
            for r in rows
        ]
    finally:
        db.close()


@app.get("/api/knowledge/tech-sections/{section_id}/preview", summary="Preview tech section content")
def preview_tech_section(section_id: str, company_id: str = DEFAULT_COMPANY_ID, render: str = ""):
    from tender_agent.core.db import get_db_session
    from data.knowledge.models import TemplateSection

    db = next(get_db_session())
    try:
        row = db.query(TemplateSection).filter(
            TemplateSection.id == section_id,
            TemplateSection.is_current == True,
            TemplateSection.deleted_at.is_(None),
        ).first()
        if row is None:
            raise HTTPException(status_code=404, detail="Tech section not found")
        if render.lower() == "pdf":
            try:
                preview_path = _ensure_template_section_preview_pdf(row, company_id=company_id)
            except Exception as exc:
                logger.warning(
                    "[knowledge] render tech section PDF failed section_id={} error={}",
                    section_id,
                    str(exc)[:200],
                )
                raise HTTPException(status_code=500, detail=f"Tech section document preview unavailable: {exc}") from exc
            return FileResponse(
                path=str(preview_path),
                filename=f"tech-section-{row.chapter_id or section_id}.pdf",
                media_type="application/pdf",
                content_disposition_type="inline",
            )
        return _extract_template_section_preview(row, company_id=company_id)
    finally:
        db.close()


@app.get("/api/knowledge/tech-sections/{section_id}/images/{image_index}", summary="Preview tech section image")
def preview_tech_section_image(section_id: str, image_index: int, company_id: str = DEFAULT_COMPANY_ID):
    from tender_agent.core.db import get_db_session
    from data.knowledge.models import TemplateSection

    db = next(get_db_session())
    try:
        row = db.query(TemplateSection).filter(
            TemplateSection.id == section_id,
            TemplateSection.is_current == True,
            TemplateSection.deleted_at.is_(None),
        ).first()
        if row is None:
            raise HTTPException(status_code=404, detail="Tech section not found")
        blob = _template_section_image_blob(row, max(0, int(image_index or 0)), company_id=company_id)
        if not blob:
            raise HTTPException(status_code=404, detail="Tech section image not found")
        content, content_type = blob
        return Response(content=content, media_type=content_type)
    finally:
        db.close()


@app.post("/api/knowledge/certificates/upload", summary="Upload certificate material")
async def upload_certificate_material(
    file: UploadFile = File(...),
    category: str = Form(...),
    name: str = Form(...),
    company_id: str = Form(DEFAULT_COMPANY_ID),
    scope: str = Form("company"),
    subcategory: str = Form(""),
    cert_number: str = Form(""),
    expire_date: str = Form(""),
    issuer: str = Form(""),
):
    from tender_agent.core.db import get_db_session
    from data.knowledge.models import Certificate

    if not category.strip() or not name.strip():
        raise HTTPException(status_code=400, detail="分类和素材名称不能为空")
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in {".png", ".jpg", ".jpeg", ".webp", ".pdf"}:
        raise HTTPException(status_code=400, detail="仅支持 png/jpg/jpeg/webp/pdf")

    save_name = _safe_upload_name(file.filename or "certificate")
    cert_dir = (Path("data/knowledge/shared/certs") if scope == "shared" else _company_knowledge_dir(company_id) / "certs")
    cert_dir.mkdir(parents=True, exist_ok=True)
    save_path = cert_dir / save_name
    content = await file.read()
    save_path.write_bytes(content)

    parsed_expire = None
    if expire_date.strip():
        try:
            parsed_expire = datetime.fromisoformat(expire_date.strip()).date()
        except Exception:
            raise HTTPException(status_code=400, detail="有效期格式应为 YYYY-MM-DD")

    db = next(get_db_session())
    try:
        row = Certificate(
            company_id=company_id if scope != "shared" else None,
            scope="shared" if scope == "shared" else "company",
            category=category.strip(),
            subcategory=subcategory.strip() or None,
            name=name.strip(),
            file_path=str(save_path),
            file_type=suffix.lstrip("."),
            file_size=len(content),
            cert_number=cert_number.strip() or None,
            expire_date=parsed_expire,
            issuer=issuer.strip() or None,
            is_current=True,
        )
        db.add(row)
        db.commit()
        return {"ok": True, "id": str(row.id), "file_path": str(save_path)}
    finally:
        db.close()


@app.put("/api/knowledge/certificates/{certificate_id}", summary="Update certificate material")
async def update_certificate_material(certificate_id: str, body: dict):
    from tender_agent.core.db import get_db_session
    from data.knowledge.models import Certificate

    db = next(get_db_session())
    try:
        row = db.query(Certificate).filter(Certificate.id == certificate_id, Certificate.deleted_at.is_(None)).first()
        if row is None:
            raise HTTPException(status_code=404, detail="证书素材不存在")
        for field in ("category", "subcategory", "name", "cert_number", "issuer"):
            if field in body:
                value = str(body.get(field) or "").strip()
                setattr(row, field, value if field in {"category", "name"} else (value or None))
        if not row.category or not row.name:
            raise HTTPException(status_code=400, detail="分类和名称不能为空")
        if "expire_date" in body:
            value = str(body.get("expire_date") or "").strip()
            row.expire_date = datetime.fromisoformat(value).date() if value else None
        row.updated_at = datetime.now()
        db.commit()
        return {"ok": True}
    finally:
        db.close()


@app.delete("/api/knowledge/certificates/{certificate_id}", summary="Delete certificate material")
async def delete_certificate_material(certificate_id: str):
    from tender_agent.core.db import get_db_session
    from data.knowledge.models import Certificate

    db = next(get_db_session())
    try:
        row = db.query(Certificate).filter(Certificate.id == certificate_id, Certificate.deleted_at.is_(None)).first()
        if row is None:
            raise HTTPException(status_code=404, detail="证书素材不存在")
        row.deleted_at = datetime.now()
        row.is_current = False
        db.commit()
        logger.info("[knowledge] certificate deleted: id={}, name={}", certificate_id, row.name)
        return {"ok": True}
    except HTTPException:
        raise
    except Exception as exc:
        db.rollback()
        logger.exception("[knowledge] certificate delete failed: id={}", certificate_id)
        raise HTTPException(status_code=500, detail="证书素材删除失败，请查看后端日志") from exc
    finally:
        db.close()


@app.post("/api/knowledge/tech-master/upload", summary="Upload and scan technical master")
async def upload_tech_master(
    file: UploadFile = File(...),
    company_id: str = Form(DEFAULT_COMPANY_ID),
    scope: str = Form("company"),
):
    if not (file.filename or "").lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="技术母版仅支持 .docx")
    content = await file.read()
    target = (
        Path("data/knowledge/shared/master/技术文件.docx")
        if scope == "shared"
        else _company_knowledge_dir(company_id) / "master" / "技术文件.docx"
    )
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_bytes(content)
    count = _scan_template_sections_from_docx(
        target,
        category="technical_master",
        replace_current=True,
        company_id=company_id if scope != "shared" else "",
        scope="shared" if scope == "shared" else "company",
    )
    return {"ok": True, "section_count": count, "file_path": str(target)}


@app.put("/api/knowledge/tech-sections/{section_id}", summary="Update tech section")
async def update_tech_section(section_id: str, body: dict):
    from tender_agent.core.db import get_db_session
    from data.knowledge.models import TemplateSection

    db = next(get_db_session())
    try:
        row = db.query(TemplateSection).filter(TemplateSection.id == section_id, TemplateSection.deleted_at.is_(None)).first()
        if row is None:
            raise HTTPException(status_code=404, detail="技术章节不存在")
        for field in ("chapter_id", "title", "full_path", "category", "display_name"):
            if field in body:
                value = str(body.get(field) or "").strip()
                setattr(row, field, value if field == "title" else (value or None))
        if "keywords" in body:
            keywords = body.get("keywords")
            if isinstance(keywords, str):
                row.keywords = [x.strip() for x in re.split(r"[,，\n]", keywords) if x.strip()]
            elif isinstance(keywords, list):
                row.keywords = [str(x).strip() for x in keywords if str(x).strip()]
        if not row.title:
            raise HTTPException(status_code=400, detail="章节标题不能为空")
        row.updated_at = datetime.now()
        db.commit()
        return {"ok": True}
    finally:
        db.close()


def _soft_delete_tech_sections(db: Any, section_ids: list[str]) -> int:
    from data.knowledge.models import TemplateSection

    normalized_ids = {str(value).strip() for value in section_ids if str(value).strip()}
    if not normalized_ids:
        raise HTTPException(status_code=400, detail="请选择需要删除的技术章节")

    roots = (
        db.query(TemplateSection)
        .filter(
            TemplateSection.id.in_(normalized_ids),
            TemplateSection.deleted_at.is_(None),
        )
        .all()
    )
    if not roots:
        raise HTTPException(status_code=404, detail="技术章节不存在")

    active_rows = (
        db.query(TemplateSection)
        .filter(
            TemplateSection.is_current.is_(True),
            TemplateSection.deleted_at.is_(None),
        )
        .all()
    )

    def source_key(row: Any) -> tuple[str, str, str, str]:
        metadata = getattr(row, "metadata_info", None) or {}
        return (
            str(getattr(row, "company_id", None) or ""),
            str(getattr(row, "scope", None) or ""),
            str(getattr(row, "category", None) or ""),
            str(metadata.get("source_file_path") or ""),
        )

    rows_by_id: dict[str, Any] = {}
    for root in roots:
        root_chapter = str(getattr(root, "chapter_id", None) or "").strip()
        root_key = source_key(root)
        for row in active_rows:
            row_chapter = str(getattr(row, "chapter_id", None) or "").strip()
            is_descendant = bool(root_chapter and row_chapter.startswith(f"{root_chapter}."))
            if source_key(row) == root_key and (str(row.id) == str(root.id) or is_descendant):
                rows_by_id[str(row.id)] = row

    deleted_at = datetime.now()
    for row in rows_by_id.values():
        row.deleted_at = deleted_at
        row.is_current = False
    db.commit()
    return len(rows_by_id)


@app.post("/api/knowledge/tech-sections/bulk-delete", summary="Delete multiple tech sections")
async def bulk_delete_tech_sections(body: TechSectionBulkDeleteRequest):
    from tender_agent.core.db import get_db_session

    db = next(get_db_session())
    try:
        return {"ok": True, "deleted_count": _soft_delete_tech_sections(db, body.section_ids)}
    finally:
        db.close()


@app.delete("/api/knowledge/tech-sections/{section_id}", summary="Delete tech section")
async def delete_tech_section(section_id: str):
    from tender_agent.core.db import get_db_session

    db = next(get_db_session())
    try:
        return {"ok": True, "deleted_count": _soft_delete_tech_sections(db, [section_id])}
    finally:
        db.close()


@app.post("/api/knowledge/history/import", summary="Import historical bid document sections")
async def import_history_bid(file: UploadFile = File(...), company_id: str = Form(DEFAULT_COMPANY_ID)):
    if not (file.filename or "").lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="历史反向入库当前仅支持 .docx")
    history_dir = _company_knowledge_dir(company_id) / "history"
    history_dir.mkdir(parents=True, exist_ok=True)
    save_path = history_dir / _safe_upload_name(file.filename or "history.docx")
    content = await file.read()
    save_path.write_bytes(content)
    count = _scan_template_sections_from_docx(
        save_path,
        category="history_bid",
        replace_current=False,
        company_id=company_id,
    )
    return {"ok": True, "section_count": count, "file_path": str(save_path)}
