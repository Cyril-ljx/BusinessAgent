"""
DOCX parser: section tree extraction + paragraph/table anchor index.
"""

from dataclasses import dataclass, field
import os
from pathlib import Path
import re
import subprocess
import sys
import tempfile
from contextlib import redirect_stderr, redirect_stdout
from io import StringIO
from typing import Any, Iterator

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table


@dataclass
class Section:
    id: str
    title: str
    level: int
    content: str = ""
    page_no: int | None = None
    has_table: bool = False
    start_item_idx: int | None = None
    end_item_idx: int | None = None
    children: list["Section"] = field(default_factory=list)


@dataclass
class ParsedDoc:
    file_name: str
    full_text: str
    sections: list[Section]
    flat_sections: list[Section]
    block_index: list[dict] = field(default_factory=list)


def parse_docx(file_path: str | Path) -> ParsedDoc:
    file_path = Path(file_path)
    try:
        doc = Document(str(file_path))
    except PackageNotFoundError as e:
        raise ValueError(f"文档解析失败: 非标准docx或文件损坏: {e}")
    except Exception as e:
        raise ValueError(f"文档解析失败: {e}")

    items: list[dict] = []
    block_idx = 0
    for kind, item in _iter_blocks(doc):
        if kind == "paragraph":
            text = item.text.strip()
            if not text:
                continue
            items.append(
                {
                    "kind": "paragraph",
                    "text": text,
                    "level": _detect_heading_level(item),
                    "block_idx": block_idx,
                }
            )
            block_idx += 1
        elif kind == "table":
            items.append(
                {
                    "kind": "table",
                    "text": _table_to_text(item),
                    "level": 0,
                    "block_idx": block_idx,
                }
            )
            block_idx += 1

    return _parsed_from_items(file_path.name, items)


def parse_pdf(file_path: str | Path) -> ParsedDoc:
    """Parse PDFs with PyMuPDF, with optional OCR for scanned documents."""
    file_path = Path(file_path)
    try:
        import fitz
    except Exception as exc:
        raise ValueError(f"PDF解析依赖 PyMuPDF 不可用: {exc}") from exc

    items: list[dict] = []
    try:
        pdf = fitz.open(str(file_path))
        pdf_toc = []
        try:
            pdf_toc = pdf.get_toc(simple=True) or []
        except Exception:
            pdf_toc = []
        for page_idx, page in enumerate(pdf, start=1):
            _extract_pdf_tables(page, page_idx, items)
            text = page.get_text("text") or ""
            for line in text.splitlines():
                cleaned = line.strip()
                if not cleaned:
                    continue
                block_idx = len(items)
                items.append(
                    {
                        "kind": "paragraph",
                        "text": cleaned,
                        "level": _detect_text_heading_level(cleaned),
                        "block_idx": block_idx,
                        "page_no": page_idx,
                    }
                )
        items.extend(_pdf_bookmark_file_directory_items(pdf_toc, len(items)))
    except Exception as exc:
        raise ValueError(f"PDF解析失败: {exc}") from exc

    if not items:
        if os.getenv("ENABLE_PDF_OCR", "true").lower() in {"1", "true", "yes", "on"}:
            ocr_items = _parse_pdf_with_ocr(file_path)
            if ocr_items:
                return _parsed_from_items(file_path.name, ocr_items)
        # Do not break the whole project creation for scanned PDFs. Downstream
        # locator will report no useful sections, but the API can still return
        # a controlled failure/warning instead of crashing in parsing.
        return ParsedDoc(
            file_name=file_path.name,
            full_text="",
            sections=[],
            flat_sections=[],
            block_index=[],
        )

    return _parsed_from_items(file_path.name, items)


def _pdf_bookmark_file_directory_items(toc: list[list[Any]], start_idx: int) -> list[dict]:
    """Convert authoritative PDF bookmarks under response-file format into blocks.

    Many PDFs expose the clean bid-file structure only in the left bookmark
    tree. The text layer may instead contain a looser sentence such as
    "供应商认为有必要提供的其他资料", so we promote the bookmark tree as the first
    authority for file-composition extraction.
    """
    if not toc:
        return []
    best: list[tuple[int, str, int | None]] = []
    best_score = -1
    for idx, entry in enumerate(toc):
        if not isinstance(entry, (list, tuple)) or len(entry) < 2:
            continue
        try:
            level = int(entry[0])
        except Exception:
            continue
        title = _clean_pdf_bookmark_title(str(entry[1] or ""))
        if not _looks_like_file_format_bookmark(title):
            continue
        children: list[tuple[int, str, int | None]] = []
        for child in toc[idx + 1 :]:
            if not isinstance(child, (list, tuple)) or len(child) < 2:
                continue
            try:
                child_level = int(child[0])
            except Exception:
                continue
            if child_level <= level:
                break
            if child_level != level + 1:
                continue
            child_title = _clean_pdf_bookmark_title(str(child[1] or ""))
            if _looks_like_file_directory_bookmark_item(child_title):
                page_no = None
                if len(child) >= 3:
                    try:
                        page_no = int(child[2])
                    except Exception:
                        page_no = None
                children.append((child_level, child_title, page_no))
        score = len(children) * 10 + (20 if "响应文件格式" in title or "投标文件格式" in title else 0)
        if len(children) >= 3 and score > best_score:
            best = children
            best_score = score
    if len(best) < 3:
        return []

    items: list[dict] = [
        {
            "kind": "paragraph",
            "text": "响应文件目录",
            "level": 2,
            "block_idx": start_idx,
            "page_no": best[0][2],
            "pdf_bookmark": True,
        }
    ]
    seen: set[str] = set()
    for _, title, page_no in best:
        key = re.sub(r"\s+", "", _strip_pdf_bookmark_number(title))
        if not key or key in seen:
            continue
        seen.add(key)
        items.append(
            {
                "kind": "paragraph",
                "text": title,
                "level": 0,
                "block_idx": start_idx + len(items),
                "page_no": page_no,
                "pdf_bookmark": True,
            }
        )
    return items if len(items) >= 4 else []


def _clean_pdf_bookmark_title(title: str) -> str:
    value = re.sub(r"\s+", " ", str(title or "")).strip()
    value = re.sub(r"\s*\d+\s*$", "", value).strip()
    return value.strip("：:；;。,.，、")


def _strip_pdf_bookmark_number(title: str) -> str:
    value = _clean_pdf_bookmark_title(title)
    value = re.sub(r"^第[一二三四五六七八九十百零\d]+[章节部分]\s*", "", value)
    value = re.sub(r"^[一二三四五六七八九十]+[、.．]\s*", "", value)
    value = re.sub(r"^\d+(?:\.\d+)*[、.．]?\s*", "", value)
    return value.strip()


def _looks_like_file_format_bookmark(title: str) -> bool:
    compact = re.sub(r"\s+", "", str(title or ""))
    if not compact:
        return False
    if any(term in compact for term in ("评分", "评审", "合同", "目录")):
        return False
    return any(
        term in compact
        for term in (
            "响应文件格式",
            "投标文件格式",
            "报价文件格式",
            "应答文件格式",
            "申请文件格式",
            "响应文件组成",
            "投标文件组成",
        )
    )


def _looks_like_file_directory_bookmark_item(title: str) -> bool:
    name = _strip_pdf_bookmark_number(title)
    compact = re.sub(r"\s+", "", name)
    if not compact or len(compact) > 60:
        return False
    if any(term in compact for term in ("评审", "评分", "合同签订", "合同履行", "质疑", "投诉", "询问函")):
        return False
    return any(
        term in compact
        for term in (
            "函",
            "报价表",
            "证明",
            "授权",
            "资格审查",
            "响应方案",
            "投标方案",
            "其他资料",
            "其他材料",
            "其他文件",
            "承诺",
            "声明",
        )
    )


def _extract_pdf_tables(page: Any, page_idx: int, items: list[dict]) -> None:
    """Extract vector PDF tables with PyMuPDF.

    Consecutive table fragments with the same header are merged so a table that
    continues on the next page stays one logical block.
    """
    try:
        with redirect_stdout(StringIO()), redirect_stderr(StringIO()):
            finder = page.find_tables()
        tables = getattr(finder, "tables", []) or []
    except Exception:
        return

    for table in tables:
        try:
            rows = table.extract()
        except Exception:
            continue
        text = _rows_to_markdown(rows)
        if text:
            _append_table_item(items, text, page_idx)


def _rows_to_markdown(rows: list[list[Any]]) -> str:
    normalized: list[list[str]] = []
    for row in rows or []:
        cells = [str(cell or "").strip().replace("\n", " ") for cell in (row or [])]
        if any(cells):
            normalized.append(cells)
    if not normalized:
        return ""
    width = max(len(row) for row in normalized)
    normalized = [row + [""] * (width - len(row)) for row in normalized]
    lines = ["| " + " | ".join(normalized[0]) + " |"]
    lines.append("|" + "|".join(["---"] * width) + "|")
    for row in normalized[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _table_header_key(markdown: str) -> str:
    for line in (markdown or "").splitlines():
        stripped = re.sub(r"\s+", "", line)
        if stripped.startswith("|") and "---" not in stripped:
            return stripped
    return ""


def _append_table_item(items: list[dict], text: str, page_idx: int) -> None:
    header = _table_header_key(text)
    if header:
        for previous in reversed(items):
            if previous.get("kind") != "table":
                continue
            if previous.get("table_header") != header:
                continue
            previous_end = previous.get("page_no_end") or previous.get("page_no")
            if previous_end is not None and int(previous_end) < page_idx - 1:
                break
            lines = text.splitlines()
            body = "\n".join(line for line in lines[2:] if line.strip())
            if body:
                previous["text"] = f"{previous['text']}\n{body}"
                previous["page_no_end"] = page_idx
            return
    items.append(
        {
            "kind": "table",
            "text": text,
            "level": 0,
            "block_idx": len(items),
            "page_no": page_idx,
            "page_no_end": page_idx,
            "table_header": header,
        }
    )


def _parse_pdf_with_ocr(file_path: Path) -> list[dict]:
    """OCR fallback for scanned PDFs.

    PaddleOCR/PP-OCR is optional. If it is not installed or model loading fails,
    return an empty list so PDF parsing degrades gracefully.
    """
    try:
        import fitz
        from paddleocr import PaddleOCR
    except Exception:
        return []

    max_pages = int(os.getenv("PDF_OCR_MAX_PAGES", "20"))
    try:
        ocr = PaddleOCR(use_doc_orientation_classify=False, use_doc_unwarping=False, use_textline_orientation=False)
    except TypeError:
        try:
            ocr = PaddleOCR(use_angle_cls=True, lang="ch")
        except Exception:
            return []
    except Exception:
        return []

    items: list[dict] = []
    try:
        pdf = fitz.open(str(file_path))
        for page_idx, page in enumerate(pdf, start=1):
            if page_idx > max_pages:
                break
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image_bytes = pix.tobytes("png")
            lines = _run_paddle_ocr(ocr, image_bytes)
            for line in lines:
                cleaned = line.strip()
                if not cleaned:
                    continue
                items.append(
                    {
                        "kind": "paragraph",
                        "text": cleaned,
                        "level": _detect_text_heading_level(cleaned),
                        "block_idx": len(items),
                        "page_no": page_idx,
                        "ocr": True,
                    }
                )
    except Exception:
        return []
    return items


def _run_paddle_ocr(ocr: Any, image_bytes: bytes) -> list[str]:
    """Normalize PaddleOCR v2/v3/v5-style results to text lines."""
    tmp_path = ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            tmp.write(image_bytes)
            tmp_path = tmp.name
        result = ocr.ocr(tmp_path)
    except Exception:
        try:
            result = ocr.predict(tmp_path or image_bytes)
        except Exception:
            return []
    finally:
        if tmp_path:
            try:
                Path(tmp_path).unlink(missing_ok=True)
            except Exception:
                pass

    lines: list[str] = []

    def visit(obj: Any) -> None:
        if obj is None:
            return
        if isinstance(obj, str):
            if obj.strip():
                lines.append(obj.strip())
            return
        if isinstance(obj, dict):
            for key in ("text", "rec_text", "transcription"):
                value = obj.get(key)
                if isinstance(value, str) and value.strip():
                    lines.append(value.strip())
            for key in ("rec_texts", "texts"):
                value = obj.get(key)
                if isinstance(value, list):
                    for text in value:
                        visit(text)
            return
        if isinstance(obj, (list, tuple)):
            # Common v2 format: [box, (text, score)]
            if len(obj) == 2 and isinstance(obj[1], (list, tuple)) and obj[1] and isinstance(obj[1][0], str):
                visit(obj[1][0])
                return
            for item in obj:
                visit(item)

    visit(result)
    return lines


def parse_document(file_path: str | Path) -> ParsedDoc:
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(file_path)
    if suffix == ".doc":
        return parse_docx(_convert_doc_to_docx(file_path))
    return parse_docx(file_path)


def _convert_doc_to_docx(doc_path: Path) -> Path:
    """Convert legacy .doc to .docx for python-docx parsing."""
    doc_path = Path(doc_path)
    output_path = doc_path.with_suffix(".docx")
    if output_path.exists() and output_path.stat().st_mtime >= doc_path.stat().st_mtime:
        return output_path

    if sys.platform == "win32":
        try:
            import win32com.client

            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(str(doc_path.absolute()))
            doc.SaveAs(str(output_path.absolute()), FileFormat=16)
            doc.Close()
            word.Quit()
            if output_path.exists():
                return output_path
        except Exception:
            pass

    try:
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(doc_path.parent),
                str(doc_path),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=120,
        )
        if output_path.exists():
            return output_path
    except Exception:
        pass

    raise ValueError("无法解析 .doc 老格式：请安装 Word/LibreOffice，或先另存为 .docx 后重试。")


def _parsed_from_items(file_name: str, items: list[dict]) -> ParsedDoc:
    items = _normalize_heading_levels(items)
    sections = _build_sections(items)
    flat_sections = _flatten(sections)
    full_text = "\n".join(it["text"] for it in items if it.get("text"))
    return ParsedDoc(
        file_name=file_name,
        full_text=full_text,
        sections=sections,
        flat_sections=flat_sections,
        block_index=[
            {
                "anchor": f"p{it['block_idx']}",
                "kind": it["kind"],
                "text": it["text"][:240],
                "page_no": it.get("page_no"),
                "page_no_end": it.get("page_no_end"),
                "ocr": bool(it.get("ocr")),
                "pdf_bookmark": bool(it.get("pdf_bookmark")),
            }
            for it in items
            if it.get("block_idx") is not None
        ],
    )


def _iter_blocks(doc: Document) -> Iterator[tuple[str, object]]:
    from docx.oxml.ns import qn

    body = doc.element.body
    paragraphs = {p._element: p for p in doc.paragraphs}
    tables = {t._element: t for t in doc.tables}
    for child in body.iterchildren():
        if child.tag == qn("w:p") and child in paragraphs:
            yield ("paragraph", paragraphs[child])
        elif child.tag == qn("w:tbl") and child in tables:
            yield ("table", tables[child])


def _detect_heading_level(paragraph) -> int:
    text = paragraph.text.strip()
    if not text or len(text) > 80:
        return 0
    if len(re.sub(r"\s+", "", text)) > 45 and _is_sentence_like(text):
        return 0

    style_name = (paragraph.style.name or "").lower()
    if "heading 1" in style_name or "标题 1" in style_name:
        return 1
    if "heading 2" in style_name or "标题 2" in style_name:
        return 2
    if "heading 3" in style_name or "标题 3" in style_name:
        return 3

    if re.match(r"^第[一二三四五六七八九十百零\d]+[章节部分]", text):
        return 1
    if re.match(r"^第[一二三四五六七八九十百零\d]+节", text) or re.match(r"^附件[一二三四五六七八九十\d]+", text):
        return 2
    if re.match(r"^[一二三四五六七八九十]+、", text):
        return 2
    numbered_level = _numbered_heading_level(text)
    if numbered_level:
        return numbered_level
    if re.match(r"^[（(][一二三四五六七八九十\d]+[）)]", text) and not _is_sentence_like(text):
        return 3

    if all(run.bold for run in paragraph.runs if run.text.strip()):
        if 2 <= len(text) <= 30 and not _is_sentence_like(text):
            return 2

    return 0


def _detect_text_heading_level(text: str) -> int:
    if not text or len(text) > 80:
        return 0
    if len(re.sub(r"\s+", "", text)) > 45 and _is_sentence_like(text):
        return 0
    if re.match(r"^第[一二三四五六七八九十百零\d]+[章节部分]", text):
        return 1
    if re.match(r"^[一二三四五六七八九十]+、", text):
        return 2
    numbered_level = _numbered_heading_level(text)
    if numbered_level:
        return numbered_level
    if re.match(r"^[（(][一二三四五六七八九十\d]+[）)]", text) and not _is_sentence_like(text):
        return 3
    return 0


def _numbered_heading_level(text: str) -> int:
    path = _heading_number_path(text)
    if not path:
        return 0
    if len(path) == 1:
        return 2
    return min(len(path), 4)


def _heading_number_path(text: str) -> list[int]:
    if _is_sentence_like(text):
        return []
    match = re.match(r"^(\d+(?:\.\d+){0,5})[\.、\s]+", text or "")
    if not match:
        match = re.match(r"^(\d+(?:\.\d+){1,5})(?=[^\d])", text or "")
    if not match:
        return []
    try:
        return [int(part) for part in match.group(1).split(".") if part != ""]
    except ValueError:
        return []


def _normalize_heading_levels(items: list[dict]) -> list[dict]:
    """Use numbering continuity to smooth heading levels.

    This avoids treating a body sentence with a number as a heading, and keeps
    1.2 / 1.2.1 style levels stable even when Word styles are missing.
    """
    last_path_by_depth: dict[int, list[int]] = {}
    normalized: list[dict] = []
    for item in items:
        copy = dict(item)
        if copy.get("kind") == "paragraph" and copy.get("level", 0) > 0:
            path = _heading_number_path(str(copy.get("text", "")))
            if path:
                depth = len(path)
                previous = last_path_by_depth.get(depth)
                parent_ok = depth == 1 or (depth - 1) in last_path_by_depth
                continuous = previous is None or path[:-1] == previous[:-1] or parent_ok
                if continuous:
                    copy["level"] = 2 if depth == 1 else min(depth, 4)
                    last_path_by_depth[depth] = path
                    for stale_depth in [d for d in last_path_by_depth if d > depth]:
                        last_path_by_depth.pop(stale_depth, None)
        normalized.append(copy)
    return normalized


def _is_sentence_like(text: str) -> bool:
    normalized = re.sub(r"\s+", "", text or "")
    if not normalized:
        return True
    if normalized.endswith(("。", "；", ";", "，", ",", "：", ":")):
        return True
    body_markers = ("应", "须", "需", "必须", "不得", "提供", "报价人", "投标人", "供应商", "采购人", "负责", "同意")
    return len(normalized) >= 18 and any(marker in normalized for marker in body_markers)


_PREAMBLE_TITLE_MARKERS = (
    "投标文件所需资料",
    "响应文件内容一览表",
    "投标文件目录表",
    "索引目录表",
)


def _normalize_section_title(text: str) -> str:
    t = re.sub(r"\s+", "", text or "")
    t = re.sub(r"^第[一二三四五六七八九十百零\d]+[章节部分]", "", t)
    t = re.sub(r"^[一二三四五六七八九十]+[、.]", "", t)
    t = re.sub(r"^\d+(?:\.\d+)*[.、]?", "", t)
    t = re.sub(r"^[（(][一二三四五六七八九十\d]+[）)]", "", t)
    return t


def _is_repeat_heading(new_title: str, current_title: str) -> bool:
    """识别表格前重复出现的居中标题变体。"""
    a = _normalize_section_title(new_title)
    b = _normalize_section_title(current_title)
    if not a or not b:
        return False
    if len(a) < 4 or len(b) < 4:
        return a == b
    if a == b:
        return True
    return min(len(a), len(b)) >= 6 and (a in b or b in a)


def _build_sections(items: list[dict]) -> list[Section]:
    root_sections: list[Section] = []
    stack: list[Section] = []
    section_counter = 0
    sub_counter: dict[str, int] = {}
    preamble = _build_preamble_section(items)
    if preamble is not None:
        root_sections.append(preamble)

    for item in items:
        if item["kind"] == "paragraph" and item["level"] > 0:
            level = item["level"]
            title = item["text"]
            if stack and _is_repeat_heading(title, stack[-1].title):
                if stack[-1].content:
                    stack[-1].content += "\n"
                stack[-1].content += title
                stack[-1].end_item_idx = item.get("block_idx")
                continue

            while stack and stack[-1].level >= level:
                stack.pop()

            if stack:
                parent = stack[-1]
                sub_counter[parent.id] = sub_counter.get(parent.id, 0) + 1
                new_id = f"{parent.id}.{sub_counter[parent.id]}"
            else:
                section_counter += 1
                new_id = str(section_counter)

            sec = Section(id=new_id, title=title, level=level, start_item_idx=item.get("block_idx"))
            if stack:
                stack[-1].children.append(sec)
            else:
                root_sections.append(sec)
            stack.append(sec)
        else:
            text = item.get("text", "")
            if not text or not stack:
                continue
            if item["kind"] == "table":
                stack[-1].has_table = True
            if stack[-1].content:
                stack[-1].content += "\n"
            stack[-1].content += text
            stack[-1].end_item_idx = item.get("block_idx")

    return root_sections


def _build_preamble_section(items: list[dict]) -> Section | None:
    """仅在文档开头确实是投标资料清单时保留为前置章节。

    大多数招标文件以封面文本或目录开头，不应成为章节。少数文件会在第一个
    标题前放置权威的投标文件索引表，仅保留这种窄场景，避免污染其他文档。
    """
    first_heading_pos: int | None = None
    for pos, item in enumerate(items):
        if item.get("kind") == "paragraph" and int(item.get("level") or 0) > 0:
            first_heading_pos = pos
            break
    if first_heading_pos is None or first_heading_pos <= 0:
        return None

    preamble_items = items[:first_heading_pos]
    text = "\n".join(str(item.get("text") or "") for item in preamble_items if item.get("text")).strip()
    if not text:
        return None
    compact = re.sub(r"\s+", "", text)
    if not any(marker in compact for marker in _PREAMBLE_TITLE_MARKERS):
        return None
    has_category = "文件类型" in compact or "类型" in compact
    has_material = "文件名称" in compact or "证明材料" in compact
    has_order = "序号" in compact or "装订顺序" in compact
    if not (has_category and has_material and has_order):
        return None

    title = "文档开头"
    for item in preamble_items:
        if item.get("kind") != "paragraph":
            continue
        candidate = str(item.get("text") or "").strip()
        candidate_compact = re.sub(r"\s+", "", candidate)
        if any(marker in candidate_compact for marker in _PREAMBLE_TITLE_MARKERS):
            title = candidate
            break

    block_indices = [item.get("block_idx") for item in preamble_items if item.get("block_idx") is not None]
    return Section(
        id="preamble",
        title=title,
        level=1,
        content=text,
        has_table=any(item.get("kind") == "table" for item in preamble_items),
        start_item_idx=min(block_indices) if block_indices else None,
        end_item_idx=max(block_indices) if block_indices else None,
    )


def _flatten(sections: list[Section]) -> list[Section]:
    out: list[Section] = []
    for s in sections:
        out.append(s)
        out.extend(_flatten(s.children))
    return out


def _table_to_text(table: Table) -> str:
    lines: list[str] = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
        if i == 0:
            lines.append("|" + "|".join(["---"] * len(cells)) + "|")
    return "\n".join(lines)
